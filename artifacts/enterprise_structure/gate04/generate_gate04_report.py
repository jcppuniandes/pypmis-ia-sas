from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_ROW_HEIGHT_RULE, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(r"C:\Users\Ricardo\Documents\GitHub\pypmis-ia-sas")
TEMPLATE = Path(
    r"C:\Users\Ricardo\Documents\P&P\P&Pmis Construction AI\Diseño\Resumen de Sprint"
    r"\Informe_Tecnico_PPMIS_Core_Correction_Gate03_2026-08-10.docx"
)
OUTPUT = Path(
    r"C:\Users\Ricardo\Documents\P&P\P&Pmis Construction AI\Diseño\Resumen de Sprint"
    r"\Informe_Tecnico_PPMIS_Workspace_Revision_Manager_Gate04_2026-08-10.docx"
)
LOGO = REPO / "frontend" / "public" / "pypmis-construction-ai-logo.png"

NAVY = "0C2638"
BLUE = "17679A"
TEAL = "129C9D"
PALE_BLUE = "EAF3F7"
PALE_TEAL = "E8F6F3"
PALE_GREEN = "E5F4EA"
GREEN = "1C7C54"
PALE_AMBER = "FFF4DB"
AMBER = "A85D00"
PALE_RED = "FCE9E7"
RED = "A13B32"
LIGHT = "F4F7F9"
MID = "D7E1E6"
DARK = RGBColor(12, 38, 56)
MUTED = RGBColor(82, 104, 116)


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def set_cell_width(cell, width: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_run(run, *, size=None, bold=None, color=None, font=None, italic=None) -> None:
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color) if isinstance(color, str) else color
    if font is not None:
        run.font.name = font
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)


def remove_body(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.62)
    section.bottom_margin = Inches(0.58)
    section.left_margin = Inches(0.67)
    section.right_margin = Inches(0.67)
    section.header_distance = Inches(0.22)
    section.footer_distance = Inches(0.22)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos")
    normal.font.size = Pt(9.2)
    normal.font.color.rgb = DARK
    normal.paragraph_format.space_after = Pt(4.5)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size, color, before, after in (
        ("Heading 1", 17, NAVY, 10, 7),
        ("Heading 2", 12.2, BLUE, 8, 4),
        ("Heading 3", 10.3, TEAL, 6, 3),
    ):
        style = doc.styles[style_name]
        style.font.name = "Aptos Display"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Aptos Display")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header
    header.is_linked_to_previous = False
    header_p = header.paragraphs[0]
    header_p.clear()
    header_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header_p.add_run("P&Pmis Construction AI  |  Gate 04 · Workspace Revision Manager")
    set_run(r, size=7.8, bold=True, color=BLUE)
    p_pr = header_p._p.get_or_add_pPr()
    bottom_border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), TEAL)
    bottom_border.append(bottom)
    p_pr.append(bottom_border)

    first_header = section.first_page_header
    first_header.paragraphs[0].clear()

    footer = section.footer
    footer.is_linked_to_previous = False
    footer_p = footer.paragraphs[0]
    footer_p.clear()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer_p.add_run("P&Pmis · 10 de agosto de 2026  |  Página ")
    set_run(r, size=7.5, color="657B87")
    add_page_number(footer_p)

    first_footer = section.first_page_footer
    first_footer_p = first_footer.paragraphs[0]
    first_footer_p.clear()
    first_footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = first_footer_p.add_run("Documento técnico de configuración · Uso interno")
    set_run(r, size=7.5, color="657B87")


def add_rule(doc: Document, color=TEAL, size=16) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None, italic=False):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_run(r, bold=True, color=NAVY)
        r = p.add_run(text[len(bold_prefix) :])
        set_run(r, italic=italic)
    else:
        r = p.add_run(text)
        set_run(r, italic=italic)
    return p


def add_bullets(doc: Document, items: list[str], *, compact=True) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.18)
        p.paragraph_format.first_line_indent = Inches(-0.12)
        p.paragraph_format.space_after = Pt(2.2 if compact else 4)
        p.add_run(item)


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.left_indent = Inches(0.22)
        p.paragraph_format.first_line_indent = Inches(-0.15)
        p.paragraph_format.space_after = Pt(2.5)
        p.add_run(item)


def add_codebox(doc: Document, text: str, *, fill=LIGHT) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 100, 130, 100, 130)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    for index, line in enumerate(text.splitlines()):
        if index:
            p.add_run().add_break()
        r = p.add_run(line)
        set_run(r, size=8.1, font="Cascadia Mono", color=NAVY)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_callout(doc: Document, title: str, body: str, *, kind="info") -> None:
    fills = {"pass": (PALE_GREEN, GREEN), "warning": (PALE_AMBER, AMBER), "risk": (PALE_RED, RED), "info": (PALE_BLUE, BLUE)}
    fill, color = fills[kind]
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 115, 145, 115, 145)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title.upper())
    set_run(r, size=8.2, bold=True, color=color)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(body)
    set_run(r, size=9, color=NAVY)
    prevent_row_split(table.rows[0])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    *,
    widths: list[float] | None = None,
    font_size=8.0,
    first_col_bold=False,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        set_cell_margins(cell, 80, 90, 80, 90)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if widths:
            set_cell_width(cell, widths[i])
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(header)
        set_run(r, size=7.8, bold=True, color="FFFFFF")
    set_repeat_header(table.rows[0])
    prevent_row_split(table.rows[0])
    for ridx, row_data in enumerate(rows):
        row = table.add_row()
        if ridx % 2:
            for cell in row.cells:
                shade(cell, LIGHT)
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            set_cell_margins(cell, 68, 85, 68, 85)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            if widths:
                set_cell_width(cell, widths[i])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(str(value))
            set_run(r, size=font_size, bold=(first_col_bold and i == 0), color=NAVY)
        prevent_row_split(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_status_table(doc: Document, rows: list[tuple[str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.style = "Table Grid"
    for i, header in enumerate(("Control", "Estado", "Evidencia")):
        cell = table.rows[0].cells[i]
        shade(cell, NAVY)
        set_cell_margins(cell)
        set_cell_width(cell, (2.35, 1.05, 3.65)[i])
        r = cell.paragraphs[0].add_run(header)
        set_run(r, size=8, bold=True, color="FFFFFF")
    set_repeat_header(table.rows[0])
    for name, status, evidence in rows:
        row = table.add_row()
        for i, value in enumerate((name, status, evidence)):
            cell = row.cells[i]
            set_cell_margins(cell)
            set_cell_width(cell, (2.35, 1.05, 3.65)[i])
            if i == 1:
                shade(cell, PALE_GREEN if status == "PASS" else PALE_AMBER)
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(value)
            set_run(r, size=7.9, bold=(i in (0, 1)), color=GREEN if i == 1 and status == "PASS" else NAVY)
        prevent_row_split(row)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def section(doc: Document, number: str, title: str, *, page_break=False) -> None:
    h = doc.add_heading(f"{number}. {title}", level=1)
    h.paragraph_format.keep_with_next = True
    h.paragraph_format.page_break_before = number == "A"


def build_report() -> Document:
    doc = Document(TEMPLATE)
    remove_body(doc)
    configure_document(doc)
    props = doc.core_properties
    props.title = "Informe Técnico Gate 04 — Workspace Structure Revision Manager"
    props.subject = "P&Pmis Construction AI · Enterprise Structure"
    props.author = "P&P Ingeniería y Proyectos · Codex"
    props.keywords = "Gate 04, Enterprise Structure, Workspace Revision, ADMIN MODE, P&Pmis"
    props.comments = "Generado a partir de evidencia reproducible del repositorio y localhost."

    # Cover — memo masthead variant.
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(7)
    p.add_run().add_picture(str(LOGO), width=Inches(0.72))
    r = p.add_run("   P&Pmis Construction AI")
    set_run(r, size=15, bold=True, color=NAVY)
    add_rule(doc, TEAL, 22)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("INFORME TÉCNICO · GATE 04")
    set_run(r, size=10.5, bold=True, color=TEAL)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("Workspace Structure\nRevision Manager")
    set_run(r, size=28, bold=True, color=NAVY, font="Aptos Display")
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(15)
    r = p.add_run("Revisiones CORE gobernadas · edición DRAFT · validación · diff · aprobación · publicación sucesora · rollback lógico")
    set_run(r, size=11.5, color=BLUE)
    add_callout(
        doc,
        "Resultado",
        "GATE 04 COMPLETO — implementación, migración PostgreSQL, regresión Gate 03, QA ADMIN/USER y despliegue localhost validados.",
        kind="pass",
    )
    add_table(
        doc,
        ["Campo", "Valor"],
        [
            ["Aplicación", "P&Pmis Construction AI"],
            ["Repositorio", r"C:\Users\Ricardo\Documents\GitHub\pypmis-ia-sas"],
            ["Rama / baseline", "main · HEAD dcea01e · cambios Gate 04 sin commit"],
            ["Fecha de verificación", "10 de agosto de 2026 · America/Bogota"],
            ["Entorno", "Docker Compose · localhost · PostgreSQL 16 · Redis 7"],
            ["Release vigente", "ES-PYP-CORE-RECONCILED-20260809 · ID 1 · published"],
            ["Fuente", "PROMPT_04_WORKSPACE_STRUCTURE_REVISION_MANAGER.md"],
            ["SHA-256 fuente", "24542e052d8cc21c0d0e347aa08f12330ba5c802294f978dd6d7f16b6f662575"],
        ],
        widths=[1.75, 5.35],
        font_size=8.2,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Condición de QA",
        "La publicación sucesora y el rollback se probaron exclusivamente en una base aislada en memoria. En PostgreSQL local no se creó ningún DRAFT ni se alteró el tenant publicado.",
        kind="info",
    )

    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    doc.add_heading("Contenido", level=1)
    contents = [
        "Resumen ejecutivo",
        "Alcance, baseline y exclusiones",
        "Arquitectura implementada",
        "Modelo de revisión",
        "Persistencia, migración e inmutabilidad",
        "Servicio de dominio y reglas de gobierno",
        "API ADMIN",
        "RBAC",
        "Auditoría",
        "UI ADMIN — Revision Manager",
        "USER MODE — Enterprise Explorer",
        "Flujo Create New Revision",
        "Operaciones DRAFT",
        "Record Code Preview y recodificación",
        "Validate Revision",
        "Compare Releases",
        "Aprobación explícita",
        "Publish Successor y rollback aislado",
        "Despliegue localhost",
        "Suites y controles de calidad",
        "Matriz de aceptación",
        "Riesgos y siguiente incremento",
        "Anexo A — archivos modificados",
        "Anexo B — endpoints, permisos y eventos",
    ]
    add_table(doc, ["N.º", "Sección"], [[str(i + 1), item] for i, item in enumerate(contents)], widths=[0.65, 6.45], font_size=8.2)
    add_callout(
        doc,
        "Lectura recomendada",
        "Las secciones 3 a 9 describen el diseño técnico; 10 a 18 documentan el flujo funcional; 19 a 21 contienen la evidencia de despliegue y aceptación.",
        kind="info",
    )

    section(doc, "1", "Resumen ejecutivo", page_break=True)
    add_paragraph(
        doc,
        "Se implementó el Workspace Structure Revision Manager para administrar cambios sucesores de la estructura empresarial sin editar el release CORE publicado. La solución reutiliza EnterpriseCoreRelease y adopta un snapshot completo editable en estado DRAFT; no introduce un motor genérico de versionado ni duplica físicamente workspaces durante la preparación de la revisión.",
    )
    add_paragraph(
        doc,
        "El flujo incorpora creación desde el último release vigente, Add/Edit/Move/Classify/Archive, cálculo backend de record_code, recodificación recursiva de descendientes, validación de gobierno, diff detallado, aprobación explícita con hashes, publicación transaccional sucesora, inmutabilidad e idempotencia, y rollback lógico con confirmación.",
    )
    add_status_table(
        doc,
        [
            ("Create New Revision", "PASS", "Clone idempotente desde release 1; REV-002 en fixture aislado."),
            ("Published release preserved", "PASS", "PostgreSQL conserva release 1 y fingerprint 58be6cce288d…"),
            ("Draft editing", "PASS", "Cinco operaciones DRAFT verificadas."),
            ("Validate / Compare / Approval", "PASS", "Guards de estado y hashes probados."),
            ("Publish Successor / Rollback", "PASS", "Ciclo completo ejecutado en SQLite aislado."),
            ("ADMIN / USER QA", "PASS", "Navegador real sobre localhost; 0 errores de consola."),
            ("Tests / Build / Migration", "PASS", "53 backend + 142 frontend; Alembic 0032 head."),
        ],
    )
    add_callout(
        doc,
        "Decisión de cierre",
        "Gate 04 satisface los criterios funcionales y técnicos. Project Creator no fue iniciado, conforme a la exclusión expresa.",
        kind="pass",
    )

    section(doc, "2", "Alcance, baseline y exclusiones")
    doc.add_heading("Baseline preservado", level=2)
    add_table(
        doc,
        ["Objeto", "Valor verificado"],
        [
            ["Tenant", "P&P Ingeniería y Proyectos · id 1"],
            ["Release", "ES-PYP-CORE-RECONCILED-20260809 · id 1 · published"],
            ["Workspaces", "14"],
            ["Strategic Objectives", "7"],
            ["Classifications", "26"],
            ["Links", "0"],
            ["PROPERTY / FACILITY", "0 / 0"],
            ["Fingerprint", "58be6cce288d4067f83f43bd9ac850da76f20784b4f113aac50398a6bf70c828"],
        ],
        widths=[2.0, 5.1],
        first_col_bold=True,
    )
    doc.add_heading("Incluido", level=2)
    add_bullets(
        doc,
        [
            "Gestión gobernada de revisiones CORE sucesoras en ADMIN MODE.",
            "Snapshot DRAFT completo, validación, diff, aprobación, publicación e historial.",
            "Formulario único por tipo y reglas de composición reutilizadas.",
            "UI accesible con texto e iconos para estados de cambio.",
            "USER Explorer limitado al release publicado vigente.",
        ],
    )
    doc.add_heading("Exclusiones respetadas", level=2)
    add_bullets(
        doc,
        [
            "Project Creator, EXPERIENCE, PROPERTY, FACILITY y Asset Manager.",
            "CPM, XML P6, workflow genérico, microservicios o framework universal de versionado.",
            "Mutación de datos reales para QA funcional del sucesor.",
        ],
    )

    section(doc, "3", "Arquitectura implementada", page_break=True)
    add_codebox(
        doc,
        "React / Vite — WorkspaceRevisionManager\n"
        "        ↓ REST + Bearer + tenant context\n"
        "FastAPI — router_admin + exact RBAC guards\n"
        "        ↓ EnterpriseStructureRevisionService\n"
        "SQLAlchemy — EnterpriseCoreRelease snapshot DRAFT\n"
        "        ↓ Alembic 0032 / PostgreSQL trigger\n"
        "PostgreSQL 16 — release history + workspaces materialized on publish\n"
        "Redis 7 / Celery — existing platform services, unchanged",
    )
    add_table(
        doc,
        ["Capa", "Componente", "Responsabilidad"],
        [
            ["Presentación", "WorkspaceRevisionManager.tsx", "Árbol DRAFT, editor, preview, estados, diff y gate actions."],
            ["Integración", "enterpriseStructureApi", "Contratos tipados para clone, CRUD DRAFT, validate, diff, approve, publish y rollback."],
            ["API", "router_admin.py", "Rutas ADMIN con permiso específico y alcance de organización."],
            ["Dominio", "revisions.py", "Reglas, hashes, snapshot, validación, diff, materialización y auditoría."],
            ["Persistencia", "EnterpriseCoreRelease", "Estado de revisión, snapshot, hashes, aprobación y vínculo al release anterior."],
            ["Seguridad", "permissions.py + PERMISSION_SEED", "Evaluación sin bypass y grants a roles administrativos existentes."],
            ["Base de datos", "Alembic 20260810_0032", "Columnas, FKs, índice y trigger PostgreSQL de inmutabilidad."],
        ],
        widths=[1.05, 2.25, 3.8],
        font_size=7.7,
    )
    add_paragraph(
        doc,
        "La alternativa de snapshot completo fue elegida por simplicidad y compatibilidad. Con 14 workspaces el costo de serialización es bajo, el diff es determinista y la publicación puede materializarse de forma transaccional usando identidades declarativas existentes.",
    )

    section(doc, "4", "Modelo de revisión")
    add_table(
        doc,
        ["Campo", "Uso en Gate 04"],
        [
            ["revision_number", "Secuencia lógica; el sucesor del release 1 inicia en 2."],
            ["previous_release_id", "FK al release base; se conserva después de publicar."],
            ["state", "draft, published, superseded o unpublished."],
            ["snapshot_json", "Snapshot editable completo mientras state=draft."],
            ["base_content_fingerprint", "Fingerprint del release fuente al crear el DRAFT."],
            ["content_fingerprint", "Hash determinista del release_code y snapshot DRAFT."],
            ["validation_json", "Resultado, checks, errores, conflictos y hashes validados."],
            ["validated_*", "Actor, fecha y hash exacto validado."],
            ["diff_hash", "Hash determinista del cambio contra el release base."],
            ["approved_*", "Actor, fecha, draft_hash y diff_hash explícitamente aprobados."],
            ["created_* / updated_at", "Trazabilidad temporal y autoría del release."],
            ["published_*", "Nulos en DRAFT; obligatorios lógicamente al publicar."],
        ],
        widths=[2.0, 5.1],
        font_size=7.8,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Identidad",
        "Technical ID y external_key se conservan para nodos existentes. Los nodos nuevos reciben external_key interno no editable y technical_id=None hasta la materialización del publish.",
        kind="info",
    )

    section(doc, "5", "Persistencia, migración e inmutabilidad", page_break=True)
    add_paragraph(
        doc,
        "La migración 20260810_0032 extiende enterprise_core_releases de manera aditiva, backfill-a el release publicado existente y deja published_at/published_by_user_id anulables para soportar DRAFT. No borra tablas ni workspaces.",
    )
    add_table(
        doc,
        ["Control", "Implementación", "Resultado"],
        [
            ["Revision metadata", "14 columnas nuevas; FKs de created/validated/approved actor.", "PASS"],
            ["Backfill Gate 03", "revision_number=id, base fingerprint, created/published actors.", "PASS"],
            ["Published immutability", "Listener SQLAlchemy y trigger PostgreSQL.", "PASS"],
            ["Delete protection", "Trigger rechaza DELETE físico de releases CORE.", "PASS"],
            ["Draft mutability", "Solo OLD.state=draft permite cambiar snapshot y hashes.", "PASS"],
            ["PostgreSQL head", "alembic current → 20260810_0032 (head).", "PASS"],
        ],
        widths=[1.7, 4.5, 0.9],
        font_size=7.8,
    )
    add_codebox(doc, "20260810_0031  →  20260810_0032 (head)\nadd governed workspace structure revisions")
    add_paragraph(
        doc,
        "En la base local, después de la migración, existe un único release: ID 1, revision_number 1, state published y conteos 14/7/26/0. No se creó un DRAFT real.",
    )

    section(doc, "6", "Servicio de dominio y reglas de gobierno")
    add_table(
        doc,
        ["Regla", "Comportamiento"],
        [
            ["Fuente vigente", "Solo se clona el release published actual; draft previo se devuelve idempotentemente."],
            ["Operaciones DRAFT", "Add, Edit, Move, Classify y Archive rechazan releases no draft."],
            ["Composición", "Los hijos permitidos se obtienen del catálogo publicado de workspace types."],
            ["Identidad", "No existe endpoint para cambiar technical_id ni external_key."],
            ["Record Code", "Backend calcula el siguiente segmento y recodifica descendientes en movimientos."],
            ["Archive", "Lógico; bloquea raíz y nodos con hijos activos."],
            ["Catálogos", "Valida existencia, estado, applicability y required categories."],
            ["Hash invalidation", "Cada modificación invalida validación y aprobación previas."],
            ["Base changed", "Publicación aborta con BASE_RELEASE_CHANGED."],
            ["Idempotencia", "Segundo publish del mismo release retorna el publicado sin mutaciones."],
        ],
        widths=[1.65, 5.45],
        font_size=7.8,
        first_col_bold=True,
    )

    section(doc, "7", "API ADMIN", page_break=True)
    prefix = "/api/v1/admin-configuration/enterprise-structure"
    api_rows = [
        ["POST", f"{prefix}/enterprise-core-releases/{{published_id}}/clone", "Crear o recuperar DRAFT sucesor."],
        ["GET", f"{prefix}/enterprise-core-releases/{{release_id}}", "Consultar release/revisión."],
        ["PATCH", f"{prefix}/enterprise-core-releases/{{release_id}}", "Editar nombre del DRAFT."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/record-code-preview", "Preview backend BEFORE/AFTER."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/workspaces", "Add Workspace."],
        ["PATCH", f"{prefix}/enterprise-core-releases/{{release_id}}/workspaces/{{key}}", "Edit Workspace."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/workspaces/{{key}}/move", "Move y recodificación."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/workspaces/{{key}}/archive", "Archive lógico."],
        ["PUT", f"{prefix}/enterprise-core-releases/{{release_id}}/workspaces/{{key}}/classifications", "Reemplazar clasificaciones."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/validate", "Validar y fijar hashes."],
        ["GET", f"{prefix}/enterprise-core-releases/{{release_id}}/diff", "Comparar contra previous release."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/approve", "Aprobación explícita."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/publish", "Publicación sucesora."],
        ["POST", f"{prefix}/enterprise-core-releases/{{release_id}}/rollback", "Rollback lógico confirmado."],
    ]
    add_table(doc, ["Método", "Ruta", "Finalidad"], api_rows, widths=[0.62, 4.55, 1.93], font_size=6.8)
    add_callout(
        doc,
        "Seguridad de rutas",
        "Todas las rutas derivan tenant_id y user_id del contexto autenticado y requieren permiso específico; las mutaciones exigen organization-wide scope.",
        kind="info",
    )

    section(doc, "8", "RBAC")
    permission_rows = [
        ["admin.enterprise_structure.revision.create", "Clone / Create New Revision"],
        ["admin.enterprise_structure.revision.edit", "PATCH, preview y operaciones workspace DRAFT"],
        ["admin.enterprise_structure.revision.validate", "Validate"],
        ["admin.enterprise_structure.revision.compare", "Diff / Compare"],
        ["admin.enterprise_structure.revision.approve", "Approve"],
        ["admin.enterprise_structure.publish", "Publish Successor"],
        ["admin.enterprise_structure.rollback", "Rollback lógico"],
    ]
    add_table(doc, ["Permiso", "Operación"], permission_rows, widths=[3.85, 3.25], font_size=7.6)
    add_paragraph(
        doc,
        "La base local contiene los siete permisos con dos grants cada uno, correspondientes a los roles administrativos existentes. require_enterprise_permission valida usuario activo, tenant, asignación vigente, permiso y alcance; no existe bypass.",
    )
    add_callout(
        doc,
        "Separación de funciones",
        "El modelo permite separar create/edit/validate/compare/approve/publish/rollback. La política local actual concede el conjunto administrativo a organization_admin y configuration_admin; una matriz de segregación más estricta es una decisión de gobierno futura.",
        kind="warning",
    )

    section(doc, "9", "Auditoría", page_break=True)
    event_rows = [
        ["enterprise_structure.revision_created", "Creación o clone del DRAFT"],
        ["enterprise_structure.revision_modified", "Add/Edit/Move/Classify/Archive o metadata"],
        ["enterprise_structure.revision_validated", "Resultado, errores, conflictos y hashes"],
        ["enterprise_structure.revision_approved", "Actor y hashes aprobados"],
        ["enterprise_structure.core_published", "Release sucesor materializado"],
        ["enterprise_structure.core_unpublished", "Rollback del release actual"],
    ]
    add_table(doc, ["Evento", "Momento"], event_rows, widths=[3.6, 3.5], font_size=7.8)
    add_paragraph(
        doc,
        "Cada evento se persiste en SecurityEvent con actor, tenant, release, previous release cuando aplica, draft hash, diff hash, timestamp y result. La prueba de rollback exige y encuentra el conjunto completo de seis eventos.",
    )

    section(doc, "10", "UI ADMIN — Revision Manager", page_break=True)
    add_paragraph(
        doc,
        "La pantalla Enterprise Structure Configuration conserva CompactModuleHeader. Cuando existe un CORE publicado, Revision Manager es la vista inicial; el baseline editable anterior permanece disponible como Published baseline y queda bloqueado por coreLocked.",
    )
    add_table(
        doc,
        ["Zona", "Contenido"],
        [
            ["Release strip", "Current Published Release, DRAFT sucesor, Based on y estado."],
            ["Gate actions", "Validate, Compare, Approve y Publish con habilitación por estado."],
            ["Status bar", "Draft hash, diff hash, resultado de validación y aprobación."],
            ["Workspace tree", "Jerarquía DRAFT con Added/Modified/Moved/Archived/Classification/Baseline."],
            ["Workspace detail", "Record code, tipo, parent, status, cambio, descripción y clasificaciones."],
            ["Editor único", "Name, Type, Parent, Description, Responsible, Status y classifications."],
            ["Move preview", "BEFORE / AFTER y lista de descendientes afectados."],
            ["Diff", "Resumen por acción y tabla detallada con old/new/parents/status/classifications."],
        ],
        widths=[1.65, 5.45],
        font_size=7.8,
        first_col_bold=True,
    )
    add_paragraph(doc, "Los estados no dependen solo del color: cada nodo muestra texto visible y un icono. Los botones de árbol usan etiquetas accesibles para expandir/contraer.")
    add_codebox(doc, "ADMIN MODE · ENTERPRISE STRUCTURE\nEnterprise Structure Configuration\nCurrent Published Release: ES-PYP-CORE-RECONCILED-20260809\n[ Create New Revision ]")

    section(doc, "11", "USER MODE — Enterprise Explorer")
    add_paragraph(
        doc,
        "Enterprise Explorer no consume el DRAFT. Durante la edición de una revisión continúa mostrando latest_core_release() publicado; después de publicar un sucesor, la misma consulta apunta al nuevo release vigente.",
    )
    add_table(
        doc,
        ["Verificación real", "Resultado"],
        [
            ["Modo", "USER MODE · Enterprise Strategy Manager"],
            ["Submódulo", "Enterprise Explorer"],
            ["Release visible", "ES-PYP-CORE-RECONCILED-20260809"],
            ["Nodos", "14"],
            ["Fingerprint visible", "58be6cce288d4067…"],
            ["DRAFT visible", "No"],
            ["Árbol jerárquico", "Sí; record codes 01, 01.01, 01.01.01…"],
            ["Errores de consola", "0"],
        ],
        widths=[2.1, 5.0],
        first_col_bold=True,
    )

    section(doc, "12", "Flujo Create New Revision", page_break=True)
    add_codebox(
        doc,
        "Published Release 1\n"
        "ES-PYP-CORE-RECONCILED-20260809\n"
        "        ↓ Create New Revision\n"
        "Draft Release 2\n"
        "ES-PYP-CORE-REV-002\n"
        "previous_release_id = 1\n"
        "        ↓ Edit → Validate → Compare → Approve\n"
        "Publish Successor",
    )
    add_numbered(
        doc,
        [
            "Confirma que el release seleccionado es published y coincide con el published actual.",
            "Devuelve el DRAFT existente si ya fue creado desde la misma base; evita duplicados.",
            "Normaliza el snapshot del release 1 y conserva los 14 workspaces lógicos.",
            "Genera revision_number 2 y release_code ES-PYP-CORE-REV-002.",
            "Registra previous_release_id, base_content_fingerprint, actor y timestamps.",
            "Calcula draft_hash y diff_hash inicial; emite revision_created.",
        ],
    )
    add_callout(
        doc,
        "Resultado aislado",
        "El test crea REV-002, conserva las claves ENT-PYP/BU-CORE/PF-ONE/PF-TWO y verifica que el release publicado no cambia.",
        kind="pass",
    )

    section(doc, "13", "Operaciones DRAFT")
    operation_rows = [
        ["Add Child", "Tipo filtrado por allowed_children; external_key generado; technical_id diferido."],
        ["Edit", "Name, description, responsible y status; no identidad técnica."],
        ["Move", "Valida ciclos/composición; conserva key/id; recodifica subárbol."],
        ["Classify", "Reemplaza valores del workspace; duplicados y catálogo controlados."],
        ["Archive", "Status archived; bloquea raíz o hijos activos; no DELETE físico."],
    ]
    add_table(doc, ["Operación", "Regla principal"], operation_rows, widths=[1.3, 5.8], font_size=7.8, first_col_bold=True)
    doc.add_heading("Ejemplo probado", level=2)
    add_bullets(
        doc,
        [
            "Add Program ‘AI Program’ bajo PF-ONE con preview 01.01.01.01.",
            "Edit a ‘AI SaaS Program’ y descripción ‘Controlled draft edit’.",
            "Classify PF-ONE de growth a operational-excellence.",
            "Archive del programa nuevo; status archived en DRAFT.",
            "Intento Portfolio bajo Portfolio rechazado por parent-child incompatibility.",
        ],
    )

    section(doc, "14", "Record Code Preview y recodificación")
    add_paragraph(
        doc,
        "El frontend nunca calcula ni permite digitar record_code. El backend usa el parent, hermanos, sort order y el código excluido cuando se mueve un nodo. El preview se muestra antes de guardar.",
    )
    add_table(
        doc,
        ["Elemento", "BEFORE", "AFTER"],
        [
            ["Portfolio PF-ONE", "01.01.01", "01.02.01"],
            ["Descendiente Program", "01.01.01.01", "01.02.01.01"],
            ["Descendiente Project", "01.01.01.01.01", "01.02.01.01.01"],
        ],
        widths=[2.6, 2.25, 2.25],
        font_size=8.0,
        first_col_bold=True,
    )
    add_callout(
        doc,
        "Cobertura reforzada",
        "La prueba crea explícitamente Portfolio → Program → Project, compara affected_descendants y valida la recodificación persistida de ambos niveles.",
        kind="pass",
    )

    section(doc, "15", "Validate Revision", page_break=True)
    checks = [
        ["single_root", "Una raíz exacta"],
        ["acyclic", "Cero ciclos"],
        ["parent_child_compatible", "Composition Rules"],
        ["record_code_unique", "Record codes únicos"],
        ["external_key_unique", "Identidad declarativa única"],
        ["required_classifications", "Categorías obligatorias"],
        ["category_applicable", "Applicability por tipo"],
        ["links_valid", "Links con origen/destino válidos"],
        ["cross_tenant_zero", "Snapshot del tenant actual"],
        ["no_orphans", "Todo parent existe"],
        ["status_transitions_valid", "Estados y archivo válidos"],
        ["codes_unique", "Códigos internos no duplicados"],
    ]
    add_table(doc, ["Check", "Control"], checks, widths=[2.85, 4.25], font_size=7.7)
    add_codebox(doc, "VALID\n0 errors\n0 conflicts\nall checks = true", fill=PALE_GREEN)
    add_paragraph(
        doc,
        "El escenario inválido fuerza tenant 999, code y record_code duplicados, un ciclo y un parent inexistente; Validate devuelve valid=false y marca individualmente los checks fallidos.",
    )

    section(doc, "16", "Compare Releases")
    add_table(
        doc,
        ["Acción", "Conteo del escenario aislado", "Ejemplo"],
        [
            ["Added", "1", "Second Business Unit"],
            ["Modified", "1", "BU-CORE → Construction Management"],
            ["Moved", "1", "PF-ONE bajo la nueva Business Unit"],
            ["Archived", "1", "PF-TWO"],
            ["Classification Changes", "1", "PF-ONE: growth → operational-excellence"],
            ["Unchanged", "1", "Enterprise root"],
        ],
        widths=[1.55, 1.65, 3.9],
        font_size=7.7,
    )
    add_paragraph(
        doc,
        "RevisionDiffItem incluye action, old/new record code, type, name, parent before/after, status before/after, classifications before/after y affected descendants. diff_hash se calcula sobre la representación normalizada y ordenada.",
    )

    section(doc, "17", "Aprobación explícita")
    add_paragraph(
        doc,
        "Validate no implica aprobación. Approve requiere que el DRAFT esté validado, que la validación sea válida y que draft_hash y diff_hash suministrados coincidan con la revisión actual.",
    )
    add_table(
        doc,
        ["Control", "Resultado probado"],
        [
            ["Publish sin aprobación", "409 · approval required"],
            ["Approve con draft hash incorrecto", "409 · HASH_MISMATCH"],
            ["Approve válido", "approved_by / approved_at / approved hashes persistidos"],
            ["Edición posterior", "Invalida validation y approval"],
        ],
        widths=[2.5, 4.6],
        first_col_bold=True,
    )

    section(doc, "18", "Publish Successor y rollback aislado", page_break=True)
    add_table(
        doc,
        ["Guard de publicación", "Comprobación"],
        [
            ["state=draft", "Rechaza releases históricos/no DRAFT."],
            ["validation=PASS", "Mismos draft_hash y diff_hash validados."],
            ["approval=PRESENT", "Aprobación explícita del contenido actual."],
            ["hash=MATCH", "Payload, release y aprobación coinciden."],
            ["diff_hash=MATCH", "Diff no cambió desde aprobación."],
            ["source still current", "Release base sigue siendo published vigente."],
        ],
        widths=[2.35, 4.75],
        font_size=7.8,
    )
    add_paragraph(
        doc,
        "En la fixture aislada, Release 2 pasa a published, Release 1 a superseded y previous_release_id permanece en 1. La materialización actualiza EnterpriseWorkspace por external_key en una transacción. Un segundo publish devuelve el mismo release sin mutaciones.",
    )
    add_paragraph(
        doc,
        "Rollback exige confirm=true y motivo. El release 2 pasa a unpublished, el release 1 vuelve a published y su snapshot se rematerializa; BU-CORE recupera ‘Core Business Unit’. No se elimina ningún release ni workspace.",
    )
    add_status_table(
        doc,
        [
            ("Publish successor", "PASS", "Release 2 published; previous_release_id=1."),
            ("Previous preserved", "PASS", "Release 1 retained as superseded/history."),
            ("Second publish", "PASS", "Safe replay returns same release."),
            ("Immutable after publish", "PASS", "SQLAlchemy raises on field mutation."),
            ("BASE_RELEASE_CHANGED", "PASS", "Competing published release blocks publish."),
            ("Logical rollback", "PASS", "Explicit confirmation restores release 1."),
        ],
    )

    section(doc, "19", "Despliegue localhost", page_break=True)
    add_table(
        doc,
        ["Servicio", "Endpoint/puerto", "Estado"],
        [
            ["Frontend React/Vite", "http://127.0.0.1:5173/app", "UP · HTTP 200"],
            ["FastAPI", "http://127.0.0.1:8000", "UP · healthy"],
            ["Readiness", "/api/v1/health/ready", "api/database/redis = ok"],
            ["PostgreSQL 16", "127.0.0.1:5432", "UP · healthy"],
            ["Redis 7", "127.0.0.1:6379", "UP · healthy"],
            ["Celery worker", "control-core queue", "UP"],
            ["Celery beat", "scheduler", "UP"],
        ],
        widths=[2.1, 3.1, 1.9],
        font_size=7.8,
    )
    add_paragraph(
        doc,
        "La imagen backend fue reconstruida después de la implementación; Alembic 0032 se aplicó antes de recrear api/worker/beat. Los volúmenes de PostgreSQL y documentos no fueron eliminados.",
    )
    add_table(
        doc,
        ["QA navegador", "Evidencia"],
        [
            ["USER", "Enterprise Explorer · release 1 · 14 nodos · no DRAFT."],
            ["ADMIN", "Enterprise Structure Configuration · CORE published · Create New Revision."],
            ["Consola", "0 errores."],
            ["Datos reales", "1 release published; 0 drafts creados por QA."],
        ],
        widths=[1.6, 5.5],
        first_col_bold=True,
    )

    section(doc, "20", "Suites y controles de calidad", page_break=True)
    quality_rows = [
        ["Backend Gate 04", "pytest test_enterprise_structure_revisions.py", "9 passed · 6.72 s"],
        ["Jerarquía/importador", "pytest Gate 2A/importer", "20 passed · 23.85 s"],
        ["Publish/rollback Gate 03", "pytest test_enterprise_structure_apply.py", "24 passed · 236.37 s"],
        ["Backend style", "ruff format --check + ruff check", "118 files · PASS"],
        ["Frontend focal", "Vitest Enterprise + Revision Manager", "10 passed"],
        ["Frontend completo", "npm run test -- --run", "24 files · 142 passed · 82.52 s"],
        ["Frontend format", "npm run format:check", "PASS"],
        ["Frontend lint", "ESLint max-warnings=10", "0 errors · 8 warnings preexistentes"],
        ["Frontend build", "tsc && vite build", "PASS · 2343 modules"],
        ["Migration", "alembic upgrade/current", "0032 head · PostgreSQL PASS"],
        ["Repository", "git diff --check", "PASS"],
        ["Browser QA", "USER + ADMIN + console", "PASS / PASS / 0 errors"],
    ]
    add_table(doc, ["Suite/control", "Comando o alcance", "Resultado"], quality_rows, widths=[1.85, 3.5, 1.75], font_size=7.2)
    add_callout(
        doc,
        "Incidente de entorno resuelto",
        "Una ejecución paralela agotó el espacio temporal de SQLite y bloqueó WSL/Docker. Se eliminaron solo directorios .pytest_tmp del repositorio, se reinició el motor sin borrar volúmenes y las suites se repitieron secuencialmente con 20/20 y 24/24 PASS.",
        kind="warning",
    )

    section(doc, "21", "Matriz de aceptación", page_break=True)
    acceptance = [
        ("1", "Crear draft desde published", "PASS", "REV-002, state=draft"),
        ("2", "Published permanece intacto", "PASS", "Snapshot/fingerprint preservados"),
        ("3", "previous_release correcto", "PASS", "previous_release_id=1"),
        ("4", "Add workspace", "PASS", "AI Program"),
        ("5", "Edit workspace", "PASS", "AI SaaS Program"),
        ("6", "Move workspace", "PASS", "PF-ONE movido"),
        ("7", "Archive workspace", "PASS", "Status archived"),
        ("8", "Classifications", "PASS", "growth → operational-excellence"),
        ("9", "Record code preview", "PASS", "01.01.01.01"),
        ("10", "Descendant recoding", "PASS", "Program + Project recodificados"),
        ("11", "Cycle blocked", "PASS", "409 cycles"),
        ("12", "Invalid parent blocked", "PASS", "Composition rule"),
        ("13", "Duplicate record code blocked", "PASS", "Validate false"),
        ("14", "external_key preserved", "PASS", "PF-ONE estable"),
        ("15", "Validate valid", "PASS", "0 errors / 0 conflicts"),
        ("16", "Validate invalid", "PASS", "5 checks fallidos"),
        ("17", "Diff added", "PASS", "added=1"),
        ("18", "Diff modified", "PASS", "modified=1"),
        ("19", "Diff moved", "PASS", "moved=1"),
        ("20", "Diff archived", "PASS", "archived=1"),
        ("21", "Diff classifications", "PASS", "classification_changes=1"),
        ("22", "Approval required", "PASS", "Publish sin approve bloqueado"),
        ("23", "Hash mismatch blocked", "PASS", "HASH_MISMATCH"),
        ("24", "Base release changed", "PASS", "BASE_RELEASE_CHANGED"),
        ("25", "Publish successor", "PASS", "Release 2 published aislado"),
        ("26", "Previous preserved", "PASS", "Release 1 superseded/histórico"),
        ("27", "Release immutable", "PASS", "Listener + trigger"),
        ("28", "Second publish safe", "PASS", "Same release, 0 mutaciones"),
        ("29", "Rollback logical", "PASS", "Release 1 restored"),
        ("30", "Cross-tenant blocked", "PASS", "cross_tenant_zero=false"),
        ("31", "RBAC", "PASS", "7 permisos exactos"),
        ("32", "Audit events", "PASS", "6 eventos verificados"),
        ("33", "ADMIN UI", "PASS", "Revision Manager real + tests"),
        ("34", "Regression Gate 03", "PASS", "44 pruebas históricas"),
        ("35", "Frontend build", "PASS", "tsc + Vite"),
        ("36", "Migrations PostgreSQL", "PASS", "Alembic 0032 head"),
    ]
    add_table(doc, ["#", "Criterio", "Estado", "Evidencia"], [list(item) for item in acceptance], widths=[0.35, 2.65, 0.75, 3.35], font_size=6.8)
    add_callout(
        doc,
        "Criterios de cierre",
        "Create New Revision, preservation, editing, preview, validate, compare, approval, publish successor, previous_release, immutability, rollback, ADMIN QA, USER published-only y tests: PASS.",
        kind="pass",
    )

    section(doc, "22", "Riesgos y siguiente incremento", page_break=True)
    add_table(
        doc,
        ["Riesgo / deuda", "Impacto", "Mitigación recomendada"],
        [
            ["Snapshot completo en estructuras muy grandes", "Crecimiento de JSON y costo de diff.", "Medir con 1k/10k nodos antes de optimizar; no introducir deltas prematuramente."],
            ["Segregación de funciones", "Dos roles locales poseen approve y publish.", "Definir matriz SoD por tenant y asignaciones separadas."],
            ["E2E mutacional PostgreSQL", "El ciclo completo se ejecutó en fixture aislado.", "Añadir base efímera PostgreSQL en CI con clone→publish→rollback."],
            ["Concurrencia de drafts", "Existe guard de base vigente, no lock de editor multiusuario.", "Agregar optimistic version/If-Match si aparece edición concurrente real."],
            ["Espacio local limitado", "Docker/pytest paralelo puede bloquear WSL.", "Mantener >10 GB libres y ejecutar suites pesadas secuencialmente."],
            ["8 warnings BIM viewer", "No bloquean Gate 04, sí son deuda técnica.", "Corregir hooks/dependencias en un sprint de hardening separado."],
        ],
        widths=[2.15, 2.1, 2.85],
        font_size=7.2,
    )
    doc.add_heading("Recomendación", level=2)
    add_paragraph(
        doc,
        "El siguiente incremento recomendado es un Gate 04 de hardening operativo: E2E transaccional con PostgreSQL efímero, prueba de dos editores concurrentes, segregación approve/publish y medición de snapshots grandes. No iniciar Project Creator automáticamente; ese trabajo requiere una autorización y un prompt propios.",
    )
    add_callout(doc, "Cierre", "Detener el alcance al finalizar Gate 04. Estado final: completo y desplegado en localhost.", kind="pass")

    section(doc, "A", "Anexo A — archivos modificados", page_break=True)
    backend_files = [
        "backend/app/modules/enterprise_structure/constants.py",
        "backend/app/modules/enterprise_structure/importer/publish.py",
        "backend/app/modules/enterprise_structure/models.py",
        "backend/app/modules/enterprise_structure/repository.py",
        "backend/app/modules/enterprise_structure/revisions.py (nuevo)",
        "backend/app/modules/enterprise_structure/router_admin.py",
        "backend/app/modules/enterprise_structure/schemas.py",
        "backend/app/modules/enterprise_structure/service.py",
        "backend/alembic/versions/20260810_0032_workspace_revision_manager.py (nuevo)",
        "backend/tests/test_enterprise_structure_revisions.py (nuevo)",
    ]
    frontend_files = [
        "frontend/src/features/enterprise-structure/api/index.ts",
        "frontend/src/features/enterprise-structure/components/WorkspaceRevisionManager.tsx (nuevo)",
        "frontend/src/features/enterprise-structure/enterpriseStructure.css",
        "frontend/src/features/enterprise-structure/pages/AdminEnterpriseStructurePage.tsx",
        "frontend/src/features/enterprise-structure/pages/EnterpriseExplorerPage.tsx",
        "frontend/src/features/enterprise-structure/types/index.ts",
        "frontend/tests/workspace-revision-manager.test.tsx (nuevo)",
    ]
    add_table(doc, ["Capa", "Archivo"], [["Backend", item] for item in backend_files] + [["Frontend", item] for item in frontend_files], widths=[1.0, 6.1], font_size=7.4)
    add_paragraph(
        doc,
        "Los artefactos de evidencia del Gate 04 se almacenan en artifacts/enterprise_structure/gate04 y no forman parte de la implementación productiva contabilizada.",
    )

    section(doc, "B", "Anexo B — endpoints, permisos y eventos", page_break=True)
    add_paragraph(doc, "Resumen de trazabilidad cruzada:")
    cross_rows = [
        ["Create", "POST …/{published_id}/clone", "revision.create", "revision_created"],
        ["Edit/Add/Move/Classify/Archive", "PATCH/POST/PUT …/workspaces", "revision.edit", "revision_modified"],
        ["Validate", "POST …/{id}/validate", "revision.validate", "revision_validated"],
        ["Compare", "GET …/{id}/diff", "revision.compare", "—"],
        ["Approve", "POST …/{id}/approve", "revision.approve", "revision_approved"],
        ["Publish", "POST …/{id}/publish", "enterprise_structure.publish", "core_published"],
        ["Rollback", "POST …/{id}/rollback", "enterprise_structure.rollback", "core_unpublished"],
    ]
    add_table(doc, ["Flujo", "API", "Permiso admin.*", "Evento enterprise_structure.*"], cross_rows, widths=[1.35, 2.45, 1.7, 1.6], font_size=6.9)
    add_callout(
        doc,
        "Evidencia reproducible",
        r"Resumen JSON: artifacts\enterprise_structure\gate04\verification_summary.json · Prompt SHA-256 verificado · localhost activo.",
        kind="info",
    )

    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    report = build_report()
    report.save(OUTPUT)
    reopened = Document(OUTPUT)
    text = "\n".join(
        [p.text for p in reopened.paragraphs]
        + [cell.text for table in reopened.tables for row in table.rows for cell in row.cells]
    )
    required = [
        "GATE 04 COMPLETO",
        "ES-PYP-CORE-RECONCILED-20260809",
        "Workspace Structure Revision Manager",
        "Create New Revision",
        "20260810_0032",
        "142 passed",
        "No iniciar Project Creator",
    ]
    missing = [item for item in required if item not in text]
    if missing:
        raise RuntimeError(f"Missing required report content: {missing}")
    if "�" in text:
        raise RuntimeError("Replacement character found in generated report")
    if len(reopened.tables) < 25 or len(reopened.paragraphs) < 120:
        raise RuntimeError("Generated report is unexpectedly small")
    print(f"OUTPUT={OUTPUT}")
    print(f"SIZE={OUTPUT.stat().st_size}")
    print(f"PARAGRAPHS={len(reopened.paragraphs)}")
    print(f"TABLES={len(reopened.tables)}")
    print(f"SECTIONS={len(reopened.sections)}")
    print("STRUCTURAL_QA=PASS")


if __name__ == "__main__":
    main()
