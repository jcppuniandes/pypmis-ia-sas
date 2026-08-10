from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(r"C:\Users\Ricardo\Documents\GitHub\pypmis-ia-sas")
OUTPUT = Path(
    r"C:\Users\Ricardo\Documents\P&P\P&Pmis Construction AI\Diseño\Resumen de Sprint"
    r"\Informe_Tecnico_PPMIS_Workspace_Revision_Manager_Gate04H_2026-08-10.docx"
)

FONT = "Calibri"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
NAVY = "0B2545"
TEAL = "137F83"
MUTED = "5C6F7A"
LIGHT = "F2F4F7"
PALE_BLUE = "E8EEF5"
PALE_GREEN = "E7F4EC"
GREEN = "1C6B43"
PALE_AMBER = "FFF3D6"
AMBER = "7A5A00"
PALE_RED = "FCE9E7"
RED = "9B1C1C"
WHITE = "FFFFFF"


def set_run_font(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:ascii"), FONT)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, *, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    if tr_pr.find(qn("w:cantSplit")) is None:
        tr_pr.append(OxmlElement("w:cantSplit"))


def set_table_geometry(table, widths_dxa: list[int]) -> None:
    if sum(widths_dxa) != 9360:
        raise ValueError(f"Table widths must sum to 9360 DXA, got {sum(widths_dxa)}")
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        prevent_row_split(row)
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    start = OxmlElement("w:fldChar")
    start.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((start, instruction, end))


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in (
        ("Heading 1", 16, BLUE, 16, 8),
        ("Heading 2", 13, BLUE, 12, 6),
        ("Heading 3", 12, DARK_BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = FONT
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for list_style in ("List Bullet", "List Number"):
        style = doc.styles[list_style]
        style.font.name = FONT
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(8)
        style.paragraph_format.line_spacing = 1.167

    header = section.header
    header.is_linked_to_previous = False
    paragraph = header.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("P&Pmis Construction AI  |  Gate 04H")
    set_run_font(run, size=9, bold=True, color=MUTED)

    first_header = section.first_page_header
    first_header.paragraphs[0].clear()

    footer = section.footer
    footer.is_linked_to_previous = False
    paragraph = footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("Informe técnico Gate 04H  |  Página ")
    set_run_font(run, size=9, color=MUTED)
    add_page_number(paragraph)

    first_footer = section.first_page_footer
    paragraph = first_footer.paragraphs[0]
    paragraph.clear()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("P&Pmis Construction AI · Uso interno · 10 de agosto de 2026")
    set_run_font(run, size=9, color=MUTED)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        prefix = paragraph.add_run(bold_prefix)
        set_run_font(prefix, bold=True, color=NAVY)
        remainder = paragraph.add_run(text[len(bold_prefix) :])
        set_run_font(remainder)
    else:
        run = paragraph.add_run(text)
        set_run_font(run)


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    run = paragraph.add_run(text)
    set_run_font(run)


def add_numbered(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Number")
    run = paragraph.add_run(text)
    set_run_font(run)


def add_callout(doc: Document, label: str, text: str, *, fill: str, border: str) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(10)
    paragraph.paragraph_format.left_indent = Inches(0.12)
    paragraph.paragraph_format.right_indent = Inches(0.12)
    paragraph.paragraph_format.line_spacing = 1.10
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    p_bdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "6")
    left.set(qn("w:color"), border)
    p_bdr.append(left)
    p_pr.append(p_bdr)
    label_run = paragraph.add_run(f"{label}: ")
    set_run_font(label_run, bold=True, color=border)
    text_run = paragraph.add_run(text)
    set_run_font(text_run, color=NAVY)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths_dxa: list[int],
    *,
    header_fill: str = LIGHT,
) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_repeat_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, header_fill)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        run = paragraph.add_run(header)
        set_run_font(run, size=9.5, bold=True, color=NAVY)
    for row_values in rows:
        row = table.add_row()
        for index, value in enumerate(row_values):
            cell = row.cells[index]
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.05
            if index == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            elif len(value) <= 18:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph.add_run(value)
            color = GREEN if value == "PASS" else RED if value in {"NOT RUN", "FAILED"} else NAVY
            set_run_font(run, size=9.3, bold=value in {"PASS", "NOT RUN", "HARDENING_REQUIRED"}, color=color)
    set_table_geometry(table, widths_dxa)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(4)
    spacer.paragraph_format.space_after = Pt(4)


def add_page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_title_page(doc: Document) -> None:
    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_before = Pt(20)
    kicker.paragraph_format.space_after = Pt(4)
    run = kicker.add_run("INFORME TÉCNICO · GATE 04H")
    set_run_font(run, size=11, bold=True, color=TEAL)

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    run = title.add_run("Workspace Revision Manager\nOperational Hardening")
    set_run_font(run, size=23, bold=True, color=NAVY)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(20)
    run = subtitle.add_run("P&Pmis Construction AI · Enterprise Structure")
    set_run_font(run, size=14, color=MUTED)

    metadata = [
        ("Aplicación", "P&Pmis Construction AI"),
        ("Gate", "04H · Workspace Revision Manager Operational Hardening"),
        ("Fecha", "10 de agosto de 2026"),
        ("Entorno", "Windows · Docker Desktop · PostgreSQL 16 · React/TypeScript"),
        ("Baseline", "ES-PYP-CORE-RECONCILED-20260809 · Release ID 1"),
        ("Estado", "HARDENING_REQUIRED"),
    ]
    for label, value in metadata:
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(3)
        paragraph.paragraph_format.line_spacing = 1.05
        run = paragraph.add_run(f"{label}: ")
        set_run_font(run, size=11, bold=True, color=NAVY)
        run = paragraph.add_run(value)
        set_run_font(run, size=11, color=RED if label == "Estado" else NAVY, bold=label == "Estado")

    add_callout(
        doc,
        "Dictamen",
        "El hardening y su arnés reproducible quedaron implementados. El cierre de aceptación permanece pendiente "
        "porque el equipo no cumple el mínimo de 10 GB libres exigido para E2E PostgreSQL, regresión completa y "
        "benchmark 100/1k/10k. No se inicia Gate 05.",
        fill=PALE_AMBER,
        border=AMBER,
    )


def build_report() -> Document:
    doc = Document()
    configure_document(doc)
    add_title_page(doc)
    add_page_break(doc)

    add_heading(doc, "1. Resumen ejecutivo")
    add_body(
        doc,
        "Gate 04H endurece el Workspace Structure Revision Manager sin agregar capacidades de negocio. El trabajo "
        "incorpora concurrencia optimista, segregación de funciones, Four-Eyes, invalidación de estados, pruebas "
        "PostgreSQL efímeras, benchmark reproducible, CI dedicado y trazabilidad del actor modificador.",
    )
    add_callout(
        doc,
        "Resultado",
        "La implementación está disponible en el repositorio y el localhost responde. La aceptación integral no se "
        "declara: el preflight encontró 1,50 GB libres frente al requisito de 10 GB, por lo que se aplicó "
        "ABORT HEAVY TESTS. Estado final: HARDENING_REQUIRED.",
        fill=PALE_RED,
        border=RED,
    )
    add_heading(doc, "Objetivos cubiertos", level=2)
    for item in (
        "Contrato If-Match y versionado atómico para evitar pérdida silenciosa de cambios.",
        "Invalidación determinista de Validate y Approval ante cualquier modificación posterior.",
        "Roles Structure Editor, Structure Approver y Structure Publisher con filtro de deber obligatorio.",
        "Regla Four-Eyes: el aprobador no puede publicar la misma revisión.",
        "E2E PostgreSQL, benchmark y job CI implementados para ejecución en un entorno con capacidad suficiente.",
        "Baseline persistente inspeccionado antes y después, sin DRAFT ni fixtures de prueba.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "Índice del informe", level=2)
    for item in (
        "Baseline, alcance y exclusiones",
        "Arquitectura de concurrencia e invalidación",
        "Segregación de funciones y Four-Eyes",
        "PostgreSQL E2E, benchmark y CI",
        "Migración, hardening y despliegue localhost",
        "Verificación, riesgos y plan de cierre",
    ):
        add_numbered(doc, item)

    add_page_break(doc)
    add_heading(doc, "2. Baseline protegido, alcance y exclusiones")
    add_body(
        doc,
        "La regla principal de 04H fue no usar la base persistente para pruebas mutacionales. Solo se aplicó la "
        "migración aditiva 0033, se inicializó trazabilidad operativa y se sembraron roles de seguridad del producto.",
    )
    add_heading(doc, "Comprobación antes/después", level=2)
    add_table(
        doc,
        ["Control", "Antes", "Después", "Resultado"],
        [
            ["Tenant", "P&P Ingeniería y Proyectos", "Sin cambio", "PASS"],
            ["Release", "ES-PYP-CORE-RECONCILED-20260809", "Sin cambio", "PASS"],
            ["Release ID / estado", "1 / published", "1 / published", "PASS"],
            ["Release count / DRAFT", "1 / 0", "1 / 0", "PASS"],
            ["Workspaces", "14", "14", "PASS"],
            ["Strategic Objectives", "7", "7", "PASS"],
            ["Classifications", "26", "26", "PASS"],
            ["Links", "0", "0", "PASS"],
        ],
        [2700, 2520, 2520, 1620],
    )
    add_body(
        doc,
        "Cambios operativos autorizados: revision_version=1; last_modified_by_user_id inicializado desde "
        "created_by_user_id; roles structure_editor, structure_approver y structure_publisher creados; el "
        "administrador vigente recibió únicamente Structure Editor para conservar edición sin crear un bypass de "
        "aprobación/publicación.",
    )
    add_heading(doc, "Exclusiones respetadas", level=2)
    for item in (
        "No se inició Project Creator ni Gate 05.",
        "No se modificaron EXPERIENCE, PROPERTY, FACILITY, Asset Manager, CPM ni XML P6.",
        "No se creó workflow engine, microservicio, almacenamiento Delta ni rediseño de Enterprise Structure.",
        "No se refactorizó BIM Manager.",
    ):
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "3. Control de concurrencia y contrato HTTP")
    add_heading(doc, "Versionado optimista", level=2)
    add_body(
        doc,
        "EnterpriseCoreRelease incorpora revision_version entero no nulo y lo declara como version_id_col de "
        "SQLAlchemy. Cada UPDATE emite una condición sobre la versión observada; si otro administrador ya confirmó "
        "N+1, el segundo UPDATE afecta cero filas y se traduce a 409 REVISION_VERSION_CONFLICT.",
    )
    add_table(
        doc,
        ["Operación", "Método", "Control", "Respuesta"],
        [
            ["Metadata edit", "PATCH release", "If-Match requerido", "ETag nuevo"],
            ["Add", "POST workspaces", "If-Match requerido", "CoreRevision"],
            ["Edit", "PATCH workspace", "If-Match requerido", "CoreRevision"],
            ["Move", "POST move", "If-Match requerido", "CoreRevision"],
            ["Classify", "PUT classifications", "If-Match requerido", "CoreRevision"],
            ["Archive", "POST archive", "If-Match requerido", "CoreRevision"],
        ],
        [2520, 2160, 2520, 2160],
        header_fill=PALE_BLUE,
    )
    add_heading(doc, "Escenario de dos editores", level=2)
    for item in (
        "Editor A y Editor B leen la revisión en versión N.",
        "Editor A guarda con If-Match N y recibe N+1.",
        "Editor B intenta guardar con If-Match N.",
        "El backend rechaza con 409 REVISION_VERSION_CONFLICT; no realiza merge silencioso.",
    ):
        add_numbered(doc, item)
    add_callout(
        doc,
        "Mensaje UI",
        "This revision changed since you opened it. Reload the latest version before continuing.",
        fill=PALE_BLUE,
        border=BLUE,
    )
    add_heading(doc, "Trazabilidad", level=2)
    add_body(
        doc,
        "last_modified_by_user_id registra el actor de la última mutación del DRAFT y se expone como "
        "last_modified_by. GET y respuestas mutacionales incluyen ETag; la lectura real del release 1 devolvió "
        "ETag \"1\", revision_version 1 y last_modified_by admin@demo.local.",
    )

    add_page_break(doc)
    add_heading(doc, "4. Invalidación de Validate y Approval")
    add_body(
        doc,
        "_refresh_draft centraliza el cambio de snapshot/metadata, recalcula hashes, actualiza el actor y elimina "
        "validation_json, validated_at, validated_draft_hash y todos los campos de aprobación. El hash del DRAFT "
        "ahora incluye release_name además de release_code y snapshot, por lo que una edición de metadata también "
        "invalida el estado previo.",
    )
    add_table(
        doc,
        ["Caso", "Secuencia", "Resultado esperado", "Implementación"],
        [
            ["A", "Validate N → mutate → approve hashes N", "HASH_MISMATCH", "Implementado"],
            ["B", "Approve N → mutate → publish", "APPROVAL_INVALIDATED", "Implementado"],
            ["C", "Approve → same actor publish", "FOUR_EYES_VIOLATION", "Implementado"],
            ["D", "Publish replay", "0 mutations", "Implementado"],
        ],
        [1080, 3420, 2700, 2160],
    )
    add_heading(doc, "Orden de validaciones", level=2)
    for item in (
        "Publicación exige aprobación vigente antes de aceptar hashes.",
        "Si la aprobación fue borrada por una mutación, responde APPROVAL_INVALIDATED.",
        "Si el aprobador coincide con el publicador, responde FOUR_EYES_VIOLATION.",
        "Después valida hashes, resultado de Validate y release base vigente.",
    ):
        add_bullet(doc, item)

    add_page_break(doc)
    add_heading(doc, "5. Segregación de funciones (SoD) y Four-Eyes")
    add_table(
        doc,
        ["Rol", "Puede", "No puede"],
        [
            ["Structure Editor", "create · edit · validate · compare", "approve · publish · rollback"],
            ["Structure Approver", "compare · approve", "edit · publish · rollback"],
            ["Structure Publisher", "compare · publish · rollback", "edit · approve"],
        ],
        [2160, 3600, 3600],
        header_fill=PALE_BLUE,
    )
    add_body(
        doc,
        "REVISION_DUTY_ROLES filtra los roles aceptados en los endpoints de revisión. Los grants amplios de "
        "organization_admin/configuration_admin no actúan como bypass. Para compatibilidad, el administrador inicial "
        "recibe Structure Editor; Approver y Publisher deben asignarse explícitamente a usuarios diferentes.",
    )
    add_heading(doc, "Controles de alcance y vigencia", level=2)
    for item in (
        "Tenant scope: el usuario debe pertenecer al tenant de la solicitud.",
        "Vigencia: starts_at y ends_at filtran asignaciones futuras o expiradas.",
        "Scope: las operaciones de revisión requieren alcance organization.",
        "Four-Eyes: approved_by_user_id debe ser distinto de published_by_user_id.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "Pruebas SoD declaradas", level=2)
    add_body(
        doc,
        "El E2E PostgreSQL contiene los 11 casos requeridos: editor edita/no aprueba/no publica; approver "
        "aprueba/no edita; publisher publica/no aprueba; mismo usuario approve+publish bloqueado; cross-tenant, "
        "asignación expirada y scope inválido bloqueados. Estado local: NOT RUN por preflight de disco.",
    )

    add_page_break(doc)
    add_heading(doc, "6. PostgreSQL E2E efímero y CI")
    add_body(
        doc,
        "docker-compose.gate04h.yml crea PostgreSQL 16 con almacenamiento tmpfs, sin volumen persistente. El "
        "contenedor de prueba aplica Alembic hasta head, ejecuta pytest sobre fixtures sintéticos y corre el benchmark. "
        "GATE04H_EPHEMERAL=true y la validación del esquema impiden usar SQLite o una URL no PostgreSQL.",
    )
    add_heading(doc, "Flujo transaccional implementado", level=2)
    for item in (
        "Seed published release sintético con Enterprise, Business Unit A/B, Portfolio A, Program A y Project A.",
        "Create New Revision; Add Portfolio/Program/Project; Edit; Move con recodificación de descendientes.",
        "Classify; Archive; Record Code Preview; Validate; Compare; Approve; Publish Successor.",
        "Verificación de previous_release; replay idempotente; segundo successor; rollback lógico.",
        "Verificación de no borrado físico y eventos de auditoría.",
    ):
        add_numbered(doc, item)
    add_heading(doc, "Job CI", level=2)
    add_table(
        doc,
        ["Paso", "Implementación"],
        [
            ["Job", "enterprise-structure-postgres-e2e"],
            ["Migración", "alembic upgrade head + current"],
            ["Pruebas", "ciclo, concurrencia, invalidación, SoD, Four-Eyes"],
            ["Benchmark", "100 / 1.000 / 10.000 nodos"],
            ["Teardown", "docker compose down; sin volumen persistente"],
            ["Artifacts", "actions/upload-artifact v4"],
        ],
        [2700, 6660],
    )
    add_callout(
        doc,
        "Estado local",
        "NOT RUN. La implementación no se confunde con evidencia de ejecución. El pipeline deberá producir los "
        "resultados PASS antes de cambiar el dictamen a READY_FOR_PROJECT_CREATOR.",
        fill=PALE_AMBER,
        border=AMBER,
    )

    add_page_break(doc)
    add_heading(doc, "7. Benchmark de snapshots y decisión Snapshot vs Delta")
    add_body(
        doc,
        "snapshot_benchmark.py crea tenants sintéticos aislados de 100, 1.000 y 10.000 nodos. Mide tamaño JSON, "
        "clone, load, preview, move, validate, diff, approve, publish, rollback, transacción total y memoria pico con "
        "tracemalloc. No se definieron umbrales antes de medir.",
    )
    add_table(
        doc,
        ["Métrica", "100", "1k", "10k"],
        [
            ["Snapshot size", "NOT RUN", "NOT RUN", "NOT RUN"],
            ["Clone ms", "NOT RUN", "NOT RUN", "NOT RUN"],
            ["Validate ms", "NOT RUN", "NOT RUN", "NOT RUN"],
            ["Diff ms", "NOT RUN", "NOT RUN", "NOT RUN"],
            ["Publish ms", "NOT RUN", "NOT RUN", "NOT RUN"],
            ["Rollback ms", "NOT RUN", "NOT RUN", "NOT RUN"],
            ["Peak memory", "NOT RUN", "NOT RUN", "NOT RUN"],
        ],
        [3600, 1920, 1920, 1920],
    )
    add_callout(
        doc,
        "Decisión",
        "PENDING_MEASUREMENT. Se mantiene snapshot completo y no se implementa Delta/Change Set. No se recomienda "
        "Delta sin observar una limitación concreta en las mediciones 10k.",
        fill=PALE_BLUE,
        border=BLUE,
    )
    add_heading(doc, "Criterio para cerrar la decisión", level=2)
    add_body(
        doc,
        "Tras ejecutar el benchmark se debe anexar snapshot_benchmark.json/csv, revisar tiempo de transacción y "
        "memoria pico, y documentar si el snapshot completo es operativamente aceptable. Solo una limitación medida "
        "justificaría abrir un diseño Delta en un gate futuro.",
    )

    add_page_break(doc)
    add_heading(doc, "8. Migración 0033 y hardening técnico")
    add_heading(doc, "Cambios de persistencia", level=2)
    for item in (
        "revision_version INTEGER NOT NULL DEFAULT 1.",
        "last_modified_by_user_id con FK e índice a user_accounts.",
        "Backfill de last_modified_by_user_id desde created_by_user_id.",
        "Mapper version_id_col para compare-and-swap atómico.",
        "Trigger de inmutabilidad recreado de forma compatible con PostgreSQL JSON.",
    ):
        add_bullet(doc, item)
    add_heading(doc, "Defecto heredado descubierto", level=2)
    add_body(
        doc,
        "La primera ejecución transaccional de 0033 reveló que el trigger 0032 usaba "
        "NEW.snapshot_json IS DISTINCT FROM OLD.snapshot_json. PostgreSQL no define igualdad para el tipo JSON, por "
        "lo que cualquier UPDATE del release publicado fallaba con UndefinedFunction. Alembic revirtió toda la "
        "transacción; no hubo cambio parcial.",
    )
    add_body(
        doc,
        "La corrección elimina temporalmente el trigger durante el backfill y lo recrea usando "
        "snapshot_json::text IS DISTINCT FROM OLD.snapshot_json::text. La segunda ejecución alcanzó "
        "20260810_0033 (head) y el baseline permaneció intacto.",
    )
    add_heading(doc, "Carrera de creación", level=2)
    add_body(
        doc,
        "create_revision captura IntegrityError por dos clones concurrentes. Si el otro administrador ya creó el "
        "DRAFT para el mismo release base, retorna ese DRAFT; en otro conflicto devuelve "
        "REVISION_CREATE_CONFLICT en lugar de error 500.",
    )

    add_page_break(doc)
    add_heading(doc, "9. Verificación ejecutada y regresión pendiente")
    add_heading(doc, "Comprobaciones ligeras ejecutadas", level=2)
    add_table(
        doc,
        ["Control", "Resultado", "Detalle"],
        [
            ["Hardening contract pytest", "PASS", "3 passed"],
            ["Workspace Revision Manager Vitest", "PASS", "3 passed"],
            ["Ruff dirigido", "PASS", "format + lint"],
            ["Python AST", "PASS", "módulos modificados"],
            ["Prettier dirigido", "PASS", "frontend + test"],
            ["ESLint dirigido", "PASS", "0 errores"],
            ["TypeScript", "PASS", "tsc --noEmit"],
            ["Alembic", "PASS", "20260810_0033 head"],
            ["API health", "PASS", "database + redis ok"],
            ["Frontend HTTP", "PASS", "127.0.0.1:5173/app → 200"],
        ],
        [3600, 1440, 4320],
    )
    add_heading(doc, "Suites no ejecutadas", level=2)
    add_body(
        doc,
        "Por ABORT HEAVY TESTS no se ejecutaron Gate 02A/02B/03/04 completos, Organization & Security, ADMIN "
        "Enterprise Structure, USER Enterprise Explorer, frontend full suite, Vite production build, Ruff full, E2E "
        "PostgreSQL ni benchmark 100/1k/10k. Estas ausencias impiden declarar Regression=PASS.",
    )

    add_page_break(doc)
    add_heading(doc, "10. Despliegue localhost y recuperación operativa")
    add_table(
        doc,
        ["Servicio", "Estado", "Acceso"],
        [
            ["Frontend", "PASS", "http://127.0.0.1:5173/app"],
            ["API", "PASS", "http://127.0.0.1:8000"],
            ["PostgreSQL", "PASS", "healthy · Alembic 0033"],
            ["Redis", "PASS", "healthy"],
            ["Worker", "PASS", "running"],
            ["Beat", "PASS", "running"],
        ],
        [2880, 1440, 5040],
    )
    add_heading(doc, "Incidente de espacio", level=2)
    add_body(
        doc,
        "El primer build con caché desplegó la versión inicial 04H. Un rebuild posterior, después de intentar limpiar "
        "caché, reinstaló dependencias y falló al desempaquetar por falta de disco; Docker Desktop llegó a 0 MB y "
        "detuvo el daemon. No se recrearon contenedores ni se tocaron volúmenes.",
    )
    add_body(
        doc,
        "Se identificó y eliminó exclusivamente una actualización incompleta de Docker Desktop bajo el directorio "
        "Temp del usuario, liberando espacio regenerable. Docker se reinició, los volúmenes reaparecieron y se "
        "sincronizaron los módulos Python finales dentro de los contenedores existentes. La API confirmó el filtro "
        "structure_publisher, ETag \"1\" y baseline 14/7/26/0.",
    )
    add_callout(
        doc,
        "Riesgo de despliegue",
        "El localhost está operativo, pero la imagen final reproducible debe reconstruirse cuando haya al menos 10 "
        "GB libres. La sincronización de runtime se pierde si los contenedores se recrean antes de ese build.",
        fill=PALE_AMBER,
        border=AMBER,
    )

    add_page_break(doc)
    add_heading(doc, "11. Evidencias y archivos entregados")
    add_table(
        doc,
        ["Área", "Archivo principal"],
        [
            ["Migración", "backend/alembic/versions/20260810_0033_workspace_revision_operational_hardening.py"],
            ["Backend", "backend/app/modules/enterprise_structure/revisions.py"],
            ["Permisos", "backend/app/modules/enterprise_structure/permissions.py"],
            ["Contrato HTTP", "backend/app/modules/enterprise_structure/router_admin.py"],
            ["Frontend", "frontend/src/features/enterprise-structure/components/WorkspaceRevisionManager.tsx"],
            ["PostgreSQL E2E", "backend/tests/postgres/test_enterprise_structure_postgres_e2e.py"],
            ["Benchmark", "backend/tests/postgres/snapshot_benchmark.py"],
            ["Entorno", "docker-compose.gate04h.yml"],
            ["CI", ".github/workflows/ci.yml"],
            ["Evidencia", "artifacts/enterprise_structure/gate04h/"],
        ],
        [2520, 6840],
    )
    add_heading(doc, "Artifacts mínimos", level=2)
    add_body(
        doc,
        "environment_health.json, production_invariance.json, technical_hardening.json, postgres_e2e_results.json, "
        "postgres_e2e_log.txt, concurrency_results.json, sod_matrix.json, sod_tests.json, snapshot_benchmark.json, "
        "snapshot_benchmark.csv y regression_results.json.",
    )

    add_page_break(doc)
    add_heading(doc, "12. Riesgos pendientes y plan de cierre")
    add_table(
        doc,
        ["Riesgo", "Impacto", "Acción de cierre", "Prioridad"],
        [
            ["Espacio libre < 10 GB", "Bloquea pruebas pesadas", "Liberar ≥10 GB antes de ejecutar 04H", "Alta"],
            ["Imagen final no reconstruida", "Runtime no persistente al recreate", "Rebuild API/worker/beat con espacio", "Alta"],
            ["E2E PostgreSQL no ejecutado", "Concurrencia/SoD sin evidencia local", "Ejecutar compose 04H o CI", "Alta"],
            ["Benchmark no medido", "Decisión Snapshot pendiente", "Completar 100/1k/10k", "Alta"],
            ["Approver/Publisher sin asignar", "Publicación bloqueada por diseño", "Asignar usuarios distintos", "Media"],
        ],
        [2520, 2520, 3240, 1080],
    )
    add_heading(doc, "Secuencia recomendada", level=2)
    for item in (
        "Liberar y verificar al menos 10 GB en C:.",
        "Reconstruir api, worker y beat desde el repositorio; verificar Alembic 0033.",
        "Asignar Structure Approver y Structure Publisher a usuarios distintos con scope organization.",
        "Ejecutar docker compose -f docker-compose.gate04h.yml up --build --abort-on-container-exit --exit-code-from gate04h-test.",
        "Revisar artifacts del E2E, benchmark y regresión completa; confirmar invariancia del baseline real.",
        "Cambiar a READY_FOR_PROJECT_CREATOR solo si todos los criterios 04H resultan PASS.",
    ):
        add_numbered(doc, item)
    add_callout(
        doc,
        "Estado final",
        "HARDENING_REQUIRED. Gate 05 no se inicia automáticamente.",
        fill=PALE_RED,
        border=RED,
    )
    return doc


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document = build_report()
    document.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
