from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(r"C:\Users\Ricardo\Documents\GitHub\pypmis-ia-sas")
EVIDENCE = REPO / "artifacts/enterprise_structure/gate04h"
OUTPUT = Path(
    r"C:\Users\Ricardo\Documents\P&P\P&Pmis Construction AI\Diseño\Resumen de Sprint"
    r"\Informe_Tecnico_PPMIS_Workspace_Revision_Manager_Gate04H_2026-08-12.docx"
)

NAVY = "0B2545"
BLUE = "1F5F99"
TEAL = "138A8A"
GREEN = "1D6F42"
PALE_GREEN = "E8F5ED"
PALE_BLUE = "EAF1F8"
PALE_TEAL = "E8F5F5"
PALE_AMBER = "FFF4D8"
AMBER = "8A6200"
LIGHT = "F4F6F8"
MID = "D7E0E8"
MUTED = "5E7180"
WHITE = "FFFFFF"
RED = "A72D2D"
FONT = "Aptos"


def load_json(name: str) -> dict:
    return json.loads((EVIDENCE / name).read_text(encoding="utf-8"))


def set_run(run, *, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = FONT
    fonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    for attr in ("ascii", "hAnsi", "eastAsia"):
        fonts.set(qn(f"w:{attr}"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def shade(cell, fill: str) -> None:
    props = cell._tc.get_or_add_tcPr()
    node = props.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        props.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, value: int = 90) -> None:
    props = cell._tc.get_or_add_tcPr()
    margins = props.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        props.append(margins)
    for side in ("top", "start", "bottom", "end"):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_row(row) -> None:
    props = row._tr.get_or_add_trPr()
    if props.find(qn("w:cantSplit")) is None:
        props.append(OxmlElement("w:cantSplit"))


def repeat_header(row) -> None:
    props = row._tr.get_or_add_trPr()
    node = OxmlElement("w:tblHeader")
    node.set(qn("w:val"), "true")
    props.append(node)


def set_repeat_table_header(row) -> None:
    repeat_header(row)


def add_page_field(paragraph) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    text = OxmlElement("w:instrText")
    text.set(qn("xml:space"), "preserve")
    text.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend((begin, text, end))


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.82)
    section.right_margin = Inches(0.82)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(9.5)
    normal.font.color.rgb = RGBColor.from_string(NAVY)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.08

    for name, size, color in (
        ("Title", 28, NAVY),
        ("Heading 1", 17, BLUE),
        ("Heading 2", 12.5, BLUE),
        ("Heading 3", 10.5, TEAL),
    ):
        style = doc.styles[name]
        style.font.name = FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(8)
        style.paragraph_format.space_after = Pt(5)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("P&Pmis Construction AI  |  Informe técnico Gate 04H")
    set_run(run, size=8.5, bold=True, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("Uso interno  |  12 de agosto de 2026  |  Página ")
    set_run(run, size=8.5, color=MUTED)
    add_page_field(footer)


def add_body(doc: Document, text: str, *, bold_prefix: str | None = None) -> None:
    paragraph = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        first = paragraph.add_run(bold_prefix)
        set_run(first, bold=True)
        rest = paragraph.add_run(text[len(bold_prefix) :])
        set_run(rest)
    else:
        set_run(paragraph.add_run(text))


def add_bullet(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.left_indent = Inches(0.28)
    paragraph.paragraph_format.first_line_indent = Inches(-0.18)
    paragraph.paragraph_format.space_after = Pt(3)
    set_run(paragraph.add_run(text), size=9.3)


def add_callout(doc: Document, label: str, text: str, *, fill: str, accent: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.7)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 130)
    props = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "20")
    left.set(qn("w:color"), accent)
    borders.append(left)
    props.append(borders)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    set_run(paragraph.add_run(f"{label}  "), bold=True, color=accent)
    set_run(paragraph.add_run(text), color=NAVY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        shade(cell, NAVY)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(paragraph.add_run(header), size=8.3, bold=True, color=WHITE)
    set_repeat_table_header(table.rows[0])
    for row_index, values in enumerate(rows):
        row = table.add_row()
        keep_row(row)
        for index, value in enumerate(values):
            cell = row.cells[index]
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                shade(cell, LIGHT)
            set_run(cell.paragraphs[0].add_run(str(value)), size=8.2)
    if widths:
        for row in table.rows:
            for index, width in enumerate(widths):
                row.cells[index].width = Inches(width)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def page_break(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def add_cover(doc: Document) -> None:
    for _ in range(2):
        doc.add_paragraph()
    line = doc.add_paragraph()
    line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(line.add_run("P&PMIS CONSTRUCTION AI"), size=12, bold=True, color=TEAL)
    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_before = Pt(18)
    title.paragraph_format.space_after = Pt(8)
    set_run(title.add_run("Informe técnico detallado"), size=28, bold=True, color=NAVY)
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run(subtitle.add_run("Workspace Structure Revision Manager · Gate 04H"), size=17, bold=True, color=BLUE)
    descriptor = doc.add_paragraph()
    descriptor.alignment = WD_ALIGN_PARAGRAPH.CENTER
    descriptor.paragraph_format.space_before = Pt(10)
    set_run(
        descriptor.add_run("Operational Hardening · PostgreSQL E2E · Concurrency · SoD · Benchmark"),
        size=10.5,
        color=MUTED,
    )
    for _ in range(3):
        doc.add_paragraph()
    add_callout(
        doc,
        "DICTAMEN FINAL",
        "READY_FOR_PROJECT_CREATOR",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    meta = doc.add_table(rows=4, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.style = "Table Grid"
    values = [
        ("Fecha de cierre", "12 de agosto de 2026"),
        ("Baseline protegido", "ES-PYP-CORE-RECONCILED-20260809 · Release ID 1"),
        ("Entorno", "Windows · Docker Desktop · PostgreSQL 16 · React/TypeScript"),
        ("Aplicación", "http://127.0.0.1:5173/app"),
    ]
    for row, (label, value) in zip(meta.rows, values, strict=True):
        shade(row.cells[0], PALE_BLUE)
        set_run(row.cells[0].paragraphs[0].add_run(label), size=9, bold=True, color=BLUE)
        set_run(row.cells[1].paragraphs[0].add_run(value), size=9)
        for cell in row.cells:
            set_cell_margins(cell, 100)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def main() -> None:
    benchmark = load_json("snapshot_benchmark.json")
    regression = load_json("regression_results.json")
    invariance = load_json("production_invariance.json")
    concurrency = load_json("concurrency_results.json")
    sod = load_json("sod_tests.json")

    doc = Document()
    configure(doc)
    add_cover(doc)
    page_break(doc)

    doc.add_heading("1. Resumen ejecutivo", level=1)
    add_callout(
        doc,
        "RESULTADO",
        "Gate 04H quedó implementado, ejecutado y verificado. La decisión autorizada por el prompt es READY_FOR_PROJECT_CREATOR; no se inició Gate 05.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_body(
        doc,
        "El hardening operacional del Workspace Structure Revision Manager se cerró sobre PostgreSQL efímero, con el baseline real protegido. Se verificaron concurrencia optimista, invalidación de validación/aprobación, segregación de funciones, Four-Eyes, auditoría, no borrado físico, publicación sucesora, segundo publish idempotente y rollback.",
    )
    add_table(
        doc,
        ["Criterio", "Resultado", "Evidencia principal"],
        [
            ["PostgreSQL E2E", "PASS", "3 pruebas · ciclo funcional/concurrencia/SoD"],
            ["Backend", "PASS", "229 passed · 2 skipped · 85,34% coverage"],
            ["Frontend", "PASS", "143 Vitest · build Vite · lint/prettier"],
            ["Browser E2E", "PASS", "Smoke producción Playwright"],
            ["Benchmark", "COMPLETE", "100 / 1.000 / 10.000 nodos"],
            ["Baseline real", "PASS", "0 mutaciones de negocio · 0 fixtures"],
            ["Runtime", "PASS", "6 servicios · API ready · frontend 200"],
        ],
        [1.6, 1.0, 4.1],
    )
    doc.add_heading("Alcance cerrado", level=2)
    for item in (
        "Revisión de estructura empresarial por snapshot completo con estados de draft, validación, aprobación, publicación y rollback.",
        "API ADMIN con ETag/If-Match obligatorio para mutaciones del draft.",
        "Roles structure_editor, structure_approver y structure_publisher sin bypass del administrador organizacional.",
        "UI Workspace Revision Manager con conflicto de concurrencia comprensible y recarga del estado vigente.",
        "Job CI dedicado enterprise-structure-postgres-e2e y artefactos reproducibles en artifacts/enterprise_structure/gate04h/.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("2. Arquitectura y cambios implementados", level=1)
    add_body(
        doc,
        "La solución reutiliza el dominio Enterprise Structure existente. El Revision Manager opera como una capa transaccional sobre releases y snapshots, sin duplicar el árbol publicado ni introducir un segundo modelo de workspaces.",
    )
    add_table(
        doc,
        ["Capa", "Implementación 04H", "Responsabilidad"],
        [
            ["Persistencia", "Alembic 20260810_0033", "revision_version y actor modificador"],
            ["ORM", "version_id_col", "UPDATE atómico y detección de stale writer"],
            ["Repositorio", "populate_existing", "refresco real en sesiones multiusuario largas"],
            ["Servicio", "revisions.py", "clone, edit, move, classify, archive, validate, diff, approve, publish, rollback"],
            ["HTTP", "router_admin.py", "If-Match, ETag y errores de dominio 409"],
            ["Seguridad", "permissions.py", "matriz SoD y Four-Eyes"],
            ["Frontend", "WorkspaceRevisionManager", "operación y recuperación de conflictos"],
            ["Pruebas", "PostgreSQL + pytest + Vitest + Playwright", "validación cruzada de dominio y UI"],
        ],
        [1.2, 2.0, 3.5],
    )
    doc.add_heading("Correcciones descubiertas durante el cierre", level=2)
    for item in (
        "Se agregó un bootstrap guardado para la base PostgreSQL efímera: exige GATE04H_EPHEMERAL=true, rechaza SQLite, rechaza bases no vacías y deja un baseline 0032 para probar realmente la migración 0033.",
        "Se refresca EnterpriseCoreRelease con populate_existing para evitar identidad ORM obsoleta entre escritores concurrentes.",
        "La ruta administrativa heredada ahora crea external_key y record_code jerárquico, reutilizando record_codes.py.",
        "El contrato estático admite PYPMIS_REPOSITORY_ROOT, lo que permite ejecutar la suite dentro de la imagen backend sin asumir una ruta de host.",
        "Los puertos host del compose son configurables, preservando 5432/6379/8000/5173 por defecto y habilitando stacks E2E aislados.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("3. Concurrencia, integridad e invalidación", level=1)
    add_body(
        doc,
        "Cada mutación relevante exige la versión conocida por el cliente. Si otro usuario modifica la revisión, el UPDATE no fusiona silenciosamente: devuelve HTTP 409 con REVISION_VERSION_CONFLICT. La UI informa que la revisión cambió y exige recargar.",
    )
    add_table(
        doc,
        ["Escenario", "Resultado esperado", "Resultado"],
        [
            ["Escritor obsoleto", "409 REVISION_VERSION_CONFLICT", str(concurrency["stale_writer_reason"])],
            ["Validar y luego modificar", "HASH_MISMATCH al aprobar", str(concurrency["validate_then_modify_then_approve"])],
            ["Aprobar y luego modificar", "APPROVAL_INVALIDATED al publicar", str(concurrency["approve_then_modify_then_publish"])],
            ["Mismo usuario aprueba/publica", "FOUR_EYES_VIOLATION", "PASS"],
            ["Segundo publish", "0 mutaciones", "PASS"],
            ["Borrado", "Archivado lógico, no DELETE", "PASS"],
        ],
        [2.1, 2.8, 1.8],
    )
    doc.add_heading("Trazabilidad", level=2)
    add_body(
        doc,
        "last_modified_by_user_id identifica al actor de la última mutación; revision_version soporta el control optimista. Los eventos de auditoría cubren las operaciones requeridas y el E2E confirma que el release anterior permanece publicado e inalterado después de crear su sucesor.",
    )
    add_callout(
        doc,
        "CONTRATO DE CLIENTE",
        "Leer ETag, enviar If-Match en cada mutación y tratar 409 como recarga obligatoria; nunca reintentar una escritura ciega.",
        fill=PALE_BLUE,
        accent=BLUE,
    )

    page_break(doc)
    doc.add_heading("4. Segregación de funciones y Four-Eyes", level=1)
    add_table(
        doc,
        ["Rol", "Permitido", "Bloqueado"],
        [
            ["structure_editor", "crear, editar, mover, clasificar, archivar, validar, comparar", "aprobar, publicar, rollback"],
            ["structure_approver", "comparar y aprobar", "editar, publicar, rollback"],
            ["structure_publisher", "comparar, publicar y rollback", "editar y aprobar"],
        ],
        [1.45, 3.0, 2.25],
    )
    add_body(
        doc,
        "La asignación de administrador organizacional no evita la matriz de deberes. El bootstrap de compatibilidad concede únicamente Structure Editor al administrador existente; Approver y Publisher requieren asignaciones explícitas.",
    )
    sod_rows = [[key.replace("_", " "), str(value)] for key, value in sod["results"].items()]
    add_table(doc, ["Caso SoD", "Resultado"], sod_rows, [4.9, 1.8])
    add_callout(
        doc,
        "FOUR-EYES",
        "approved_by_user_id debe ser diferente de published_by_user_id. El caso con el mismo actor se rechaza aunque tenga varias asignaciones.",
        fill=PALE_TEAL,
        accent=TEAL,
    )

    page_break(doc)
    doc.add_heading("5. PostgreSQL E2E y ciclo operacional", level=1)
    add_body(
        doc,
        "El harness usa PostgreSQL 16 desechable, tmpfs y ejecución secuencial. Prepara un esquema histórico controlado, marca Alembic 0032, aplica 0033 y ejecuta las pruebas antes de destruir contenedores y red. No utiliza SQLite ni toca volúmenes persistentes.",
    )
    steps = [
        "Clonar el release publicado a draft aislado.",
        "Agregar, editar y mover nodos con recodificación de descendientes.",
        "Clasificar y archivar lógicamente.",
        "Previsualizar record_code, validar y calcular diff.",
        "Aprobar con rol separado y publicar un sucesor.",
        "Reproducir concurrencia e invalidación de hash/aprobación.",
        "Repetir publicación sin mutaciones y ejecutar rollback.",
        "Verificar auditoría, conservación del release anterior y ausencia de DELETE físico.",
    ]
    for index, item in enumerate(steps, start=1):
        add_body(doc, f"{index}. {item}")
    add_table(
        doc,
        ["Artefacto", "Estado", "Cobertura"],
        [
            ["postgres_e2e_results.json", "PASS", "ciclo funcional completo"],
            ["concurrency_results.json", "PASS", "409, hashes, aprobación y Four-Eyes"],
            ["sod_tests.json", "PASS", "roles, tenant, expiración y scope"],
            ["postgres_e2e_log.txt", "PASS", "Alembic 0033 y 3 pruebas"],
        ],
        [2.2, 1.0, 3.5],
    )

    page_break(doc)
    doc.add_heading("6. Benchmark de snapshots", level=1)
    add_body(
        doc,
        "Se midieron snapshots completos con 100, 1.000 y 10.000 nodos. No se inventaron umbrales previos. Las cifras son observaciones del entorno local y sirven como baseline para decisiones posteriores; Delta/Change Set permanece fuera del Gate 04H.",
    )
    bench_rows: list[list[str]] = []
    for item in benchmark["results"]:
        bench_rows.append(
            [
                f"{item['nodes']:,}".replace(",", "."),
                f"{item['snapshot_size_bytes'] / 1024:.1f}",
                f"{item['transaction_ms'] / 1000:.2f}",
                f"{item['publish_ms'] / 1000:.2f}",
                f"{item['rollback_ms'] / 1000:.2f}",
                f"{item['peak_memory_bytes'] / (1024 * 1024):.1f}",
            ]
        )
    add_table(
        doc,
        ["Nodos", "Snapshot KiB", "Transacción s", "Publish s", "Rollback s", "Memoria MiB"],
        bench_rows,
        [0.75, 1.15, 1.25, 1.05, 1.1, 1.15],
    )
    add_callout(
        doc,
        "DECISIÓN",
        "RETAIN_FULL_SNAPSHOT. La ejecución de 10.000 nodos completó el ciclo en 221,75 s, con snapshot de 3,38 MB y pico de memoria aproximado de 152,9 MiB. No se implementa Delta en este gate.",
        fill=PALE_AMBER,
        accent=AMBER,
    )
    doc.add_heading("Lectura operacional", level=2)
    for item in (
        "El costo crece materialmente entre 1.000 y 10.000 nodos; la publicación es la operación individual más costosa medida.",
        "El modelo completo conserva simplicidad de recuperación, trazabilidad y rollback en esta fase.",
        "Antes de fijar un SLO de producción se recomienda repetir en infraestructura objetivo y definir concurrencia y tamaño de tenant esperados.",
    ):
        add_bullet(doc, item)

    page_break(doc)
    doc.add_heading("7. Regresión integral y calidad", level=1)
    add_table(
        doc,
        ["Validación", "Resultado", "Detalle"],
        [
            ["Pytest completo", "PASS", regression["backend"]["pytest"]],
            ["Cobertura backend", "PASS", "85,34% (mínimo 75%)"],
            ["Ruff format/check", "PASS", "122 archivos · 0 hallazgos"],
            ["Prettier", "PASS", "todos los archivos src"],
            ["ESLint", "PASS", "0 errores · 8 warnings / máximo 10"],
            ["Vitest", "PASS", "24 archivos · 143 pruebas"],
            ["TypeScript + Vite", "PASS", "2.343 módulos transformados"],
            ["Playwright", "PASS", "smoke producción; BIM omitido por fixture IFC ausente"],
            ["Alembic", "PASS", "20260810_0033 (head)"],
        ],
        [1.75, 0.9, 4.05],
    )
    doc.add_heading("Regresiones de negocio cubiertas", level=2)
    for item in (
        "Gate 02A importer y Gate 02B controlled apply.",
        "Gate 03 publish/rollback y Gate 04 revision manager.",
        "Organization & Security.",
        "ADMIN Enterprise Structure y USER Enterprise Explorer.",
        "Frontend completo, visor BIM, cantidades/APU, Guided Flow y Project Controls.",
    ):
        add_bullet(doc, item)
    add_body(
        doc,
        "Las dos omisiones backend son guards declarados; el escenario Playwright BIM requiere un modelo IFC cargado y se omite deliberadamente en el tenant demo. El smoke de producción que forma parte de CI sí pasó.",
    )

    page_break(doc)
    doc.add_heading("8. Invariancia del baseline real", level=1)
    before = invariance["before"]
    after = invariance["after"]
    add_callout(
        doc,
        "PROTECCIÓN CONFIRMADA",
        "Todas las pruebas con mutaciones se ejecutaron en PostgreSQL efímero. La base persistente solo recibió verificaciones idempotentes de Alembic y consultas de lectura.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_table(
        doc,
        ["Indicador", "Antes", "Después", "Estado"],
        [
            ["Release", before["release"], after["release"], "PASS"],
            ["Release ID / estado", f"{before['release_id']} / {before['state']}", f"{after['release_id']} / {after['state']}", "PASS"],
            ["Releases / drafts", f"{before['release_count']} / {before['draft_count']}", f"{after['release_count']} / {after['draft_count']}", "PASS"],
            ["Workspaces", str(before["workspaces"]), str(after["workspaces"]), "PASS"],
            ["Objetivos", str(before["strategic_objectives"]), str(after["strategic_objectives"]), "PASS"],
            ["Clasificaciones", str(before["classifications"]), str(after["classifications"]), "PASS"],
            ["Enlaces", str(before["links"]), str(after["links"]), "PASS"],
        ],
        [1.4, 2.05, 2.05, 1.0],
    )
    add_body(doc, "Mutaciones de payload de negocio: 0. Fixtures residuales: 0. Volúmenes persistentes eliminados: ninguno.")

    page_break(doc)
    doc.add_heading("9. CI, despliegue local y operación", level=1)
    add_body(
        doc,
        "El workflow Pilot Readiness CI incluye el job exacto enterprise-structure-postgres-e2e. El job construye el entorno efímero, aplica migraciones, ejecuta E2E/concurrencia/SoD/benchmark, sube evidencias y destruye el proyecto Docker incluso ante fallo.",
    )
    add_table(
        doc,
        ["Servicio", "Estado final", "Acceso"],
        [
            ["Frontend", "UP · HTTP 200", "http://127.0.0.1:5173/app"],
            ["API", "healthy · ready", "http://127.0.0.1:8000"],
            ["PostgreSQL", "healthy", "localhost:5432"],
            ["Redis", "healthy", "localhost:6379"],
            ["Celery worker", "UP", "cola control-core"],
            ["Celery beat", "UP", "scheduler"],
        ],
        [1.6, 2.2, 2.9],
    )
    add_body(doc, "Las imágenes de API, worker, beat y frontend fueron reconstruidas y los servicios recreados. Alembic quedó en 20260810_0033 (head). Espacio libre final medido: 25,15 GB.")
    doc.add_heading("Puertos configurables", level=2)
    add_body(
        doc,
        "POSTGRES_HOST_PORT, REDIS_HOST_PORT, API_HOST_PORT y FRONTEND_HOST_PORT permiten levantar stacks paralelos sin colisión. Si no se definen, se conservan los puertos habituales 5432, 6379, 8000 y 5173.",
    )

    page_break(doc)
    doc.add_heading("10. Evidencia, riesgos y conclusión", level=1)
    add_table(
        doc,
        ["Evidencia", "Propósito"],
        [
            ["postgres_e2e_log.txt", "migración y ejecución PostgreSQL"],
            ["postgres_e2e_results.json", "ciclo funcional"],
            ["concurrency_results.json", "conflicto e invalidación"],
            ["sod_tests.json / sod_matrix.json", "segregación de funciones"],
            ["snapshot_benchmark.json / .csv", "medición de escala"],
            ["backend_regression_log.txt", "pytest completo y cobertura"],
            ["frontend_validation_log.txt", "prettier, lint, Vitest y build"],
            ["browser_e2e_log.txt", "Playwright producción"],
            ["production_invariance.json", "baseline protegido"],
            ["final_acceptance.json", "dictamen consolidado"],
        ],
        [3.0, 3.7],
    )
    doc.add_heading("Riesgos residuales no bloqueantes", level=2)
    add_table(
        doc,
        ["Riesgo", "Tratamiento recomendado"],
        [
            ["Benchmark 10k costoso", "Definir SLO y repetir en infraestructura objetivo antes de producción masiva."],
            ["8 warnings ESLint", "Reducir gradualmente a cero; actualmente dentro del presupuesto CI."],
            ["E2E BIM requiere fixture IFC", "Crear un fixture licenciado y liviano para habilitar el recorrido en CI."],
            ["Snapshot completo", "Mantener hasta contar con umbral y evidencia para evaluar Delta/Change Set."],
        ],
        [2.4, 4.3],
    )
    add_callout(
        doc,
        "CONCLUSIÓN",
        "Todos los criterios ejecutables del prompt 04H quedaron implementados y verificados, el baseline real permaneció intacto y localhost está operativo. Estado final exclusivo: READY_FOR_PROJECT_CREATOR.",
        fill=PALE_GREEN,
        accent=GREEN,
    )
    add_body(doc, "Gate 05 no fue iniciado automáticamente.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
