from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUT = REPO / "outputs" / "technical_reports" / "Informe_Tecnico_Gate_06C_Physical_Workspace_Initialization_Activation_20260813.docx"

# Preset: standard_business_brief. First-page pattern: memo_masthead.
NAVY = "0D2A3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "16222B"
MUTED = "5B6870"
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
PALE_TEAL = "EAF7F6"


def set_run(run, *, size: float = 11, color: str = INK, bold: bool | None = None, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(item, widths: list[int], indent: int = 120) -> None:
    item.autofit = False
    total = sum(widths)
    tbl_pr = item._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = item._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in item.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int]) -> None:
    item = doc.add_table(rows=1, cols=len(headers))
    item.style = "Table Grid"
    item.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, header in enumerate(headers):
        cell = item.rows[0].cells[index]
        cell.text = header
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, color=WHITE, bold=True)
    set_repeat_header(item.rows[0])
    for row_index, values in enumerate(rows):
        cells = item.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            shade(cells[index], WHITE if row_index % 2 == 0 else LIGHT)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cells[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=9.1)
    set_table_geometry(item, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(doc: Document, title: str, body: str, fill: str = PALE_TEAL) -> None:
    item = doc.add_table(rows=1, cols=1)
    item.style = "Table Grid"
    cell = item.cell(0, 0)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    set_run(cell.paragraphs[0].add_run(title.upper()), size=9, color=NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run(run, size=9.5)
    set_repeat_header(item.rows[0])
    set_table_geometry(item, [9360], indent=160)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Title", 23, NAVY, 0, 4),
    ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(11)
    styles[name].paragraph_format.space_after = Pt(8)
    styles[name].paragraph_format.line_spacing = 1.167
    styles[name].paragraph_format.left_indent = Inches(0.5)
    styles[name].paragraph_format.first_line_indent = Inches(-0.25)

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(0)
set_run(header.add_run("P&Pmis Construction AI  |  Gate 06C Technical Report"), size=8.5, color=MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer.add_run("Controlled technical baseline  |  13 Aug 2026  |  Page "), size=8, color=MUTED)
page_run = footer.add_run()
set_run(page_run, size=8, color=MUTED)
for field_type, value in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
    if field_type:
        field = OxmlElement("w:fldChar")
        field.set(qn("w:fldCharType"), field_type)
        page_run._r.append(field)
    else:
        text = OxmlElement("w:instrText" if value == " PAGE " else "w:t")
        text.set(qn("xml:space"), "preserve")
        text.text = value
        page_run._r.append(text)

doc.core_properties.title = "Informe técnico Gate 06C - Physical Workspace Initialization & Activation"
doc.core_properties.subject = "PROPERTY, FACILITY y WAREHOUSE - preparación y activación gobernada"
doc.core_properties.author = "P&Pmis Construction AI Engineering"
doc.core_properties.comments = "Gate 06C cerrado; operación física profunda no iniciada."

doc.add_paragraph("INFORME TÉCNICO", style="Title")
doc.add_paragraph("Gate 06C · Physical Workspace Initialization & Activation", style="Subtitle")
for label, value in (
    ("Tenant", "P&P Ingeniería y Proyectos"),
    ("Fecha de cierre", "13 de agosto de 2026"),
    ("Alcance", "PROPERTY · FACILITY · WAREHOUSE"),
    ("Estado final", "READY_FOR_PHYSICAL_WORKSPACE_OPERATION"),
    ("Límite", "Sin Asset Manager, Inventory Manager ni operación física profunda"),
):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    set_run(paragraph.add_run(f"{label}: "), bold=True)
    set_run(paragraph.add_run(value))
callout(
    doc,
    "Resultado de cierre",
    "Se implementó un único motor parametrizado por workspace_type_code para inicializar, validar y activar "
    "PROPERTY, FACILITY y WAREHOUSE materializados por Gate 06B. Preview no persiste; Activation es una acción "
    "separada y atómica bajo segregación de funciones.",
    PALE_GREEN,
)

heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph(
    "Gate 06C convierte los Physical Workspaces PENDING en unidades preparadas y, sólo después de una activación "
    "independiente, en Workspaces ACTIVE. La implementación conserva enterprise_workspaces como identidad canónica, "
    "usa el snapshot exacto del template materializado, aplica defaults con precedencia explícita y mantiene "
    "trazabilidad de checklists, módulos, hashes, actores, eventos y versiones."
)
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Motor único", "PASS · Property / Facility / Warehouse"],
        ["Preview", "PASS · evaluación no persistente"],
        ["Start / Retry / Validate", "PASS · BLOCKED o READY_FOR_ACTIVATION"],
        ["Activation", "PASS · PENDING→ACTIVE en transacción separada"],
        ["Segregación", "PASS · activator distinto de requestor/initializer/modifier"],
        ["PostgreSQL E2E", "PASS · 31/31, incluida concurrencia y reversibilidad"],
        ["Regresión backend", "PASS · 114/114; 4 skips previstos"],
        ["Frontend", "PASS · 185/185; TypeScript, build, lint y format"],
        ["Baseline real", "PASS · 14 Workspaces, 0 físicos, 0 inicializaciones"],
    ],
    [2850, 6510],
)

heading(doc, "2. REUSE BEFORE CREATE")
bullets(
    doc,
    [
        "Se reutilizó el patrón probado de Project Workspace Initialization 05C sin duplicar sus semánticas específicas.",
        "Se consumen los Workspaces PENDING y snapshots materializados por el proceso físico 06B.",
        "Se reutilizan EnterpriseWorkspace, AdminConfiguration, SecurityAccessAssignment, SecurityEvent y WorkspaceModuleSetting.",
        "Se creó una única tabla física porque checklists, responsables, campos y módulos difieren del dominio Project.",
        "El release CORE publicado no se crea, revisa ni modifica durante inicialización o activación.",
    ],
)

heading(doc, "3. Alcance y exclusiones")
add_table(
    doc,
    ["Incluido", "Excluido"],
    [
        ["PROPERTY / FACILITY / WAREHOUSE PENDING", "REGION / DISTRICT / SITE / PROJECT"],
        ["Preview, Start, Retry, Validate, Activate", "LINEAR_ASSET / ASSET"],
        ["Checklists, defaults, acceso, módulos", "Activos, inventario, espacios, mantenimiento"],
        ["Overview e inventario USER MODE", "Operaciones inmobiliarias o de facilities"],
    ],
    [4680, 4680],
)

page_break(doc)
heading(doc, "4. Arquitectura implementada")
add_table(
    doc,
    ["Capa", "Responsabilidad"],
    [
        ["React USER MODE", "My Physical Workspaces, Overview, preview, retry, validate y activate"],
        ["FastAPI router", "6 contratos lifecycle, autorización RBAC y ETag/If-Match"],
        ["Servicio genérico", "Una fuente de verdad para los tres tipos físicos"],
        ["Persistencia", "Singleton tenant/workspace con snapshot, estados, hashes y actores"],
        ["Enterprise models", "Identidad, parent, clasificaciones, módulos y assignments"],
        ["PostgreSQL", "Locks, unicidad, rollback y transición atómica"],
    ],
    [2150, 7210],
)

heading(doc, "5. Modelo de persistencia")
doc.add_paragraph(
    "La migración 20260813_0038 agrega physical_workspace_initializations con unicidad tenant_id/workspace_id. "
    "El registro conserva workspace type, estado, snapshot exacto de template, initialization_version, revision_version, "
    "actores y timestamps, validation/checklist hashes, checklists común/específico, defaults, assignments, estados de "
    "módulo y diagnóstico de falla."
)
callout(
    doc,
    "Identidad inmutable",
    "Business Number, Record Code y external_key permanecen inmutables durante Start, Validate y Activate. La "
    "activación sólo cambia status y version del EnterpriseWorkspace, además del estado lifecycle.",
)

heading(doc, "6. Máquina de estados")
add_table(
    doc,
    ["Estado", "Transición / significado"],
    [
        ["NOT_STARTED", "No existe registro; Preview disponible y no persistente"],
        ["INITIALIZING", "Estado transitorio durante aplicación de defaults/asignaciones"],
        ["BLOCKED", "Uno o más controles obligatorios fallan; Workspace sigue PENDING"],
        ["READY_FOR_ACTIVATION", "Checklists y hashes válidos; requiere activator independiente"],
        ["ACTIVATED", "Workspace ACTIVE; activación registrada e idempotente"],
        ["FAILED", "Falla técnica persistida; Workspace PENDING y reintento permitido"],
    ],
    [2300, 7060],
)

heading(doc, "7. Checklist común")
common = [
    "workspace_identity_valid", "workspace_type_supported", "workspace_status_pending", "parent_valid",
    "business_number_valid", "record_code_valid", "external_key_valid", "template_assigned",
    "template_snapshot_valid", "responsible_assigned", "responsible_access_valid",
    "required_attributes_complete", "required_classifications_valid", "module_settings_valid",
    "tenant_scope_valid", "no_core_revision_required",
]
add_table(doc, ["Código", "Severidad"], [[item, "Blocking"] for item in common], [7000, 2360])

heading(doc, "8. Checklists por tipo")
add_table(
    doc,
    ["Tipo", "Controles"],
    [
        ["PROPERTY", "type, manager, tenure, legal, geographic, area y value"],
        ["FACILITY", "type, responsible, operational status, geographic, capacity, area, criticality y commissioning"],
        ["WAREHOUSE", "type, manager, geographic, storage capacity/unit, criticality y parent context"],
    ],
    [1900, 7460],
)
doc.add_paragraph(
    "Los campos configurados como obligatorios producen FAIL bloqueante. Los campos opcionales incompletos producen "
    "WARNING y no impiden READY_FOR_ACTIVATION. El backend, no la UI, decide severidad y evidencia."
)

page_break(doc)
heading(doc, "9. Defaults, responsable y módulos")
bullets(
    doc,
    [
        "Precedencia: valor explícito del request > default del template exacto > default tenant/workspace type.",
        "Los defaults aplicados conservan value y source; los explícitos nunca se sobrescriben.",
        "El Responsible debe ser activo y tenant-scoped; se garantiza una assignment mínima workspace-scoped.",
        "Módulos habilitados existentes requieren definición PUBLISHED y WorkspaceModuleSetting enabled.",
        "Módulos futuros quedan PLANNED, operational_module_created=false y sin persistencia operativa profunda.",
    ],
)

heading(doc, "10. Preview, Start, Validate y Activate")
add_table(
    doc,
    ["Acción", "Contrato"],
    [
        ["Preview", "Evalúa snapshot/checklists/módulos; persisted=false; 0 mutaciones"],
        ["Start", "Crea singleton, aplica defaults/acceso, valida y termina BLOCKED o READY"],
        ["Retry", "Reutiliza el singleton FAILED/BLOCKED; no duplica assignments"],
        ["Validate", "Recalcula checklists/hashes con If-Match; limpia fallas superadas"],
        ["Activate", "Bloquea filas, revalida hashes y SoD, cambia PENDING→ACTIVE y registra evento"],
        ["Repetición", "Start READY y Activate ACTIVE retornan respuesta idempotente"],
    ],
    [1850, 7510],
)

heading(doc, "11. Concurrencia y rollback")
doc.add_paragraph(
    "Start y Activate usan row locks y revision_version. La unicidad tenant/workspace evita inicializaciones dobles; "
    "un conflicto concurrente devuelve el registro ya creado. Activation revalida validation_hash y checklist_hash "
    "antes del cambio de estado. Una falla inyectada después de modificar el Workspace revierte tanto status como "
    "estado lifecycle y registra activation_failed en una transacción controlada."
)
callout(
    doc,
    "Separación de funciones",
    "El activator no puede ser el requestor de Gate 06B, el actor que inició la inicialización ni el último actor que "
    "la modificó. organization_admin conserva permisos, pero no evita el control SoD.",
    PALE_AMBER,
)

heading(doc, "12. RBAC")
add_table(
    doc,
    ["Rol", "Capacidad"],
    [
        ["physical_workspace_initializer", "read + initialization.execute"],
        ["physical_workspace_activator", "read + activation.execute"],
        ["physical_workspace_responsible", "read en su Workspace asignado"],
        ["organization_admin", "capacidades lifecycle sujetas a SoD"],
    ],
    [3200, 6160],
)

heading(doc, "13. Auditoría")
bullets(
    doc,
    [
        "physical_workspace.initialization_started",
        "physical_workspace.initialization_validated",
        "physical_workspace.initialization_blocked / ready_for_activation / initialization_failed",
        "physical_workspace.activated / activation_failed",
        "Metadatos: tenant, workspace, type, Business/Record Code, actor, responsible, template/revision, estados, hashes, blockers, warnings, enabled/planned modules y timestamp.",
    ],
)

page_break(doc)
heading(doc, "14. APIs implementadas")
add_table(
    doc,
    ["Método", "Ruta"],
    [
        ["GET", "/api/v1/physical-workspaces"],
        ["GET", "/api/v1/physical-workspaces/{id}/initialization"],
        ["POST", "/api/v1/physical-workspaces/{id}/initialization/preview"],
        ["POST", "/api/v1/physical-workspaces/{id}/initialization/start"],
        ["POST", "/api/v1/physical-workspaces/{id}/initialization/validate"],
        ["POST", "/api/v1/physical-workspaces/{id}/activate"],
        ["GET", "/api/v1/physical-workspaces/{id}/overview"],
    ],
    [1500, 7860],
)

heading(doc, "15. USER MODE")
bullets(
    doc,
    [
        "Enterprise Explorer incorpora la acción My Physical Workspaces.",
        "Inventario filtrable por Workspace Type, Workspace Status e Initialization Status; el API soporta ocho filtros.",
        "Overview muestra identidad, parent, responsible, snapshot, avance, blockers, warnings y activación.",
        "Initialization Checklist separa controles comunes y específicos por tipo.",
        "Module Readiness distingue READY de PLANNED y explica que no existe módulo operacional profundo.",
        "Las acciones visibles se derivan de can_initialize/can_activate calculadas por el backend.",
        "Explorer mantiene indicadores PENDING/ACTIVE en árbol y tabla.",
    ],
)

heading(doc, "16. Evidencia de pruebas")
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Gate 06C PostgreSQL E2E", "31 passed · concurrencia + rollback + migration cycle"],
        ["Regresión backend 05C–06C", "114 passed · 4 skipped previstos"],
        ["Frontend full suite", "29 files · 185 passed"],
        ["Ruff scope Gate 06C", "PASS · check y format"],
        ["Frontend Prettier", "PASS"],
        ["Frontend ESLint", "PASS · 0 errores; 9 warnings históricos permitidos"],
        ["TypeScript / Vite", "PASS · producción compilada"],
        ["Browser localhost", "PASS · login, Explorer, My Physical Workspaces, filtros, 0 errores consola"],
    ],
    [2900, 6460],
)

heading(doc, "17. Migración y reversibilidad")
bullets(
    doc,
    [
        "Alembic 20260813_0038 depende de 20260812_0037 y es aditiva.",
        "El ciclo PostgreSQL upgrade 0037→0038, downgrade 0038→0037 y re-upgrade fue exitoso.",
        "La migración no crea ni altera registros de inicialización en el tenant persistente.",
        "El entorno principal quedó en 20260813_0038 (head).",
    ],
)

heading(doc, "18. Protección del baseline real")
add_table(
    doc,
    ["Control", "Estado al cierre"],
    [
        ["CORE publicado", "ES-PYP-CORE-RECONCILED-20260809 · sin cambio"],
        ["Enterprise Workspaces", "14"],
        ["Project Workspaces", "1"],
        ["Physical Workspaces", "0"],
        ["Project/Physical initializations", "0 / 0"],
        ["Project/Physical creation requests", "0 / 0"],
        ["Physical Templates", "5 DRAFT / 0 PUBLISHED"],
        ["Physical Creation Policies", "3 DRAFT"],
        ["Physical Number sequences", "todas next_value = 1"],
        ["Project Number sequence", "next_value = 1"],
    ],
    [3450, 5910],
)

page_break(doc)
heading(doc, "19. Riesgos y deuda controlada")
bullets(
    doc,
    [
        "Las plantillas y policies físicas reales siguen DRAFT; por ello no hay Workspaces reales elegibles para inicializar.",
        "Los nueve warnings ESLint preexistentes están en BIM Viewer/Project Creation y dentro del presupuesto CI; Gate 06C no agrega warnings.",
        "El motor no materializa Asset Manager ni Inventory Manager: sólo conserva readiness PLANNED.",
        "Los fixtures de aceptación son efímeros y se ejecutan en PostgreSQL tmpfs o SQLite aislado.",
    ],
)

heading(doc, "20. Matriz de aceptación agrupada")
criteria = [
    "Motor genérico para PROPERTY / FACILITY / WAREHOUSE",
    "Exclusión de tipos no elegibles y exigencia PENDING",
    "Preview no persistente y snapshot exacto sin auto-upgrade",
    "Defaults con precedencia y acceso mínimo responsable",
    "Checklists común + específico con blockers/warnings",
    "Módulos existentes READY y futuros PLANNED sin operación",
    "Retry/idempotencia/ETag/hashes/rollback/concurrencia",
    "Activation separada con Four-Eyes/SoD",
    "RBAC, SecurityEvent y tenant isolation",
    "My Physical Workspaces, Overview y estados Explorer",
    "Migration reversible y PostgreSQL E2E",
    "Baseline real sin mutaciones funcionales",
]
add_table(doc, ["Criterio", "Resultado"], [[item, "PASS"] for item in criteria], [7850, 1510])
callout(doc, "Estado final exclusivo", "READY_FOR_PHYSICAL_WORKSPACE_OPERATION", PALE_GREEN)
doc.add_paragraph(
    "Gate 06C termina aquí. Asset Manager, Inventory Manager, mantenimiento, spaces, utilities, inventario, activos y "
    "operación profunda de Property/Facility/Warehouse no fueron iniciados."
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
