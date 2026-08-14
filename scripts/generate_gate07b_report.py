from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
OUTPUT = (
    REPO
    / "outputs"
    / "technical_reports"
    / "Informe_Tecnico_Gate_07B_Project_Proposal_20260813.docx"
)

# Preset: standard_business_brief. Header archetype: memo_masthead.
NAVY = "0D2A3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "16222B"
MUTED = "5B6870"
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
PALE_TEAL = "EAF7F6"


def set_run(run, *, size=11, color=INK, bold=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", indent)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row = table.rows[0]
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_row._tr.get_or_add_trPr().append(header_flag)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, NAVY)
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=9.1)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(doc, title, body, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_flag)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    set_run(cell.paragraphs[0].add_run(title.upper()), size=9, color=NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run(run, size=9.5)
    set_table_geometry(table, [9360], indent=160)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def bullets(doc, values):
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("P&Pmis Construction AI  |  "), size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, end])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Title", 23, NAVY, 0, 4),
    ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(11)
    styles[name].paragraph_format.space_after = Pt(6)
    styles[name].paragraph_format.line_spacing = 1.10
    styles[name].paragraph_format.left_indent = Inches(0.5)
    styles[name].paragraph_format.first_line_indent = Inches(-0.25)

header = section.header.paragraphs[0]
set_run(
    header.add_run("P&Pmis Construction AI  |  Gate 07B Technical Report"),
    size=8.5,
    color=MUTED,
    bold=True,
)
add_page_number(section.footer.paragraphs[0])

title = doc.add_paragraph(style="Title")
set_run(title.add_run("Informe técnico — Gate 07B"), size=23, color=NAVY, bold=True)
subtitle = doc.add_paragraph(style="Subtitle")
set_run(
    subtitle.add_run("Project Proposal Foundation & Lifecycle"),
    size=13,
    color=MUTED,
)

add_table(
    doc,
    ["Campo", "Detalle"],
    [
        ["Aplicación", "P&Pmis Construction AI"],
        ["Módulo", "Enterprise Strategy Manager / Idea & Demand Manager"],
        ["Submódulo", "Project Proposal"],
        ["Baseline", "Gate 07A Idea Lifecycle aceptado y preparado"],
        ["Fecha de corte", "13 de agosto de 2026"],
        ["Estado", "Implementado, probado y operativo en localhost"],
    ],
    [2400, 6960],
)

callout(
    doc,
    "Resultado ejecutivo",
    "Gate 07B quedó configurado como un proceso estratégico independiente que transforma una Idea aceptada en una Project Proposal gobernada. Conserva trazabilidad completa con Gate 07A, no crea Project Workspace y termina exclusivamente en READY_FOR_STRATEGIC_GATE_DECISION.",
    PALE_GREEN,
)

heading(doc, "1. Objetivo y alcance")
doc.add_paragraph(
    "La entrega implementa el registro Project Proposal, su ciclo de vida, revisión, evaluación y contrato de preparación para Gate 07C. La solución reutiliza autenticación, tenant isolation, Workspaces empresariales, configuración versionada, auditoría, control de concurrencia y navegación existentes."
)
bullets(
    doc,
    [
        "Entrada válida: Idea en estado ACCEPTED con evaluación aceptada y readiness READY_FOR_PROJECT_PROPOSAL.",
        "Salida válida: Proposal en READY_FOR_STRATEGIC_GATE con readiness READY_FOR_STRATEGIC_GATE_DECISION.",
        "Relación: una Idea puede originar varias Proposals históricas, con máximo una activa por defecto.",
        "Límite: no se implementaron decisiones approve/reject/defer de Gate 07C ni creación de Project Workspace.",
    ],
)

heading(doc, "2. Integración arquitectónica")
add_table(
    doc,
    ["Capa", "Implementación", "Reutilización"],
    [
        ["Dominio", "project_proposal", "Patrones de Idea Lifecycle y configuración administrada"],
        ["API", "FastAPI /api/v1/project-proposals", "AuthContext, permisos, ETag, idempotencia"],
        ["Persistencia", "SQLAlchemy + Alembic 0041", "Tenant UUID, JSONB, timestamps, SecurityEvent"],
        ["USER UI", "Workspace Project Proposal", "Navegación de macroproceso/módulo/submódulo"],
        ["ADMIN UI", "Project Proposal Configuration", "Clone, edit, publish y preview"],
        ["Operación", "Docker Compose", "API, PostgreSQL, Redis y Vite existentes"],
    ],
    [1700, 3300, 4360],
)

heading(doc, "3. Invariantes del dominio")
bullets(
    doc,
    [
        "Project Proposal es un registro estratégico; no hereda ni materializa EnterpriseWorkspace.",
        "source_idea_id y accepted_evaluation_id preservan el origen y la decisión aceptada.",
        "owning_workspace_id sólo admite Enterprise, Business Unit o Portfolio; target_portfolio_id es independiente.",
        "Los objetivos estratégicos se copian como snapshot y no dependen de cambios posteriores del catálogo.",
        "Mapping, revisión y evaluación se resuelven contra revisiones publicadas y quedan congelados en la Proposal.",
        "Una evaluación completada es inmutable; una reevaluación crea una nueva versión.",
    ],
)

heading(doc, "4. Modelo de datos y migración")
doc.add_paragraph(
    "La revisión 20260813_0041 crea únicamente las dos tablas de negocio requeridas y enlaza la Proposal con Idea, evaluación aceptada, tenant, Workspaces, usuarios y configuración. La migración fue verificada en PostgreSQL mediante upgrade, downgrade y re-upgrade."
)
add_table(
    doc,
    ["Tabla", "Responsabilidad", "Controles relevantes"],
    [
        ["project_proposals", "Identidad, origen, ownership, snapshots, estado, version y readiness", "FK, índices, tenant_id, lock_version, timestamps"],
        ["project_proposal_evaluations", "Versiones de matriz, ratings, scores, hallazgos y resultado", "FK, versión única por Proposal, snapshot inmutable"],
    ],
    [2600, 3900, 2860],
)

heading(doc, "5. Ciclo de vida implementado")
add_table(
    doc,
    ["Secuencia", "Estado", "Control"],
    [
        ["1", "DRAFT", "Edición controlada y source snapshot"],
        ["2", "SUBMITTED", "Campos requeridos completos"],
        ["3", "UNDER_REVIEW", "Checklist publicado y asignación de reviewer"],
        ["4", "UNDER_EVALUATION", "Matriz vigente y evaluador autorizado"],
        ["5", "EVALUATED", "Evaluación inmutable completada"],
        ["6", "READY_FOR_STRATEGIC_GATE", "Readiness sin blockers para Gate 07C"],
    ],
    [1200, 3100, 5060],
)
doc.add_paragraph(
    "Los flujos de excepción incluyen RETURNED con comentario obligatorio, resubmission sin pérdida de historial y CANCELLED como cierre controlado. Las acciones disponibles se calculan en backend y se exponen a la UI."
)

heading(doc, "6. Numeración, idempotencia y concurrencia")
bullets(
    doc,
    [
        "AdminNumberSequence emite PROP-00001, PROP-00002 y siguientes dentro del tenant.",
        "Preview calcula el siguiente número pero no lo reserva ni incrementa la secuencia.",
        "Create bloquea la Idea fuente y la secuencia para evitar duplicados en concurrencia.",
        "Idempotency-Key permite reintentar la creación sin producir una Proposal adicional.",
        "If-Match/ETag valida lock_version; una versión obsoleta retorna HTTP 412.",
    ],
)

heading(doc, "7. Mapeo Idea → Project Proposal")
doc.add_paragraph(
    "La creación resuelve la configuración publicada, registra configuration_id/revision/hash/source y persiste valores copiados. Los datos originados en Gate 07A se muestran como sólo lectura; los campos propios de Proposal pueden enriquecerse antes del submit."
)
add_table(
    doc,
    ["Origen Idea", "Destino Proposal", "Tratamiento"],
    [
        ["name / description", "proposal_name / business_need", "Snapshot de origen"],
        ["expected_benefit", "justification / expected_benefits", "Snapshot trazable"],
        ["estimated_value", "rom_cost", "Valor ROM inicial"],
        ["objectives", "strategic_objectives_snapshot", "Copia versionada"],
        ["owning_workspace", "owning_workspace_id/path", "Herencia y validación"],
        ["accepted evaluation", "accepted_evaluation_id", "Referencia inmutable"],
    ],
    [2600, 3300, 3460],
)

heading(doc, "8. Revisión y checklist")
bullets(
    doc,
    [
        "Checklist configurable y versionado con snapshot dentro del registro.",
        "Asignación de reviewer validada contra usuarios activos del tenant.",
        "Hallazgos y comentarios quedan en el historial auditado.",
        "El reviewer puede continuar a evaluación o devolver para ajuste, sin alterar la Idea fuente.",
    ],
)

heading(doc, "9. Evaluación de Proposal")
doc.add_paragraph("La matriz publicada incorpora nueve criterios con pesos gobernados:")
bullets(
    doc,
    [
        "Strategic alignment, business value y benefits clarity.",
        "Scope maturity, delivery feasibility y schedule realism.",
        "Cost confidence, risk exposure y organizational readiness.",
        "Ratings, score ponderado, ranking, observaciones y recommendation quedan congelados por versión.",
    ],
)

heading(doc, "10. Contrato de preparación Gate 07C")
add_table(
    doc,
    ["Elemento", "Contenido"],
    [
        ["Estado previo", "EVALUATED"],
        ["Resultado", "READY_FOR_STRATEGIC_GATE_DECISION"],
        ["Bloqueadores", "Campos, review, evaluación, ownership, portfolio y configuración"],
        ["Warnings", "Condiciones no bloqueantes para la decisión posterior"],
        ["Integridad", "Hash del payload y referencias source/config/evaluation"],
        ["Exclusión", "No aprueba, rechaza, difiere ni crea Workspace"],
    ],
    [2600, 6760],
)

heading(doc, "11. Superficie API")
add_table(
    doc,
    ["Grupo", "Operaciones"],
    [
        ["Registro", "options, preview, create, list, detail, update"],
        ["Lifecycle", "submit, start-review, complete-review, return, resubmit, cancel"],
        ["Evaluation", "start, complete, versions"],
        ["Gate readiness", "readiness, mark-ready"],
        ["Idea integration", "GET /ideas/{id}/project-proposals"],
        ["Admin", "list configs, clone draft, edit, publish, preview"],
    ],
    [2600, 6760],
)

heading(doc, "12. Seguridad y segregación")
doc.add_paragraph("Se añadieron once permisos de mínima granularidad y se distribuyeron entre cinco perfiles funcionales.")
add_table(
    doc,
    ["Perfil", "Alcance principal"],
    [
        ["proposal_requestor", "Crear y mantener borradores autorizados"],
        ["proposal_reviewer", "Consultar, revisar, devolver y enrutar"],
        ["proposal_evaluator", "Ejecutar evaluación versionada"],
        ["proposal_gate_preparer", "Verificar y marcar readiness para Gate 07C"],
        ["organization_admin", "Administrar configuración y operación transversal"],
    ],
    [3000, 6360],
)
bullets(
    doc,
    [
        "Cada consulta y mutación filtra por tenant_id y alcance de Workspace.",
        "La autorización usa el mismo permission/context resolver del núcleo empresarial.",
        "SecurityEvent registra creación, transición, actualización, evaluación y publicación de configuración.",
    ],
)

heading(doc, "13. USER MODE")
bullets(
    doc,
    [
        "Nuevo submódulo Project Proposal dentro de Enterprise Strategy Manager → Idea & Demand Manager.",
        "Colas: My Project Proposals, All Authorized, Draft, To Review, Assigned to Me, Under Evaluation, Ready for Strategic Gate, Returned y Cancelled.",
        "Registro con métricas, búsqueda y detalle gobernado por allowed_actions.",
        "Paneles de origen Idea, contenido Proposal, review, evaluation, historial y readiness.",
        "New Proposal sólo acepta Ideas elegibles y ofrece Preview antes de persistir.",
    ],
)

heading(doc, "14. ADMIN MODE")
bullets(
    doc,
    [
        "Nuevo submódulo Project Proposal Configuration dentro de Enterprise Strategy Manager.",
        "Dos revisiones iniciales publicadas: lifecycle/mapping y evaluation matrix.",
        "Secciones para numbering, mapping, required fields, checklist, ownership, target portfolio, matrix, readiness, inheritance, permissions y preview.",
        "Clone to draft conserva published inmutable; edit valida JSON; publish activa una nueva revisión gobernada.",
    ],
)

heading(doc, "15. Integración con Idea Lifecycle")
doc.add_paragraph(
    "El detalle de una Idea ACCEPTED incorpora la lista de Proposals relacionadas y la acción Preview/Create cuando el backend lo permite. La respuesta de readiness de Gate 07A declara can_create_project_proposal y conserva source_idea_id y accepted_evaluation_id. No se modifican evaluaciones históricas de Idea."
)

heading(doc, "16. Navegación por Workspace")
doc.add_paragraph(
    "El navegador operacional incorpora Project Proposals después de Ideas para Workspaces Enterprise, Business Unit y Portfolio. El componente ProjectProposalWorkspace se reutiliza en la ruta contextual y en el acceso global; no se duplicó lógica de negocio ni una segunda pantalla transaccional."
)

heading(doc, "17. Evidencia de pruebas")
add_table(
    doc,
    ["Verificación", "Resultado", "Evidencia"],
    [
        ["Ruff alcance CI", "PASS", "backend/app + backend/tests"],
        ["Backend focal", "PASS", "12 pruebas Gate 07A/07B/workspace context"],
        ["PostgreSQL Gate 07B", "PASS", "6 pruebas Gate 07A + 07B"],
        ["Alembic", "PASS", "upgrade 0040→0041, downgrade, re-upgrade"],
        ["TypeScript", "PASS", "tsc --noEmit"],
        ["ESLint focal", "PASS", "0 warnings"],
        ["Vitest", "PASS", "4 pruebas / 2 archivos"],
        ["Vite build", "PASS", "2.362 módulos transformados"],
        ["AppFlow", "PASS", "27 casos + 1 timeout ambiental aprobado en ejecución aislada"],
        ["Browser localhost", "PASS", "USER + ADMIN; consola sin errores"],
    ],
    [3300, 1500, 4560],
)

heading(doc, "18. Validación operativa en localhost")
add_table(
    doc,
    ["Servicio", "Dirección", "Estado"],
    [
        ["Frontend", "http://127.0.0.1:5173/app", "Up"],
        ["API", "http://127.0.0.1:8000", "Healthy"],
        ["Readiness", "/api/v1/health/ready", "api/database/redis: ok"],
        ["PostgreSQL", "Docker internal", "Migration 0041 aplicada"],
        ["Redis", "Docker internal", "Ready"],
    ],
    [2700, 4300, 2360],
)

heading(doc, "19. Archivos principales entregados")
bullets(
    doc,
    [
        "backend/app/modules/project_proposal/{models,schemas,service,router}.py",
        "backend/alembic/versions/20260813_0041_gate07b_project_proposal.py",
        "backend/tests/test_project_proposal_gate07b.py y docker-compose.gate07b.yml",
        "frontend/src/features/project-proposal/{types,api,ProjectProposalWorkspace,ProjectProposalAdminView,projectProposal.css}.tsx/ts/css",
        "Integraciones en App.tsx, applicationNavigation.ts, IdeaLifecycleWorkspace.tsx y WorkspaceOperationalPage.tsx.",
    ],
)

heading(doc, "20. Riesgos y deuda controlada")
bullets(
    doc,
    [
        "Gate 07C todavía no existe; el contrato sólo entrega readiness y debe ser consumido sin reinterpretar la decisión.",
        "La matriz y reglas iniciales son defaults gobernados; negocio debe ratificar pesos y checklist antes de producción.",
        "La ejecución AppFlow completa presentó un timeout bajo carga Docker paralela; el mismo caso pasó en ejecución aislada ampliada.",
        "El almacenamiento de anexos reutiliza referencias; la gestión binaria empresarial permanece fuera de este gate.",
    ],
)

heading(doc, "21. Matriz de aceptación")
criteria = [
    "Proposal separada de Idea y Workspace",
    "Sólo Ideas ACCEPTED + READY",
    "Trazabilidad source/evaluation",
    "Máximo una activa por defecto",
    "Numeración PROP",
    "Preview no consume número",
    "Idempotencia",
    "ETag / 412",
    "Owning Workspace válido",
    "Target Portfolio separado",
    "Snapshot de objetivos",
    "Mapping config/revision/hash",
    "Dos tablas de negocio",
    "Migración reversible",
    "Ciclo de seis estados",
    "Return/resubmit/cancel",
    "Checklist versionado",
    "Evaluación de nueve criterios",
    "Evaluación inmutable",
    "Reevaluación versionada",
    "Readiness Gate 07C",
    "Sin decisión Gate 07C",
    "Sin creación de Workspace",
    "11 permisos",
    "5 perfiles",
    "SecurityEvent",
    "USER queues",
    "ADMIN configuration",
    "Integración Idea",
    "Navegación contextual",
    "Pruebas PostgreSQL",
    "Frontend build/tests",
    "Localhost healthy",
]
add_table(doc, ["Criterio", "Resultado"], [[item, "PASS"] for item in criteria], [7850, 1510])

heading(doc, "22. Estado final y siguiente gate")
callout(doc, "Estado final exclusivo", "READY_FOR_STRATEGIC_GATE_DECISION", PALE_GREEN)
doc.add_paragraph(
    "Gate 07B termina con una Project Proposal evaluada y preparada para decisión estratégica. La siguiente implementación permitida es Gate 07C, que deberá consumir el payload de readiness, preservar sus hashes y referencias, y registrar la decisión sin convertir retroactivamente Idea o Proposal en Workspace. Si existen blockers, el resultado permitido es GATE07B_REWORK_REQUIRED."
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
