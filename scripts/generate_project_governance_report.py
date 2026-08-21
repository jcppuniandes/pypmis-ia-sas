from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUT = (
    REPO
    / "outputs"
    / "technical_reports"
    / "Informe_Tecnico_Project_Governance_Multi_Source_Creation_20260820.docx"
)
LOGO = REPO / "frontend" / "public" / "pypmis-construction-ai-logo.png"

# Design preset: standard_business_brief. Header: memo_masthead.
NAVY = "0D2A3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "0D8F8B"
INK = "16222B"
MUTED = "5B6870"
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
PALE_TEAL = "EAF7F6"


def set_run(run, *, size=10.0, color=INK, bold=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margin(cell, top=70, start=105, bottom=70, end=105):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_geometry(table, widths, indent=105):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", indent)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header._tr.get_or_add_trPr().append(header_flag)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, NAVY)
        cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=8.6, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT)
            cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=8.7)
    table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def callout(doc, title, body, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_flag)
    cell = table.cell(0, 0)
    shade(cell, fill)
    cell_margin(cell, 120, 145, 120, 145)
    set_run(cell.paragraphs[0].add_run(title.upper()), size=8.7, color=NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run(run, size=9.2)
    table_geometry(table, [9360], indent=145)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1, page_break=False):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    return paragraph


def bullets(doc, values):
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2.5)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("P&Pmis Construction AI  |  "), size=8.0, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, end])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.72)
section.bottom_margin = Inches(0.82)
section.left_margin = section.right_margin = Inches(0.82)
section.header_distance = section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(9.8)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(4)
normal.paragraph_format.line_spacing = 1.04
for name, size, color, before, after in (
    ("Title", 21, NAVY, 0, 3),
    ("Subtitle", 12, MUTED, 0, 12),
    ("Heading 1", 13.4, BLUE, 10, 4),
    ("Heading 2", 11.2, DARK_BLUE, 7, 3),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(9.7)
    styles[name].paragraph_format.left_indent = Inches(0.45)
    styles[name].paragraph_format.first_line_indent = Inches(-0.22)

doc.core_properties.title = "Informe técnico — Project Governance Model & Multi-Source Project Creation"
doc.core_properties.subject = "Implementación de creación de Project multi-fuente en P&Pmis Construction AI"
doc.core_properties.author = "P&Pmis Construction AI / Codex"
doc.core_properties.keywords = "Project Governance, Gate 05B, Gate 05C, Gate 07D, PostgreSQL, P&Pmis"
language = OxmlElement("w:lang")
language.set(qn("w:val"), "es-CO")
styles["Normal"]._element.get_or_add_rPr().append(language)

header = section.header.paragraphs[0]
set_run(header.add_run("P&Pmis Construction AI  |  Project Governance Release"), size=8.0, color=MUTED, bold=True)
add_page_number(section.footer.paragraphs[0])

masthead = doc.add_table(rows=1, cols=2)
masthead.alignment = WD_TABLE_ALIGNMENT.LEFT
masthead_header = OxmlElement("w:tblHeader")
masthead_header.set(qn("w:val"), "true")
masthead.rows[0]._tr.get_or_add_trPr().append(masthead_header)
logo_cell, title_cell = masthead.rows[0].cells
if LOGO.exists():
    picture = logo_cell.paragraphs[0].add_run().add_picture(str(LOGO), width=Inches(0.78))
    picture._inline.docPr.set("descr", "Logotipo P&Pmis Construction AI")
title = title_cell.paragraphs[0]
set_run(title.add_run("INFORME TÉCNICO"), size=8.8, color=TEAL, bold=True)
subtitle = title_cell.add_paragraph()
set_run(subtitle.add_run("Project Governance Model &\nMulti-Source Project Creation"), size=19, color=NAVY, bold=True)
table_geometry(masthead, [1350, 8010], indent=120)

add_table(
    doc,
    ["Campo", "Detalle"],
    [
        ["Aplicación", "P&Pmis Construction AI"],
        ["Ámbito", "USER MODE + ADMIN MODE · Enterprise / Project Workspace"],
        ["Baseline reutilizado", "Gate 05B, Gate 05C y Gate 07D"],
        ["Migración", "20260820_0044, sucesora aditiva de 20260820_0043"],
        ["Fecha de corte", "20 de agosto de 2026"],
        ["Estado final", "READY_FOR_MULTI_SOURCE_PROJECT_CREATION"],
    ],
    [2200, 7160],
)

callout(
    doc,
    "Resultado ejecutivo",
    "Se implementó una sola capa transversal de gobernanza para crear Projects desde una Strategic Gate Decision, un Contract Award o una Direct Authorization. Se conserva una única identidad canónica PROJECT, un solo ProjectCreationRequest, el ciclo Four-eyes existente y la materialización idempotente. El localhost quedó operativo en la revisión 0044 y el Project Workspace ID 14 permaneció inalterado.",
    PALE_GREEN,
)

heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph(
    "La entrega amplía el proceso Gate 05B sin crear un segundo Project Creator ni una entidad paralela. Cada solicitud registra modelo de gobierno, fuente normalizada, snapshot/hash y la revisión exacta de política utilizada. Gate 05C continúa siendo el motor común de inicialización/activación y Gate 07D conserva la ruta Capital Owner."
)

heading(doc, "2. ADR: Multi-Source Project Creation")
doc.add_paragraph(
    "ADR 44 acepta una arquitectura de adaptadores de fuente sobre ProjectCreationService. EnterpriseWorkspace con workspace_type_code=project sigue siendo la identidad Project; ProjectCreationRequest es el proceso; AdminConfiguration es la autoridad de política. La decisión está documentada en docs/44-adr-project-governance-multi-source-creation.md."
)

heading(doc, "3. Definición del Project Governance Model")
add_table(
    doc,
    ["Modelo", "Fuente", "Propósito"],
    [
        ["CAPITAL_OWNER", "STRATEGIC_GATE_DECISION", "Inversión promovida por estrategia/Portfolio"],
        ["CONTRACTOR_DELIVERY", "CONTRACT_AWARD", "Proyecto de entrega originado en contrato adjudicado"],
        ["DIRECT_INTERNAL", "DIRECT_AUTHORIZATION", "Proyecto interno autorizado directamente"],
    ],
    [2500, 3050, 3810],
)

heading(doc, "4. Project Type vs Governance Model")
doc.add_paragraph(
    "Project Type permanece como clasificación de negocio opcional; Governance Model define origen, guardrails y readiness. No existe inferencia automática entre ambas dimensiones. Las pruebas crean DIRECT_INTERNAL con project_type nulo y mantienen el modelo explícito."
)

heading(doc, "5. Source Context model")
bullets(
    doc,
    [
        "source_context_type + source_context_id identifica fuentes internas tipadas.",
        "source_external_key identifica Contract Award o Direct Authorization externas.",
        "idempotency_key protege reintentos y source_hash verifica la huella normalizada.",
        "source_snapshot_json preserva evidencia sin introducir motores Contract/Procurement.",
    ],
)

heading(doc, "6. REUSE BEFORE CREATE")
add_table(
    doc,
    ["Capacidad", "Reutilización", "No duplicado"],
    [
        ["Identidad", "enterprise_workspaces / PROJECT", "No segunda tabla Project"],
        ["Proceso", "ProjectCreationRequest Gate 05B", "No segundo creator"],
        ["Lifecycle", "submit/review/approve/materialize", "No state machine paralela"],
        ["Inicialización", "Gate 05C", "No inicializador por modelo"],
        ["Capital Owner", "Gate 07D", "No reemplazo del strategic entry"],
        ["Política/auditoría", "AdminConfiguration + SecurityEvent", "No framework alterno"],
    ],
    [2100, 3650, 3610],
)

heading(doc, "7. Extensión de ProjectCreationRequest")
doc.add_paragraph(
    "Se añadieron governance_model, source_context_id, source_external_key, idempotency_key, source_snapshot_json, source_hash y creation_policy_id/revision/hash. Los campos Gate 07D permanecen. Submission/approval hashes incorporan lineage para congelar la fuente aprobada."
)

heading(doc, "8. Diseño Source Adapter")
doc.add_paragraph(
    "StrategicGateSourceAdapter, ContractAwardSourceAdapter y DirectAuthorizationSourceAdapter normalizan cada fuente a NormalizedProjectCreationSource. El adaptador valida cardinalidad y campos mínimos; después delega al mismo ProjectCreationService."
)

heading(doc, "9. Política CAPITAL_OWNER")
doc.add_paragraph(
    "Requiere Strategic Gate Decision, Portfolio, objetivo estratégico, template, PM y readiness Gate 07D. Mantiene Portfolio/FEL como trabajo posterior y nunca interpreta APPROVE como FID."
)

heading(doc, "10. Política CONTRACTOR_DELIVERY")
doc.add_paragraph(
    "Requiere client, contract_number y contractual_scope. Por defecto no exige Idea, Proposal, Strategic Gate, Portfolio, FEL, PDRI ni FID. La activación depende de movilización autorizada o Notice to Proceed cuando la política lo configure."
)

heading(doc, "11. Política DIRECT_INTERNAL")
doc.add_paragraph(
    "Requiere authorization_reference, sponsor e idempotency key. No exige Contract Award ni Strategic Gate. La autorización aprobada es condición de activación, no de creación del Project PENDING."
)

heading(doc, "12. Estrategia Contract Award source")
doc.add_paragraph(
    "Se optó por snapshot controlado, porque el módulo Contracts existente es operacional y presupone un Project ya creado. Esta decisión evita dependencia circular y no implementa Contract Management profundo."
)

heading(doc, "13. Estrategia Direct Authorization source")
doc.add_paragraph(
    "La referencia y patrocinador se registran en un snapshot tipado. source_external_key e idempotency_key comparten la referencia por defecto, sin impedir que un integrador suministre claves distintas."
)

heading(doc, "14. Source snapshot")
doc.add_paragraph(
    "El snapshot se serializa canónicamente y su SHA-256 incluye governance_model, source_context_type, identificadores y contenido. Se replica al metadata _project en materialización, junto con la política exacta y su hash."
)

heading(doc, "15. Resolución de Creation Policy", page_break=True)
doc.add_paragraph(
    "ProjectGovernancePolicyService resuelve primero overrides publicados en la cadena de workspaces y después el default tenant. La solicitud guarda configuration_id, revision y content_hash; validaciones históricas usan esa revisión y no reinterpretan una solicitud aprobada con una política futura."
)

heading(doc, "16. Resolución de Initialization Policy")
doc.add_paragraph(
    "Gate 05C consulta el governance metadata materializado. CONTRACTOR_DELIVERY y DIRECT_INTERNAL no heredan el requisito global de objetivo estratégico; validan fuente, PM, template y campos/readiness propios. Legacy conserva su comportamiento anterior."
)

heading(doc, "17. Resolución de Activation Policy")
doc.add_paragraph(
    "La activación sigue siendo una transición común. CAPITAL_OWNER requiere execution_authorized; CONTRACTOR_DELIVERY, mobilization_authorized o NTP; DIRECT_INTERNAL, authorization_approved. Ningún Project se autoactiva."
)

heading(doc, "18. Semántica PENDING por modelo")
add_table(
    doc,
    ["Modelo", "pending_reason", "planning_stage"],
    [
        ["CAPITAL_OWNER", "PORTFOLIO_AND_PROJECT_DEFINITION_REQUIRED", "PORTFOLIO_AND_FEL_PLANNING"],
        ["CONTRACTOR_DELIVERY", "INITIALIZATION_AND_MOBILIZATION_REQUIRED", "CONTRACT_MOBILIZATION"],
        ["DIRECT_INTERNAL", "INITIALIZATION_REQUIRED", "DIRECT_AUTHORIZATION"],
    ],
    [2450, 3900, 3010],
)

heading(doc, "19. Flujo UI")
doc.add_paragraph(
    "Create Project muestra primero Governance Model y luego un único formulario gobernado. Direct Internal presenta autorización/patrocinador; Contractor Delivery presenta cliente/contrato/alcance; Capital Owner redirige conceptualmente a Strategic Project Planning Entry y bloquea el submit genérico."
)

heading(doc, "20. Configuración ADMIN")
doc.add_paragraph(
    "Creation Policies incluye tres tabs de política, revisión/hash, fuente de herencia y guardrails editables. Guardar publica una nueva revisión tenant/workspace mediante el API; la pantalla no crea ni materializa Projects."
)

heading(doc, "21. Herencia de configuración")
doc.add_paragraph(
    "Los códigos project-governance-{model}-{tenant|workspace-id} permiten default tenant y override por workspace. Preview informa resolution_chain, source_workspace_id, required/optional fields, approvals, warnings y blockers sin persistir."
)

heading(doc, "22. APIs", page_break=True)
add_table(
    doc,
    ["Ámbito", "Endpoints"],
    [
        ["Opciones", "GET /project-creation/options"],
        ["Contrato", "POST /project-creation/from-contract/preview y /from-contract"],
        ["Directo", "POST /project-creation/direct/preview y /direct"],
        ["Lifecycle", "Rutas existentes /project-creation-requests/*"],
        ["ADMIN", "PUT /project-governance-models/{model}; POST /preview"],
        ["Overview", "GET /project-workspaces/{id}/overview"],
    ],
    [2200, 7160],
)

heading(doc, "23. Permisos y roles")
bullets(
    doc,
    [
        "project_governance_model.read/configure/publish",
        "project_creation.contract_source.create y project_creation.direct.create",
        "project_requestor: contrato/directo; portfolio_intake_planner: Capital Owner; organization_admin: todas.",
        "Reviewer, approver y materializer conservan permisos y segregación Gate 05B.",
    ],
)

heading(doc, "24. ETag e idempotencia")
doc.add_paragraph(
    "If-Match protege edición y transiciones. Índices parciales tenant/source impiden más de una solicitud activa por fuente; reintentos activos retornan la solicitud existente. Requests cancelled/rejected permiten una nueva ruta contract/direct sin colisión. Project Number se reserva únicamente al materializar."
)

heading(doc, "25. Migración")
doc.add_paragraph(
    "Alembic 20260820_0044 es aditiva, indexada y reversible a 0043. Agrega FK de policy, source_snapshot NOT NULL tras backfill y tres unique indexes parciales. No reescribe 0043. El ciclo PostgreSQL probado fue 0042→0043→0044→0043→0044→0042→0044."
)

heading(doc, "26. Compatibilidad hacia atrás")
doc.add_paragraph(
    "Requests Gate 07D con evidencia se backfillan CAPITAL_OWNER; requests Gate 05B ordinarias quedan governance_model nullable y se muestran Legacy / Not Classified. No se infiere una clasificación sin evidencia."
)

heading(doc, "27. SecurityEvent y auditoría")
doc.add_paragraph(
    "Se emiten source_validated y governance_model_selected, además de request_created/submitted/review_started/approved/workspace_created. Metadata incluye tenant, actor, request number, governance model, source context y hashes."
)

heading(doc, "28. PostgreSQL E2E", page_break=True)
add_table(
    doc,
    ["Evidencia", "Resultado"],
    [
        ["Schema 0042", "Gate 07D ausente — PASS"],
        ["Schema 0043", "Gate 07D presente; multi-source ausente — PASS"],
        ["Schema 0044", "9 columnas, FK, índices parciales y NOT NULL — PASS"],
        ["Primera corrida", "69 pruebas PASS"],
        ["Rebuild final", "52 pruebas PASS"],
        ["Reversibilidad", "0044→0043→0044 y 0044→0042→0044 — PASS"],
        ["Backend Docker completo", "477 PASS, 13 skips, 0 fallos"],
    ],
    [3400, 5960],
)

heading(doc, "29. Pruebas frontend")
doc.add_paragraph(
    "TypeScript compila sin errores; Vite generó el bundle de producción. La suite completa cerró 35 archivos y 194 pruebas PASS; las 33 pruebas focalizadas Project/Enterprise también pasaron. Se validaron selector de gobernanza, campos directos, source preview, overview y políticas ADMIN, además de las regresiones BIM/Enterprise existentes."
)

heading(doc, "30. Browser E2E")
doc.add_paragraph(
    "Chromium recorrió login ADMIN, Enterprise Structure Configuration, tres tabs de policy, cambio a USER, Enterprise Explorer, Create Project y los formularios dinámicos Direct/Contract/Capital. Resultado final: 1 escenario PASS en 19.0 s contra el localhost actualizado."
)

heading(doc, "31. Regresión Gate 07D")
doc.add_paragraph(
    "La ruta StrategicGateDecision APPROVE → CAPITAL_OWNER → Project PENDING → target Portfolio Membership → READY_FOR_PORTFOLIO_PLANNING permanece en el servicio y fue incluida en las suites SQLite/PostgreSQL. No se agregaron FID, scoring ni Portfolio Evaluation."
)

heading(doc, "32. Baseline antes/después")
add_table(
    doc,
    ["Control", "Antes", "Después", "Resultado"],
    [
        ["Alembic persistente", "20260813_0042", "20260820_0044", "PASS"],
        ["Project Workspace ID 14", "MD5 eea3648c6908a267d1e0651828ff4459", "Mismo MD5", "Sin mutación"],
        ["API ready", "Operativa", "HTTP 200 API/DB/Redis", "PASS"],
        ["Frontend", "localhost:5173", "HTTP 200 /app", "PASS"],
        ["Template", "DRAFT validada", "PYP-PRJ-CONSTRUCTION PUBLISHED", "Flujo habilitado"],
    ],
    [2350, 2900, 2830, 1280],
)

heading(doc, "33. Riesgos y deuda")
bullets(
    doc,
    [
        "El adapter Contract Award usa snapshot hasta que exista una autoridad contractual pre-Project; integrar después sin cambiar la identidad canónica.",
        "Las políticas workspace-level están soportadas por backend; la UI actual edita el default tenant y puede ampliarse con selector de scope.",
        "Mantener los índices parciales y la congelación de source hash en futuras migraciones.",
        "El bundle BIM es grande pero se mantiene lazy-loaded; no es deuda introducida por esta entrega.",
    ],
)

heading(doc, "34. Informe técnico actualizado")
doc.add_paragraph(
    "Este documento consolida arquitectura, datos, seguridad, UI, migración y evidencias verificadas de la implementación. El ADR 44 y los tests quedan versionados con el código."
)

heading(doc, "35. Estado final")
callout(
    doc,
    "READY_FOR_MULTI_SOURCE_PROJECT_CREATION",
    "Todos los criterios de aceptación de esta capa fundacional están cubiertos: identidad/proceso únicos, tres modelos, Project Type independiente, PENDING común, readiness por política, RBAC/Four-eyes/ETag/idempotencia/auditoría, PostgreSQL, frontend, navegador, regresión Gate 07D y baseline protegido.",
    PALE_GREEN,
)

heading(doc, "STOP boundary")
callout(
    doc,
    "Alcance detenido",
    "No se inició Gate 07E, Portfolio Evaluation, priorización, FEL/PDRI/FID, Contract Management profundo, Procurement ni módulos operativos Project. Se requiere un prompt explícito para continuar cualquiera de esas áreas.",
    PALE_AMBER,
)

heading(doc, "Anexo A. Archivos principales")
add_table(
    doc,
    ["Área", "Archivo"],
    [
        ["Backend", "backend/app/modules/project_creation/{governance,service,router,schemas,models}.py"],
        ["ADMIN / readiness", "enterprise_structure/project_configuration.py; project_workspace_initialization/service.py; workspace_context/navigator.py"],
        ["Migración / ADR", "backend/alembic/versions/20260820_0044_project_governance_multi_source.py; docs/44-adr-project-governance-multi-source-creation.md"],
        ["Frontend", "ProjectCreationWorkspace.tsx; ProjectWorkspaceConfigurationPanel.tsx"],
        ["E2E", "backend/tests/test_project_creation_process.py; frontend/e2e/project-governance-multi-source.spec.ts"],
    ],
    [2200, 7160],
)

heading(doc, "Anexo B. Matriz de aceptación")
criteria = [
    "Una identidad PROJECT y un ProjectCreationRequest compartido",
    "CAPITAL_OWNER, CONTRACTOR_DELIVERY y DIRECT_INTERNAL",
    "Project Type independiente",
    "Gate 07D sin regresión",
    "Contrato sin Idea/Proposal/Gate/Portfolio/FEL/PDRI/FID por defecto",
    "PENDING común, inicialización/activación comunes y readiness por modelo",
    "Sin creator o entidad paralelos",
    "RBAC, Four-eyes, ETag, idempotencia, tenant scope y audit",
    "PostgreSQL, frontend, navegador, regresión y baseline protegido",
]
add_table(doc, ["Criterio", "Estado"], [[item, "PASS"] for item in criteria], [7900, 1460])

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
