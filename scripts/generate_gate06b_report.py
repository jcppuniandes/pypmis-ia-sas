from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUT = REPO / "outputs" / "technical_reports" / "Informe_Tecnico_Gate_06B_Physical_Workspace_Creation_20260813.docx"

NAVY = "0D2A3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
TEAL = "149A94"
INK = "16222B"
MUTED = "5B6870"
WHITE = "FFFFFF"
PALE_BLUE = "F2F4F7"
PALE_TEAL = "EAF7F6"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
LINE = "CAD5DC"


def set_run(run, *, size: float = 11, color: str = INK, bold: bool | None = None, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_table_geometry(table, widths: list[int], indent: int = 120) -> None:
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int]) -> None:
    item = doc.add_table(rows=1, cols=len(headers))
    item.style = "Table Grid"
    item.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, header in enumerate(headers):
        cell = item.rows[0].cells[index]
        cell.text = header
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, color=WHITE, bold=True)
    set_repeat_header(item.rows[0])
    for row_index, values in enumerate(rows):
        cells = item.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            shade(cells[index], WHITE if row_index % 2 == 0 else PALE_BLUE)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cells[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=9.2)
    set_table_geometry(item, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(2)


def callout(doc: Document, title: str, body: str, fill: str = PALE_TEAL) -> None:
    item = doc.add_table(rows=1, cols=1)
    item.style = "Table Grid"
    cell = item.cell(0, 0)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    title_run = cell.paragraphs[0].add_run(title.upper())
    set_run(title_run, size=9, color=NAVY, bold=True)
    p = cell.add_paragraph(body)
    p.paragraph_format.space_after = Pt(0)
    for run in p.runs:
        set_run(run, size=9.5)
    set_repeat_header(item.rows[0])
    set_table_geometry(item, [9360])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    p.paragraph_format.keep_with_next = True


def bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        p = doc.add_paragraph(value, style="List Bullet")
        p.paragraph_format.space_after = Pt(4)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Title", 23, NAVY, 0, 4),
    ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(11)
    styles[name].paragraph_format.space_after = Pt(8)
    styles[name].paragraph_format.line_spacing = 1.167
    styles[name].paragraph_format.left_indent = Inches(0.5)
    styles[name].paragraph_format.first_line_indent = Inches(-0.25)

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.LEFT
header.paragraph_format.space_after = Pt(0)
set_run(header.add_run("P&Pmis Construction AI  |  Gate 06B Technical Report"), size=8.5, color=MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer.add_run("Controlled technical baseline  |  13 Aug 2026  |  Page "), size=8, color=MUTED)
page_run = footer.add_run()
set_run(page_run, size=8, color=MUTED)
for field_type, value in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
    if field_type is not None:
        field = OxmlElement("w:fldChar")
        field.set(qn("w:fldCharType"), field_type)
        page_run._r.append(field)
    else:
        text = OxmlElement("w:instrText" if value == " PAGE " else "w:t")
        text.set(qn("xml:space"), "preserve")
        text.text = value
        page_run._r.append(text)

doc.core_properties.title = "Informe técnico Gate 06B - Physical Workspace Creation Process"
doc.core_properties.subject = "PROPERTY, FACILITY y WAREHOUSE - proceso único gobernado"
doc.core_properties.author = "P&Pmis Construction AI Engineering"
doc.core_properties.comments = "Gate 06B cerrado; Gate 06C no iniciado."

doc.add_paragraph("INFORME TÉCNICO", style="Title")
doc.add_paragraph("Gate 06B · Physical Workspace Creation Process", style="Subtitle")
for label, value in (
    ("Tenant", "P&P Ingeniería y Proyectos"),
    ("Fecha de cierre", "13 de agosto de 2026"),
    ("Alcance", "PROPERTY · FACILITY · WAREHOUSE"),
    ("Estado", "READY_FOR_PHYSICAL_WORKSPACE_INITIALIZATION"),
    ("Siguiente gate", "Gate 06C · Initialization & Activation (no iniciado)"),
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    set_run(p.add_run(f"{label}: "), bold=True)
    set_run(p.add_run(value))
callout(
    doc,
    "Resultado de cierre",
    "Se implementó un solo proceso parametrizado por workspace_type_code para PROPERTY, FACILITY y WAREHOUSE. "
    "La materialización transaccional produce un enterprise_workspace PENDING; no inicializa ni activa el Workspace.",
    PALE_GREEN,
)

heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph(
    "Gate 06B amplía la arquitectura abierta existente con un proceso gobernado, tenant-scoped y reutilizable para "
    "solicitar, revisar, aprobar y materializar Workspaces físicos operativos. La identidad canónica continúa en "
    "enterprise_workspaces; la nueva entidad conserva únicamente el ciclo de solicitud, snapshots, hashes, decisión, "
    "trazabilidad y vínculo hacia el Workspace final."
)
table(
    doc,
    ["Criterio", "Resultado"],
    [
        ["Motor único", "PASS · PROPERTY / FACILITY / WAREHOUSE"],
        ["Tipos excluidos", "PASS · REGION / DISTRICT / SITE / LINEAR_ASSET / ASSET / PROJECT"],
        ["Preview", "PASS · no persistente; no consume Business Number"],
        ["Gobierno", "PASS · RBAC, Four-Eyes, ETag/If-Match, snapshots y hashes"],
        ["Materialización", "PASS · una transacción; Workspace PENDING"],
        ["Concurrencia", "PASS · Business Numbers y Record Codes únicos"],
        ["Idempotencia", "PASS · segundo materialize = ALREADY_CREATED / 0 mutaciones"],
        ["PostgreSQL E2E", "PASS · 31 pruebas y migración reversible"],
        ["Frontend", "PASS · 182/182 tests; build, lint y Prettier"],
        ["Baseline real", "PASS · 14 Workspaces; 0 solicitudes físicas; secuencias sin consumo"],
    ],
    [2900, 6460],
)

heading(doc, "2. REUSE BEFORE CREATE")
doc.add_paragraph(
    "Se inspeccionaron y reutilizaron las capacidades probadas en Gates 05B, 05C y 06A. No se generalizó "
    "ProjectCreationRequest porque habría mezclado semánticas de Project con dominios físicos; en su lugar se creó "
    "una única entidad física y un único servicio parametrizado."
)
table(
    doc,
    ["Capacidad reutilizada", "Uso en Gate 06B"],
    [
        ["enterprise_workspaces", "Identidad canónica y jerarquía del Workspace final"],
        ["AdminConfiguration", "Workspace Type, Composition Rule, Template, Policy, Catalog y Numbering Rule"],
        ["AdminNumberSequence", "Request Number y Business Number con UPDATE ... RETURNING"],
        ["Record Code engine", "Cálculo jerárquico con bloqueo de parent para concurrencia"],
        ["SecurityAccessAssignment", "Asignación mínima del responsable en scope Workspace"],
        ["SecurityEvent", "Auditoría de cada transición y resultado materializado/fallido"],
        ["ETag / If-Match", "Control optimista de todas las mutaciones, incluida Materialization"],
        ["Enterprise Explorer", "CORE publicado + Project operativo + Workspaces físicos operativos"],
    ],
    [2800, 6560],
)

heading(doc, "3. Decisión de alcance estructural vs. operacional")
bullets(
    doc,
    [
        "REGION, DISTRICT y SITE siguen administrándose en ADMIN MODE mediante Workspace Revision Manager.",
        "PROPERTY, FACILITY y WAREHOUSE usan el nuevo Physical Workspace Creation Process en USER MODE.",
        "LINEAR_ASSET permanece reservado/inactivo y ASSET continúa fuera de Enterprise Workspace Structure.",
        "Gate 06B no crea inventario, assets, mantenimiento, spaces, utilities ni transacciones inmobiliarias.",
    ],
)

page_break(doc)
heading(doc, "4. Arquitectura del proceso único parametrizado")
table(
    doc,
    ["Capa", "Responsabilidad"],
    [
        ["UI React", "Selector de tipo, intake común, atributos dinámicos, solicitudes, review queue y overview"],
        ["FastAPI Router", "Contratos USER/REVIEW/MATERIALIZATION/OVERVIEW y autorización por permiso"],
        ["Domain Service", "Reglas, transiciones, snapshots, hashes, preview, materialización e idempotencia"],
        ["SQLAlchemy Model", "Una sola tabla physical_workspace_creation_requests"],
        ["Governance Config", "Backend source of truth para parents, templates, policy, atributos y clasificaciones"],
        ["PostgreSQL", "Locks, secuencias, unicidad, FKs y transacción atómica"],
    ],
    [2200, 7160],
)

heading(doc, "5. Modelo PhysicalWorkspaceCreationRequest")
doc.add_paragraph(
    "La tabla tenant-scoped conserva request_number, workspace_type_code, state, requestor, parent, template, nombre, "
    "descripción, responsable, previews, atributos, clasificaciones, snapshot y hash de submit, decisión de approval, "
    "hash de aprobación, vínculo materializado, números finales, revision_version y trazabilidad temporal/actor."
)
callout(
    doc,
    "Separación de identidades",
    "Request Number (PWR-00001) != Business Number (PYP-PROP-00001) != Record Code (01.05.01). "
    "Preview y Approval no consumen el Business Number; solo Materialization lo reserva.",
)

heading(doc, "6. Máquina de estados y controles")
table(
    doc,
    ["Estado origen", "Transición permitida", "Control principal"],
    [
        ["DRAFT", "SUBMITTED / CANCELLED", "Owner, If-Match y revalidación total"],
        ["SUBMITTED", "UNDER_REVIEW / RETURNED / CANCELLED", "Reviewer y razón para Return"],
        ["UNDER_REVIEW", "APPROVED / RETURNED / REJECTED", "Approver, Four-Eyes y hash"],
        ["RETURNED", "DRAFT / CANCELLED", "Owner edita y crea nueva versión"],
        ["APPROVED", "MATERIALIZING", "If-Match, approval_hash y locks"],
        ["MATERIALIZING", "CREATED / FAILED", "Commit único o rollback"],
        ["CREATED", "ALREADY_CREATED", "Idempotencia; 0 mutaciones"],
    ],
    [1700, 2800, 4860],
)

heading(doc, "7. Type, Location, Template y Policy Pickers")
bullets(
    doc,
    [
        "Workspace Type Picker expone exclusivamente Property, Facility y Warehouse; el backend rechaza cualquier otro código.",
        "Workspace Location Picker consume Composition Rules, Policy publicada, tenant, status active y árbol CORE + operacional.",
        "La UI no contiene una matriz de parents: consulta la elegibilidad al backend para cada acción contextual.",
        "Template Picker acepta solo configuraciones PUBLISHED del mismo tenant, tipo y parent aplicable.",
        "Creation Policy gobierna allowed parents, template/responsible/approval, numeración, Record Code y status PENDING.",
        "Las configuraciones reales de 06A siguen DRAFT; la pantalla muestra un bloqueo controlado y no publica automáticamente.",
    ],
)

heading(doc, "8. Common Intake y atributos dinámicos")
doc.add_paragraph(
    "El formulario es común y parametrizado: Workspace Type, Name, Description, Parent, Template, Responsible, estado "
    "previsto y números proyectados readonly. Los campos específicos se derivan del tipo vigente y abarcan Property "
    "Type/tenure/legal/geographic/value/area; Facility Type/commissioning/area/capacity/criticality; y Warehouse "
    "Type/location/storage capacity/unit/criticality. Las clasificaciones se obtienen de catálogos publicados."
)

page_break(doc)
heading(doc, "9. Preview, Submit y Approval")
table(
    doc,
    ["Fase", "Comportamiento implementado"],
    [
        ["Preview", "Calcula parent path, Business Number, Record Code, template, policy, módulos y status; persisted=false"],
        ["Submit", "Revalida tenant/type/parent/rule/template/policy/responsable/atributos/clasificaciones/módulos/versión"],
        ["Snapshot", "Persiste fingerprint lógico de request y gobernanza; submitted_hash SHA-256"],
        ["Review", "Start Review, Return con reason, Reject con reason y Resubmit"],
        ["Approve", "Four-Eyes requestor != approver y last_modified_by != approver; approval_hash"],
        ["Concurrency", "409 PHYSICAL_WORKSPACE_REQUEST_VERSION_CONFLICT; no merge silencioso"],
    ],
    [1800, 7560],
)

heading(doc, "10. Materialization transaccional")
doc.add_paragraph(
    "Materialization bloquea la solicitud y el parent, revalida gobernanza publicada y aprobación, reserva la secuencia "
    "por tipo, calcula un Record Code jerárquico único, crea el EnterpriseWorkspace PENDING, persiste metadatos, "
    "clasificaciones, module settings y assignment del responsable, enlaza la solicitud y registra SecurityEvent. "
    "Todo ocurre dentro de una única transacción PostgreSQL."
)
table(
    doc,
    ["Resultado", "Garantía"],
    [
        ["Éxito", "Request CREATED; enterprise_workspace PENDING; números finales persistidos"],
        ["Error antes de commit", "ROLLBACK; no Workspace parcial; la secuencia tampoco se consume"],
        ["IntegrityError", "Solicitud marcada FAILED en transacción separada y evento failure"],
        ["Falla genérica inyectada", "Solicitud vuelve a APPROVED; evento failure; reintento posible"],
        ["Segundo materialize", "ALREADY_CREATED; mismo workspace_id; mutation_count=0"],
        ["Concurrencia", "Números por tipo y Record Codes hermanos únicos"],
    ],
    [2100, 7260],
)
callout(
    doc,
    "No-reuse de numeración",
    "La reserva usa la misma transacción que la creación; un rollback revierte next_value. No se deja un hueco por una "
    "materialización fallida antes del commit. Un número confirmado no se reutiliza.",
    PALE_AMBER,
)

heading(doc, "11. Enterprise Explorer y Physical Workspace Overview")
bullets(
    doc,
    [
        "Enterprise Explorer combina el release CORE publicado con Project y Workspaces físicos operativos, sin duplicar nodos.",
        "Creation Requests no aparecen como nodos; solo el Workspace materializado se incorpora.",
        "Overview genérico muestra tipo, nombre, Business Number, Record Code, PENDING, parent, responsable, template, request y fecha.",
        "Atributos, clasificaciones, módulos habilitados y módulos planificados se presentan por la misma vista parametrizada.",
    ],
)

heading(doc, "12. USER MODE")
table(
    doc,
    ["Vista/acción", "Contenido"],
    [
        ["Create Physical Workspace", "Selector de tres tipos, location, template, responsible y dynamic intake"],
        ["Acciones contextuales", "Create Property / Facility / Warehouse según elegibilidad backend"],
        ["My Physical Requests", "8 filtros y acciones por estado: edit, preview, submit, cancel y open"],
        ["Physical Review Queue", "Start review, return, reject, approve y materialize con RBAC"],
        ["Physical Overview", "Landing común para PROPERTY / FACILITY / WAREHOUSE"],
        ["CompactModuleHeader", "Conservado dentro del Enterprise Explorer"],
    ],
    [2500, 6860],
)

page_break(doc)
heading(doc, "13. APIs implementadas")
table(
    doc,
    ["Método", "Ruta"],
    [
        ["GET", "/api/v1/physical-workspace-creation-requests/options"],
        ["POST / GET", "/api/v1/physical-workspace-creation-requests"],
        ["GET / PUT", "/api/v1/physical-workspace-creation-requests/{id}"],
        ["POST", "/{id}/preview · /submit · /cancel"],
        ["POST", "/{id}/start-review · /return · /reject · /approve"],
        ["POST", "/{id}/materialize"],
        ["GET", "/api/v1/physical-workspaces/{workspace_id}/overview"],
    ],
    [1600, 7760],
)

heading(doc, "14. RBAC y SecurityEvent")
table(
    doc,
    ["Rol", "Permisos / responsabilidad"],
    [
        ["Physical Workspace Requestor", "create, edit, submit, read"],
        ["Physical Workspace Reviewer", "read y review"],
        ["Physical Workspace Approver", "read y approve; sujeto a Four-Eyes"],
        ["Materialization Service", "read y materialize"],
        ["Physical Workspace Responsible", "assignment mínimo sobre el Workspace creado"],
    ],
    [2850, 6510],
)
doc.add_paragraph(
    "Los eventos cubren creación, edición, submit, cancel, start review, return, reject, approve, materialization_started, "
    "workspace_created y materialization_failed. Incluyen tenant, actor, request, tipo, parent, estado anterior/posterior "
    "y resultado."
)

heading(doc, "15. Migración")
bullets(
    doc,
    [
        "Alembic 20260812_0037 es aditiva, tenant-scoped, reversible y no modifica el release CORE.",
        "Define FKs explícitas, revision_version, unicidad tenant/request_number y materialized_workspace_id único.",
        "PostgreSQL efímero validó upgrade 0036→0037, downgrade 0037→0036 y re-upgrade a 0037.",
        "PostgreSQL persistente quedó en 20260812_0037 (head).",
    ],
)

heading(doc, "16. Pruebas y regresión")
table(
    doc,
    ["Control", "Resultado"],
    [
        ["Gate 06B PostgreSQL E2E", "31 passed · concurrencia + rollback + reversibilidad"],
        ["Gate 06B SQLite", "30 passed · 1 skipped por contrato PostgreSQL"],
        ["Frontend componente Gate 06B", "8 passed"],
        ["Regresión Gates 02A–06A + Organization", "217 passed · 4 skipped; 1 diagnóstico de montaje, luego validado"],
        ["Backend full suite", "417 passed · 6 skipped · coverage 85,92 %"],
        ["Frontend full suite", "28 files · 182 passed"],
        ["TypeScript / Vite", "PASS · producción compilada"],
        ["ESLint", "PASS · 0 errores; 9 warnings históricos dentro del budget de CI"],
        ["Prettier", "PASS"],
        ["Ruff backend completo", "PASS · check y format"],
        ["Browser localhost", "PASS · login, Explorer, acciones, selector y bloque DRAFT"],
    ],
    [2850, 6510],
)

heading(doc, "17. Protección del baseline real")
table(
    doc,
    ["Control", "Estado al cierre"],
    [
        ["Tenant", "P&P Ingeniería y Proyectos · sin cambio"],
        ["CORE Release", "ES-PYP-CORE-RECONCILED-20260809 · published · sin cambio"],
        ["CORE drafts", "0"],
        ["Workspaces", "14"],
        ["Project Workspaces", "1"],
        ["REGION / DISTRICT / SITE reales", "0 / 0 / 0"],
        ["PROPERTY / FACILITY / WAREHOUSE reales", "0 / 0 / 0"],
        ["Physical Creation Requests reales", "0"],
        ["Physical Templates", "5 DRAFT / 0 PUBLISHED"],
        ["Physical Creation Policies", "3 DRAFT"],
        ["Physical Business Number sequences", "PROPERTY/FACILITY/WAREHOUSE next_value = 1"],
        ["Project Number sequence", "next_value = 1"],
    ],
    [3300, 6060],
)

heading(doc, "18. Riesgos y deuda")
bullets(
    doc,
    [
        "Las policies y templates reales permanecen DRAFT; el flujo real debe seguir bloqueado hasta aprobación/publicación administrada.",
        "El runner completo backend en un contenedor sin Redis arroja readiness 503; al ejecutarlo en la red real la prueba pasa. Es configuración del runner, no defecto Gate 06B.",
        "La suite frontend AppFlow requiere 45 s de margen bajo carga paralela; el caso tarda 27,8 s aislado y 41,4 s en suite completa.",
        "Persisten 9 warnings ESLint históricos en el visor BIM/Project Creation; están dentro del presupuesto de CI y fuera de Gate 06B.",
        "La autorización futura puede ampliar el authorization context con restricciones por workspace_type_code más granulares sin crear permisos por tipo.",
    ],
)

heading(doc, "19. Recomendación del siguiente gate")
doc.add_paragraph(
    "Abrir Gate 06C como incremento independiente para Physical Workspace Initialization & Activation. Debe reutilizar "
    "el patrón probado de Gate 05C, mantener Four-Eyes, definir readiness, asignaciones mínimas, módulos inicializables y "
    "la transición PENDING→ACTIVE. No iniciar Asset Manager, Inventory Manager ni operación profunda de Facility/Property."
)

heading(doc, "20. Matriz de aceptación")
criteria = [
    "Generic process / PROPERTY / FACILITY / WAREHOUSE",
    "Exclusiones REGION / DISTRICT / SITE / LINEAR_ASSET / ASSET",
    "Type / Location / Template / Policy pickers",
    "Dynamic Intake / Preview / Submit / Review / Return / Reject",
    "Approval / Four-Eyes / ETag-If-Match",
    "Transactional Materialization / Business Number / Record Code",
    "Concurrency / Idempotency / Failure rollback",
    "Attributes / Classifications / Modules / Responsible Assignment",
    "Enterprise Explorer / Physical Overview / USER MODE",
    "RBAC / SecurityEvent / Migration / PostgreSQL E2E",
    "Real baseline unchanged",
]
table(doc, ["Criterio agrupado", "Resultado"], [[item, "PASS"] for item in criteria], [7900, 1460])
callout(
    doc,
    "Estado final exclusivo",
    "READY_FOR_PHYSICAL_WORKSPACE_INITIALIZATION",
    PALE_GREEN,
)
doc.add_paragraph(
    "Gate 06B termina aquí. Gate 06C, Asset Manager, Inventory Manager y cualquier operación física profunda no fueron iniciados.",
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
