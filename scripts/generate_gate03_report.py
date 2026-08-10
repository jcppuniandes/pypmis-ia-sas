from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
EVIDENCE = REPO / "artifacts" / "enterprise_structure" / "gate03"
REPORT_ROOT = Path.home() / "Documents" / "P&P" / "P&Pmis Construction AI"
SOURCE = next(
    REPORT_ROOT.glob(
        "Dise*/Resumen de Sprint/Informe_Tecnico_PPMIS_Core_Correction_Gate02B_2026-08-09.docx"
    )
)
OUTPUT = SOURCE.with_name("Informe_Tecnico_PPMIS_Core_Correction_Gate03_2026-08-10.docx")

NAVY = "0D2A3A"
TEAL = "0F9D9A"
PALE_TEAL = "EAF7F6"
PALE_BLUE = "EFF5F9"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
WHITE = "FFFFFF"
TEXT = "16222B"


def load(name: str) -> dict:
    return json.loads((EVIDENCE / name).read_text(encoding="utf-8"))


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color: str, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.bold = bold


def set_cell_margin(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        cell.text = header
        shade(cell, NAVY)
        set_cell_text_color(cell, WHITE, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margin(cell)
        if widths:
            cell.width = Inches(widths[index])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 == 0 else PALE_BLUE
        for index, value in enumerate(values):
            cells[index].text = str(value)
            shade(cells[index], fill)
            set_cell_text_color(cells[index], TEXT)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cells[index])
            if widths:
                cells[index].width = Inches(widths[index])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc: Document, title: str, body: str, fill: str = PALE_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    title_run = cell.paragraphs[0].add_run(title.upper())
    title_run.bold = True
    title_run.font.color.rgb = RGBColor.from_string(NAVY)
    title_run.font.size = Pt(9)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        run.font.color.rgb = RGBColor.from_string(TEXT)
        run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def add_bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)


pre = load("gate03_pre_publish_fingerprints.json")
post = load("gate03_post_publish_fingerprints.json")
hashes = load("gate03_hash_validation.json")
review = load("gate03_functional_review.json")
publish = load("gate03_publish.json")
second = load("gate03_publish_second.json")
audit = load("gate03_audit_validation.json")
qa = load("gate03_qa.json")
tests = load("gate03_tests.json")

doc = Document(SOURCE)
doc.core_properties.title = "Informe técnico P&Pmis Construction AI — Publish CORE Gate 03"
doc.core_properties.subject = "Publicación CORE inmutable, tenant-scoped y auditable"
doc.core_properties.comments = "Documento acumulativo Gates 00, 01, 02A, 02B y 03."

doc.paragraphs[3].text = "Publish CORE — Gates 00, 01, 02A, 02B y 03"
doc.paragraphs[4].text = (
    "Release inmutable · RBAC específico · hashes verificados · auditoría · "
    "segundo publish sin mutaciones · QA ADMIN/USER"
)
cover = doc.tables[0].cell(0, 0)
cover.text = (
    "RESULTADO ACTUAL\n"
    "GATE 03 COMPLETO · PUBLISH CORE SUCCESS · QA ADMIN/USER PASS\n"
    "RELEASE ES-PYP-CORE-RECONCILED-20260809 · 14 / 7 / 26 / 0"
)
shade(cover, PALE_GREEN)
set_cell_text_color(cover, NAVY, bold=True)
set_cell_margin(cover, 180, 180, 180, 180)

contents = doc.tables[2]
row = contents.add_row().cells
row[0].text = "42–50"
row[1].text = "Gate 03 — Publish CORE, inmutabilidad, auditoría, QA y cierre"
for cell in row:
    shade(cell, PALE_GREEN)
    set_cell_margin(cell)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

add_heading(doc, "42. Gate 03 — resultado ejecutivo")
doc.add_paragraph(
    "Se cerró formalmente el CORE de Enterprise Structure mediante una publicación explícita, separada del apply, "
    "tenant-scoped, protegida por RBAC e inmutable. La transacción no reejecutó Gate 02B ni modificó workspaces, "
    "objetivos, clasificaciones, catálogos, tenant, usuarios o estados operativos."
)
add_table(
    doc,
    ["Gate de salida", "Resultado"],
    [
        ["Prepublish inventory", "PASS · 14 workspaces / 1 raíz / 7 objetivos / 26 clasificaciones / 0 links"],
        ["Hash validation", "PASS · raw y canonical exactos"],
        ["RBAC publish", "PASS · admin.enterprise_structure.publish · organization"],
        ["Publish transaction", "SUCCESS · release ID 1 · SecurityEvent ID 11"],
        ["Inmutabilidad", "PASS · ORM + trigger PostgreSQL + bloqueo API 409"],
        ["Segundo publish", "ALREADY_PUBLISHED · 0 mutaciones · 0 eventos nuevos"],
        ["Rollback", "Plan documentado y prueba lógica PASS en base aislada"],
        ["QA ADMIN / USER", "PASS / PASS"],
        ["Pruebas", "184 pruebas automatizadas reportadas · PASS"],
    ],
    [2.25, 4.15],
)
add_callout(
    doc,
    "Alcance respetado",
    "No se cargaron EXPERIENCE, PROPERTY o FACILITY; no se inició Project Creator, CPM, XML P6 ni otro módulo. "
    "Gate 03 termina con la publicación CORE y no inicia automáticamente un nivel posterior.",
    PALE_AMBER,
)

add_heading(doc, "43. Inspección arquitectónica y extensión mínima")
doc.add_paragraph(
    "La publicación existente en ADMIN MODE gestiona revisiones draft/published de tipos de workspace y catálogos, "
    "pero no conserva un snapshot estructural CORE de 14/7/26/0. Se reutilizaron sus patrones de hashes, actor, "
    "SecurityEvent e inmutabilidad y se añadió una entidad específica EnterpriseCoreRelease. No se creó un motor "
    "universal de releases."
)
add_table(
    doc,
    ["Componente", "Responsabilidad Gate 03"],
    [
        ["EnterpriseCoreRelease", "Snapshot JSON, hashes, fingerprint, conteos, actor, fecha y release anterior"],
        ["Migration 20260810_0031", "Tabla aditiva, índices, unicidad tenant/release y trigger PostgreSQL"],
        ["importer/publish.py", "Locks, hashes, RBAC, validación exacta, publicación, replay y rollback lógico"],
        ["importer/security.py", "Autorización compartida por actor, permiso y alcance organizacional"],
        ["CLI publish", "Aprobación explícita y evidencia humana/JSON desde la misma transacción"],
        ["API ADMIN/USER", "Metadata del release publicada en las dos vistas"],
        ["Frontend", "Banner de release, actor, fecha, fingerprint y acciones estructurales deshabilitadas"],
    ],
    [2.2, 4.2],
)
add_bullets(
    doc,
    [
        "El release permite rollback lógico mediante estado unpublished; los campos de contenido permanecen inmutables.",
        "El fingerprint protegido previo no incluye la tabla de releases, evitando confundir schema aditivo con cambio de fuente.",
        "Las mutaciones estructurales API exigen una nueva revisión aprobada después de publicar.",
        "Los estados operativos de workspaces y el estado de publicación permanecen separados.",
    ],
)

add_heading(doc, "44. Prepublicación, hashes y revisión funcional")
add_heading(doc, "Fingerprint de fuente aprobado", 2)
doc.add_paragraph(pre["protected_source_hash"], style="Code Block")
add_table(
    doc,
    ["Tabla protegida", "Filas", "SHA-256 prepublish"],
    [[name, value["rows"], value["sha256"]] for name, value in pre["tables"].items()],
    [2.4, 0.7, 3.3],
)
add_heading(doc, "Validación de hashes", 2)
add_table(
    doc,
    ["Control", "Esperado", "Observado", "Resultado"],
    [
        ["Raw YAML", hashes["raw_sha256"]["expected"], hashes["raw_sha256"]["observed"], "PASS"],
        ["Canonical", hashes["canonical_hash"]["expected"], hashes["canonical_hash"]["observed"], "PASS"],
    ],
    [1.0, 2.35, 2.35, 0.7],
)
add_heading(doc, "Controles funcionales", 2)
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Tenant", "ID 1 · P&P Ingeniería y Proyectos · pyp-ingenieria-proyectos · COP"],
        ["Workspaces / raíz", "14 / 1"],
        ["External key / Record Code únicos", "14 / 14"],
        ["Objetivos / clasificaciones / links", "7 / 26 / 0"],
        ["Ciclos / padres rotos", "0 / 0"],
        ["Responsible-area obligatorias", "4 / 4"],
        ["Strategic-objective", "Solo Portfolio, Program y Project"],
        ["PROPERTY / FACILITY / EXPERIENCE", "0 / 0 / NOT LOADED"],
    ],
    [2.5, 3.9],
)

add_heading(doc, "45. Transacción de publicación")
doc.add_paragraph(
    "La publicación se ejecutó después de Alembic 0031 y de repetir el preflight. La CLI abrió una transacción "
    "independiente: bloqueó tenant y tablas del dominio, revalidó actor, hashes, fingerprint, conteos, jerarquía e "
    "identidades; creó el snapshot de release y el SecurityEvent y confirmó ambos en un solo commit."
)
add_table(
    doc,
    ["Dato publicado", "Valor"],
    [
        ["Release ID / código", f"{publish['release_id']} · {publish['release_code']}"],
        ["Estado", publish["state"]],
        ["Actor", publish["actor"]],
        ["Publicado en", publish["published_at"]],
        ["Raw hash", publish["input_hash"]],
        ["Canonical hash", publish["canonical_input_hash"]],
        ["Content fingerprint", publish["content_fingerprint"]],
        ["Snapshot", f"{publish['workspace_count']} / {publish['objective_count']} / {publish['classification_count']} / {publish['link_count']}"],
        ["SecurityEvent", publish["audit_event_id"]],
        ["Mutaciones de workspaces", 0],
    ],
    [2.0, 4.4],
)
add_callout(
    doc,
    "Fingerprint publicado",
    publish["content_fingerprint"],
    PALE_GREEN,
)

add_heading(doc, "46. Matriz de identidades y estados BEFORE / AFTER")
doc.add_paragraph(
    "Publish no cambió el estado operativo de ningún workspace: 13 nodos continúan draft y el Project ID 14 "
    "continúa active. Los IDs 1, 2 y 3 adoptados y los once IDs creados en Gate 02B permanecen intactos."
)
add_table(
    doc,
    ["ID", "External key", "Record Code", "Tipo", "BEFORE", "AFTER"],
    [
        [node["id"], node["external_key"], node["record_code"], node["type"], node["status_before"], node["status_after"]]
        for node in review["nodes"]
    ],
    [0.45, 1.55, 1.15, 1.25, 0.85, 0.85],
)
add_table(
    doc,
    ["Resumen de estados", "Antes", "Después", "Transiciones"],
    [["Workspaces", "13 draft · 1 active", "13 draft · 1 active", "0"]],
    [2.1, 1.5, 1.5, 1.3],
)

add_heading(doc, "47. RBAC, auditoría, inmutabilidad e idempotencia")
add_table(
    doc,
    ["Control", "Evidencia"],
    [
        ["Actor", "admin@demo.local · existente · active"],
        ["Permiso", "admin.enterprise_structure.publish"],
        ["Alcance", "organization"],
        ["Evento", "ID 11 · enterprise_structure.core_published · success"],
        ["Target", "EnterpriseCoreRelease ID 1"],
        ["Campos inmutables", "Protegidos en ORM y trigger PostgreSQL"],
        ["Borrado físico", "Bloqueado"],
        ["Edición estructural API", "HTTP 409 · nueva revisión requerida"],
    ],
    [2.1, 4.3],
)
add_heading(doc, "Segundo publish", 2)
add_table(
    doc,
    ["Métrica", "Primera solicitud", "Segunda solicitud"],
    [
        ["Outcome", publish["outcome"], second["outcome"]],
        ["Release ID", publish["release_id"], second["release_id"]],
        ["Mutaciones", publish["mutation_count"], second["mutation_count"]],
        ["SecurityEvent", publish["audit_event_id"], "NONE"],
        ["Fingerprint", publish["content_fingerprint"], second["content_fingerprint"]],
    ],
    [1.55, 2.4, 2.4],
)
add_callout(
    doc,
    "Idempotencia demostrada",
    "El segundo publish idéntico devolvió ALREADY_PUBLISHED, mantuvo un solo release, no creó un segundo evento y "
    "reportó mutation_count = 0.",
    PALE_GREEN,
)

add_heading(doc, "48. Rollback lógico")
doc.add_paragraph(
    "El rollback de publicación es distinto del rollback de datos aplicados. La operación lógica cambia solo el "
    "estado del release a unpublished y registra actor, fecha, razón y SecurityEvent. No elimina ni modifica los 14 "
    "workspaces. Como este es el primer release CORE, no existe release anterior."
)
add_bullets(
    doc,
    [
        "Plan completo: artifacts/enterprise_structure/gate03/gate03_rollback_plan.md.",
        "Prueba aislada: test_logical_rollback_preserves_all_applied_workspaces.",
        "Resultado: snapshot de workspaces idéntico antes y después; evento core_unpublished creado.",
        "El release real de Gate 03 permanece published; no se ejecutó rollback real durante QA.",
        "La recuperación futura debe seleccionar un release previo o crear una nueva revisión aprobada.",
    ],
)

add_heading(doc, "49. QA y calidad técnica")
add_heading(doc, "QA ADMIN MODE", 2)
add_bullets(
    doc,
    [
        "Ruta: ADMIN MODE → Enterprise Structure → Enterprise Structure Configuration.",
        "14 treeitems, una raíz, Record Codes y external keys intactos; cero errores de consola.",
        "Release, fecha, actor y fingerprint visibles en un banner compacto.",
        "Agregar nodo, Nuevo nodo y editores estructurales deshabilitados después de publicar.",
        "Una solicitud directa de creación fue rechazada HTTP 409, sin mutación.",
        "CompactModuleHeader se conserva compacto y la cinta lateral mantiene módulos y submódulos.",
    ],
)
add_heading(doc, "QA USER MODE", 2)
add_bullets(
    doc,
    [
        "Ruta: USER MODE → Enterprise Strategy Manager → Enterprise Structure & Workspace Manager → Enterprise Explorer.",
        "14 nodos, 1 Project, 0 Properties y 0 Facilities; jerarquía y Record Codes coincidentes.",
        "Release publicado y fingerprint visibles sin exponer actor de edición ni acciones ADMIN.",
        "Filtro Project validado con un resultado y nodo Desarrollo P&Pmis Construction AI visible.",
        "Objetivos, categorías, vistas Árbol/Tabla, búsqueda y filtros permanecen disponibles.",
    ],
)
add_heading(doc, "Pruebas ejecutadas", 2)
add_table(
    doc,
    ["Suite / control", "Resultado"],
    [
        ["Backend apply + publish", "24 passed"],
        ["Backend Enterprise Structure + importer", "20 passed"],
        ["Casos publish obligatorios", "20 controles PASS"],
        ["Ruff", "PASS"],
        ["Frontend Vitest", "23 files · 140 passed"],
        ["ESLint", "PASS · 0 errores; 8 warnings preexistentes del visor BIM"],
        ["Prettier", "PASS"],
        ["TypeScript + Vite", "PASS"],
        ["Alembic PostgreSQL 0030 → 0031", "PASS en base temporal; 0031 head en localhost"],
    ],
    [2.8, 3.6],
)

add_heading(doc, "50. Evidencias, riesgos y cierre")
add_heading(doc, "Paquete de evidencias Gate 03", 2)
add_table(
    doc,
    ["Evidencia", "Propósito"],
    [
        ["gate03_pre_publish_inventory.json / fingerprints.json", "Baseline real y hash de fuente aprobado"],
        ["gate03_hash_validation.json", "Raw/canonical hashes exactos"],
        ["gate03_functional_review.json", "Árbol, identidades, conteos y estados BEFORE/AFTER"],
        ["gate03_publish.txt / .json / .sha256", "Primera publicación y checksum"],
        ["gate03_publish_second.txt / .json", "Replay ALREADY_PUBLISHED sin mutaciones"],
        ["gate03_post_publish_inventory.json / fingerprints.json", "Estado persistido después de publish y QA"],
        ["gate03_audit_validation.json", "Evento, RBAC, fingerprint e inmutabilidad"],
        ["gate03_rollback_plan.md", "Recuperación lógica no destructiva"],
        ["gate03_qa.json / gate03_tests.json", "QA API/UI y suites automatizadas"],
    ],
    [3.0, 3.4],
)
add_heading(doc, "Riesgos pendientes", 2)
add_bullets(
    doc,
    [
        "Los estados operativos siguen 13 draft y 1 active por diseño; no deben confundirse con el estado published del release.",
        "Modificar la estructura exige definir y aprobar formalmente un nuevo release; no existe edición in-place.",
        "Esta es la primera publicación CORE y no existe release anterior para retorno automático.",
        "El repositorio conserva ocho warnings ESLint no bloqueantes en BimIfcModelViewer, fuera del alcance Gate 03.",
        "La cadena Alembic histórica parte de un schema base creado por la aplicación; la migración específica 0030→0031 sí fue validada limpiamente en PostgreSQL.",
    ],
)
add_heading(doc, "Recomendación del siguiente incremento", 2)
doc.add_paragraph(
    "El siguiente incremento recomendado es diseñar el flujo formal de nueva revisión CORE: crear draft desde el "
    "release publicado, comparar diff, aprobar hashes y publicar un release sucesor con previous_release. Debe abrirse "
    "como gate independiente. No cargar EXPERIENCE, PROPERTY_FACILITY ni iniciar otros módulos sin autorización expresa."
)
add_table(
    doc,
    ["Gate de salida", "Estado final"],
    [
        ["Gate 03", "COMPLETO"],
        ["Publish CORE", "EJECUTADO Y VALIDADO"],
        ["Release", "ES-PYP-CORE-RECONCILED-20260809 · published"],
        ["Idempotencia", "ALREADY_PUBLISHED · 0 mutaciones"],
        ["Inmutabilidad / auditoría", "PASS / PASS"],
        ["QA ADMIN / USER", "PASS / PASS"],
        ["EXPERIENCE / PROPERTY_FACILITY", "NOT LOADED / NOT LOADED"],
        ["Próxima autoridad requerida", "Nuevo gate para revisión CORE sucesora"],
    ],
    [2.35, 4.05],
)

doc.save(OUTPUT)
print(OUTPUT)
