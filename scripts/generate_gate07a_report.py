from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUT = REPO / "outputs" / "technical_reports" / "Informe_Tecnico_Gate_07A_Idea_Demand_Manager_20260813.docx"

# Preset: standard_business_brief. First-page pattern: memo_masthead.
NAVY, BLUE, DARK_BLUE, INK, MUTED = "0D2A3A", "2E74B5", "1F4D78", "16222B", "5B6870"
WHITE, LIGHT, PALE_GREEN, PALE_AMBER, PALE_TEAL = "FFFFFF", "F2F4F7", "EAF7EF", "FFF5E6", "EAF7F6"


def set_run(run, *, size=11, color=INK, bold=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size, run.font.color.rgb = Pt(size), RGBColor.from_string(color)
    run.bold, run.italic = bold, italic


def shade(cell, fill):
    shd = cell._tc.get_or_add_tcPr().find(qn("w:shd")) or OxmlElement("w:shd")
    if shd.getparent() is None:
        cell._tc.get_or_add_tcPr().append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for name, value in (("tblW", total), ("tblInd", indent)):
        node = tbl_pr.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style, table.alignment = "Table Grid", WD_TABLE_ALIGNMENT.LEFT
    header_row = table.rows[0]
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_row._tr.get_or_add_trPr().append(header_flag)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text, cell.vertical_alignment = header, WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, NAVY)
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cells[index], WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margin(cells[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=9.1)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(doc, title, body, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_flag)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    set_run(cell.paragraphs[0].add_run(title.upper()), size=9, color=NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run(run, size=9.5)
    set_table_geometry(table, [9360], indent=160)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def bullets(doc, values):
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)


doc = Document()
section = doc.sections[0]
section.page_width, section.page_height = Inches(8.5), Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)
styles = doc.styles
normal = styles["Normal"]
normal.font.name, normal.font.size, normal.font.color.rgb = "Calibri", Pt(11), RGBColor.from_string(INK)
normal.paragraph_format.space_after, normal.paragraph_format.line_spacing = Pt(6), 1.10
for name, size, color, before, after in (
    ("Title", 23, NAVY, 0, 4), ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 16, BLUE, 16, 8), ("Heading 2", 13, BLUE, 12, 6), ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name, style.font.size, style.font.color.rgb = "Calibri", Pt(size), RGBColor.from_string(color)
    style.paragraph_format.space_before, style.paragraph_format.space_after = Pt(before), Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name, styles[name].font.size = "Calibri", Pt(11)
    styles[name].paragraph_format.space_after, styles[name].paragraph_format.line_spacing = Pt(8), 1.167
    styles[name].paragraph_format.left_indent, styles[name].paragraph_format.first_line_indent = Inches(0.5), Inches(-0.25)

header = section.header.paragraphs[0]
set_run(header.add_run("P&Pmis Construction AI  |  Gate 07A Technical Report"), size=8.5, color=MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer.add_run("Controlled technical baseline  |  13 Aug 2026  |  Page "), size=8, color=MUTED)
page_run = footer.add_run()
set_run(page_run, size=8, color=MUTED)
for field_type, value in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
    node = OxmlElement("w:fldChar" if field_type else ("w:instrText" if value == " PAGE " else "w:t"))
    if field_type:
        node.set(qn("w:fldCharType"), field_type)
    else:
        node.set(qn("xml:space"), "preserve")
        node.text = value
    page_run._r.append(node)

doc.core_properties.title = "Informe técnico Gate 07A - Idea & Demand Manager"
doc.core_properties.subject = "Foundation and governed Idea Lifecycle"
doc.core_properties.author = "P&Pmis Construction AI Engineering"
doc.core_properties.comments = "Gate 07A cerrado; sin Project Proposal, SECOP Bidder ni Idea Workspace."

doc.add_paragraph("INFORME TÉCNICO", style="Title")
doc.add_paragraph("Gate 07A · Idea & Demand Manager Foundation / Idea Lifecycle", style="Subtitle")
for label, value in (
    ("Tenant", "P&P Ingeniería y Proyectos"),
    ("Fecha de cierre", "13 de agosto de 2026"),
    ("Ruta USER", "Enterprise Strategy Manager → Idea & Demand Manager → Idea Lifecycle"),
    ("Ruta ADMIN", "Enterprise Strategy Manager → Idea & Demand Manager"),
    ("Estado final", "READY_FOR_PROJECT_PROPOSAL"),
):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    set_run(paragraph.add_run(f"{label}: "), bold=True)
    set_run(paragraph.add_run(value))
callout(doc, "Resultado de cierre", "Se implementó un único ciclo gobernado para registrar, someter, filtrar, enrutar, asignar, evaluar y decidir ideas. La salida aceptada sólo declara readiness para una futura Project Proposal; Gate 07A no crea propuestas ni workspaces.", PALE_GREEN)

heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph("Gate 07A incorpora Idea & Demand Manager sobre la arquitectura empresarial existente. Cada Idea pertenece a un Enterprise Workspace, Business Unit o Portfolio mediante owning_workspace_id, conserva trazabilidad tenant/workspace, aplica configuración publicada con herencia y registra evaluaciones inmutables por versión.")
add_table(doc, ["Control", "Resultado"], [
    ["Ciclo único", "PASS · 11 estados controlados"], ["Persistencia", "PASS · sólo ideas + idea_evaluations"],
    ["RBAC", "PASS · 12 permisos + 5 roles"], ["Concurrencia", "PASS · ETag + secuencia atómica"],
    ["PostgreSQL", "PASS · 3/3 + upgrade/downgrade/upgrade"], ["Frontend", "PASS · TypeScript, lint, build, 2/2 tests"],
    ["Browser localhost", "PASS · Idea Lifecycle y formulario validados"], ["Límite", "PASS · sin SECOP, Proposal ni Idea Workspace"],
], [2850, 6510])

heading(doc, "2. REUSE BEFORE CREATE")
bullets(doc, [
    "EnterpriseWorkspace sigue siendo la identidad canónica; Idea es un registro de proceso relacionado, no un Workspace.",
    "AdminConfiguration conserva draft/publish, revisión e inmutabilidad; se añadieron dos kinds especializados.",
    "AdminNumberSequence emite IDEA-00001 de forma tenant-scoped y atómica; el preview no consume números.",
    "EnterpriseStrategicObjective se reutiliza como clasificación no jerárquica one/multiple.",
    "SecurityRole, PermissionCatalog, assignments y SecurityEvent se reutilizan para acceso y auditoría.",
    "Los adjuntos se referencian en la Idea; no se duplicó el gestor documental Project-scoped existente.",
])

heading(doc, "3. Alcance y límites")
add_table(doc, ["Incluido", "Excluido explícitamente"], [
    ["Register, intake, screening, routing y owner", "SECOP Bidder"],
    ["Matriz de evaluación versionada", "Project Proposal Creation"],
    ["Accept / reject / return", "Idea Workspace"],
    ["Proposal readiness + mapping preview", "Operación profunda de Project Proposal"],
], [4680, 4680])

heading(doc, "4. Arquitectura implementada")
doc.add_paragraph("El módulo backend app/modules/idea_demand concentra modelos, schemas, servicio y router. El frontend features/idea-demand separa API, contratos, USER workspace, ADMIN view y estilos. La navegación de App.tsx consume el blueprint sin crear motores paralelos.")

heading(doc, "5. Modelo de datos")
add_table(doc, ["Tabla", "Propósito", "Controles"], [
    ["ideas", "Cabecera y estado del ciclo", "tenant, owning workspace, version, timestamps"],
    ["idea_evaluations", "Snapshot por versión", "unique idea/version; update/delete bloqueados"],
], [2100, 3400, 3860])

heading(doc, "6. Identidad y pertenencia")
doc.add_paragraph("La Idea usa idea_number como número de negocio y owning_workspace_id como dueño operativo. El Workspace permitido es ENTERPRISE, BUSINESS-UNIT o PORTFOLIO. target_portfolio_workspace_id es opcional e independiente, por lo que enrutar no cambia el dueño.")

heading(doc, "7. Ciclo de vida")
add_table(doc, ["Estado", "Transición principal"], [
    ["DRAFT", "create/edit → submit"], ["SUBMITTED", "screen o return"], ["SCREENING", "checklist → route → assign owner"],
    ["RETURNED", "edit → DRAFT → resubmit"], ["OWNER_ASSIGNED", "start evaluation"], ["UNDER_EVALUATION", "complete immutable snapshot"],
    ["EVALUATED", "accept / reject / return"], ["ACCEPTED", "proposal readiness only"], ["REJECTED", "terminal"],
    ["CANCELLED", "terminal"], ["ARCHIVED", "terminal/read only"],
], [2800, 6560])

heading(doc, "8. Intake y screening")
doc.add_paragraph("El screening usa un checklist configurado. Todos los ítems blocking deben aprobarse antes del routing o asignación. La devolución exige razón, conserva la evidencia anterior y permite reingreso mediante edición y resubmit.")

heading(doc, "9. Routing y owner")
doc.add_paragraph("El routing persiste route_code, notas, actor y timestamp. La asignación valida usuario activo, tenant, permiso idea.evaluate y alcance organization/owning workspace. Se impide asignar antes de completar screening y routing.")

heading(doc, "10. Evaluación versionada")
doc.add_paragraph("La matriz publicada define criterios, escala, pesos y threshold. Complete evaluation exige exactamente una calificación por criterio, calcula score ponderado y guarda matrix_configuration_id, revisión, matriz completa, ratings y resultado. La tabla no admite UPDATE ni DELETE.")

heading(doc, "11. Decisión y re-evaluación")
doc.add_paragraph("Aceptar/rechazar exige estado EVALUATED. Return preserva el snapshot y reinicia el flujo sin sobrescribirlo; una evaluación posterior obtiene evaluation_version incremental. accepted_evaluation_id enlaza el snapshot vigente para decisión/readiness.")

heading(doc, "12. Proposal readiness")
doc.add_paragraph("El endpoint de readiness requiere ACCEPTED, evaluación aceptada y objetivo estratégico. Entrega mapping_preview determinista y status READY_FOR_PROJECT_PROPOSAL o GATE07A_REWORK_REQUIRED. can_create_project_proposal permanece false: no existe endpoint de creación.")

heading(doc, "13. Configuración ADMIN")
bullets(doc, [
    "idea_demand_configuration gobierna tipos, categorías, objetivos, checklist, routing y proposal mapping.",
    "idea_evaluation_matrix gobierna escala, criterios, pesos y recommendation threshold.",
    "Published es inmutable; ADMIN puede clone-to-draft, editar con If-Match y publicar una nueva revisión.",
    "Overrides workspace-{id} se resuelven desde raíz hacia owning Workspace; el ancestro más cercano tiene precedencia.",
])

heading(doc, "14. RBAC")
add_table(doc, ["Rol", "Permisos principales"], [
    ["idea_requestor", "read/create/edit/submit/cancel"], ["idea_intake_reviewer", "read/screen/route/assign_owner"],
    ["idea_owner", "read/evaluate"], ["idea_decision_maker", "read/decide"],
    ["idea_configuration_admin", "read/admin.configure/admin.publish"],
], [3200, 6160])

heading(doc, "15. Seguridad e aislamiento")
bullets(doc, [
    "Toda consulta filtra tenant_id; lecturas no autorizadas se presentan como 404 para evitar enumeración.",
    "Organization scope ve todo el tenant; workspace scope sólo ve owning workspace, solicitudes propias o asignadas.",
    "If-Match obligatorio retorna 412 ETAG_MISMATCH ante versión obsoleta.",
    "Mutaciones de transición son idempotentes cuando el estado/actor objetivo ya coincide.",
])

heading(doc, "16. Auditoría")
doc.add_paragraph("SecurityEvent registra created, updated, submitted, screened, routed, owner_assigned, evaluation_started, evaluated, accepted, rejected, returned y cancelled, además de clone/update/publish de configuración. Cada evento conserva estado anterior/posterior y metadata pertinente.")

heading(doc, "17. Contratos API USER")
add_table(doc, ["Método/ruta", "Función"], [
    ["GET/POST /api/v1/ideas", "listar / crear"], ["GET/PUT /ideas/{id}", "detalle / editar con ETag"],
    ["POST /submit|screen|route|assign-owner", "intake controlado"], ["POST /evaluation/start|complete", "evaluación"],
    ["POST /accept|reject|return|cancel", "decisión/transiciones"], ["GET /proposal-readiness|history", "readiness y auditoría"],
], [4200, 5160])

heading(doc, "18. Contratos API ADMIN")
add_table(doc, ["Método/ruta", "Función"], [
    ["GET /ideas/admin/configurations/list", "revisiones Idea/matrix"], ["POST /admin/configuration/preview", "configuración efectiva y fuentes"],
    ["POST /admin/configurations/{id}/clone", "nueva DRAFT"], ["PUT /admin/configurations/{id}", "editar DRAFT con ETag"],
    ["POST /admin/configurations/{id}/publish", "publicar revisión validada"],
], [5000, 4360])

heading(doc, "19. Frontend USER MODE")
doc.add_paragraph("Idea Lifecycle reemplaza las pantallas vacías separadas. Incluye ocho queues, búsqueda server-side, registro, detalle, métricas, drawer de creación, número preview, selectores de Workspace/Portfolio/objetivos y acciones dependientes del estado/rol.")

heading(doc, "20. Frontend ADMIN MODE")
doc.add_paragraph("Enterprise Strategy Manager → Idea & Demand Manager muestra configuraciones, estado, revisiones y acciones clone/edit/publish. La matriz editada se valida en backend; la publicación conserva la inmutabilidad de revisiones usadas.")

heading(doc, "21. Integración Workspace")
doc.add_paragraph("Workspace Navigator ahora soporta Enterprise, Business Unit y Portfolio con la ruta Ideas protegida por idea.read. Desde estas rutas el componente aplica owning_workspace_id, conservando el mismo dominio y evitando un segundo módulo.")

heading(doc, "22. Navegación y limpieza funcional")
bullets(doc, [
    "USER: un único submódulo Idea Lifecycle dentro de Idea & Demand Manager.",
    "ADMIN: Idea & Demand Manager agrupado bajo Enterprise Strategy Manager.",
    "SECOP Bidder fue eliminado del tipo ControlFlowView, blueprint, pantallas vacías y tests.",
    "Idea Register, Intake y Evaluation Matrix dejaron de ser submódulos fragmentados; son etapas de una pantalla principal.",
])

heading(doc, "23. Migración")
doc.add_paragraph("Alembic 20260813_0040 es aditiva y revisa 0039. Upgrade crea índices, FKs, unique tenant/number y unique tenant/idea/version. Downgrade elimina primero idea_evaluations y después ideas. El compose Gate07A probó upgrade, downgrade y re-upgrade en PostgreSQL 16.")

heading(doc, "24. Evidencia backend")
add_table(doc, ["Ejecución", "Resultado"], [
    ["Ruff focal", "PASS"], ["pytest test_idea_demand_gate07a.py", "3 passed"],
    ["PostgreSQL compose gate07a", "3 passed · migrations reversible"], ["App import", "PASS"],
], [5000, 4360])

heading(doc, "25. Evidencia frontend")
add_table(doc, ["Ejecución", "Resultado"], [
    ["ESLint focal", "PASS"], ["TypeScript --noEmit", "PASS"], ["Vite production build", "PASS · 2358 modules"],
    ["IdeaLifecycleWorkspace.test", "2 passed"], ["AppFlow assertions", "actualizadas a ciclo único; SECOP ausente"],
], [5000, 4360])

heading(doc, "26. Evidencia PostgreSQL")
doc.add_paragraph("docker-compose.gate07a.yml levantó una base efímera PostgreSQL 16, creó baseline, retiró las tablas nuevas, selló 0039, ejecutó 0040, corrió pruebas, bajó a 0039 y volvió a 0040. Resultado de proceso: exit code 0.")

heading(doc, "27. Evidencia browser E2E")
bullets(doc, [
    "Login admin en http://localhost:5173/app.", "Idea & Demand Manager despliega exclusivamente Idea Lifecycle.",
    "La pantalla presenta métricas, ocho queues, búsqueda, preview y estados vacíos explícitos.",
    "New Idea abre drawer con número no editable, owning workspace, portfolio, tipo, categoría, valor y objetivos.",
])

heading(doc, "28. Despliegue local")
add_table(doc, ["Servicio", "Estado"], [
    ["Frontend", "HTTP 200 · http://127.0.0.1:5173/app"], ["API", "HTTP 200 · http://127.0.0.1:8000/api/v1/health"],
    ["PostgreSQL", "healthy · migration 0040"], ["Redis", "healthy"], ["Worker / Beat", "running"],
], [2600, 6760])

heading(doc, "29. Compatibilidad con baseline 06D")
doc.add_paragraph("No se modificaron Project Creation, Physical Workspace Creation, initialization/activation, Enterprise CORE releases ni datos de Project Controls. Workspace Context se amplió de forma aditiva para los tres tipos estratégicos y conserva el comportamiento 06D de Project/Property/Facility/Warehouse.")

heading(doc, "30. Integridad y concurrencia")
bullets(doc, [
    "AdminNumberSequence usa UPDATE ... RETURNING para reserva atómica.", "Preview consulta next_value sin mutarlo.",
    "SQLAlchemy version_id_col detecta escrituras concurrentes adicionales.", "API exige ETag en todas las mutaciones de ciclo/configuración.",
    "Unique constraints impiden números y versiones de evaluación duplicados.",
])

heading(doc, "31. Attachment reuse decision")
doc.add_paragraph("DocumentAttachment requiere project_id/document_id y no es un repositorio empresarial genérico. Para no acoplar Ideas a Project ni crear una tercera tabla, Gate 07A conserva attachment_refs_json como referencias/evidencia. Una evolución documental empresarial deberá generalizarse deliberadamente en un gate posterior.")

heading(doc, "32. Configuración heredada")
doc.add_paragraph("La resolución comienza en default publicado, recorre el path Enterprise → Business Unit → Portfolio y toma el override workspace-{id} más cercano con inherit_to_descendants. El preview retorna effective, path y source configuration/revision para explicar el resultado.")

heading(doc, "33. Objectives classification")
doc.add_paragraph("strategic_objective_codes es una lista no jerárquica contra EnterpriseStrategicObjective activos. objective_selection puede ser one o multiple; la API rechaza códigos inactivos/desconocidos y combinaciones superiores a la política efectiva.")

heading(doc, "34. Screening checklist")
add_table(doc, ["Ítem default", "Bloquea"], [
    ["complete_description", "Sí"], ["benefit_identified", "Sí"], ["workspace_confirmed", "Sí"], ["no_duplicate", "Sí"],
], [7000, 2360])

heading(doc, "35. Matriz default")
add_table(doc, ["Criterio", "Peso"], [["Strategic alignment", "30%"], ["Expected value", "25%"], ["Feasibility", "20%"], ["Risk response", "15%"], ["Urgency", "10%"]], [7000, 2360])

heading(doc, "36. Idempotencia")
doc.add_paragraph("Submit, owner assignment, start/complete evaluation, accept/reject, return y cancel retornan la representación vigente cuando una repetición ya alcanzó el mismo resultado. No duplican evaluaciones ni auditoría por simple reintento después de una respuesta perdida.")

heading(doc, "37. Observabilidad")
doc.add_paragraph("Las rutas heredan request_id, logs JSON y middleware de tiempos. SecurityEvent habilita historia por Idea y seguimiento de cambios de configuración sin añadir un subsistema de auditoría alternativo.")

heading(doc, "38. Riesgos y deuda controlada")
bullets(doc, [
    "El tenant local conserva Enterprise/BU/Portfolio en DRAFT; Gate 07A permite capturar ideas en ese baseline configurado.",
    "attachment_refs_json exige que el consumidor valide la existencia de evidencia externa; no gestiona binarios.",
    "La UI usa acciones rápidas con valores default para el piloto; formularios avanzados de rating/justificación pueden enriquecerse después.",
    "No se ejecutó ni modeló Project Proposal, conforme al stop boundary.",
])

heading(doc, "39. Matriz de aceptación")
criteria = [
    "Ciclo único de Idea", "Idea no es Workspace", "Owning Workspace 1:N", "Tipos de owning válidos", "Target Portfolio separado",
    "Objetivos one/multiple", "Herencia con precedencia", "Number preview no consume", "Reserva concurrente segura", "11 estados",
    "Screening checklist", "Return/resubmit", "Routing persistente", "Owner validation", "Matriz versionada publicada",
    "Evaluation snapshots inmutables", "Re-evaluation sin overwrite", "Accept/reject/return", "Proposal readiness only", "Attachment reuse",
    "12 permisos", "5 roles", "ETag mutaciones", "Idempotencia", "SecurityEvent history", "USER queues", "Create drawer",
    "Detail/allowed actions", "ADMIN clone/edit/publish", "Workspace Ideas route", "SECOP excluido", "Project Proposal excluida",
    "Migración reversible", "PostgreSQL gate", "Backend tests", "Frontend tests", "Production build", "Browser localhost",
    "Servicios healthy", "Informe técnico Word", "Stop boundary 07A",
]
add_table(doc, ["Criterio", "Resultado"], [[item, "PASS"] for item in criteria], [7850, 1510])

heading(doc, "40. Archivos principales")
bullets(doc, [
    "backend/app/modules/idea_demand/{models,schemas,service,router}.py",
    "backend/alembic/versions/20260813_0040_gate07a_idea_lifecycle.py",
    "backend/tests/test_idea_demand_gate07a.py y docker-compose.gate07a.yml",
    "frontend/src/features/idea-demand/*",
    "frontend/src/navigation/applicationNavigation.ts y App.tsx",
])

heading(doc, "41. Operación básica")
bullets(doc, [
    "USER MODE: abrir Enterprise Strategy Manager → Idea & Demand Manager → Idea Lifecycle.",
    "New Idea: escoger owning Workspace, clasificaciones y objetivos; Create draft asigna número definitivo.",
    "Intake: Submit → Complete screening → Route → Assign owner.",
    "Owner: Start evaluation → Complete evaluation; Decision Maker: Accept/Reject/Return.",
    "ADMIN MODE: clonar la revisión publicada, editar JSON validado y publicar.",
])

heading(doc, "42. Invariantes")
bullets(doc, [
    "Una Idea tiene exactamente un owning_workspace_id.", "Una evaluación pertenece a una sola Idea y versión.",
    "Una matriz publicada nunca se modifica in-place.", "Una Idea no materializa Workspace ni Project Proposal.",
    "El tenant actual nunca cruza consultas o acciones.",
])

heading(doc, "43. Decisiones de diseño")
add_table(doc, ["Decisión", "Motivo"], [
    ["Dos tablas", "Cumplir mínimo persistente y evitar duplicación"], ["JSON para evidencia/config snapshots", "Trazabilidad sin tablas accesorias"],
    ["Uppercase states", "Contrato explícito del prompt"], ["ETag 412", "Semántica HTTP para precondición fallida"],
    ["Shared component global/workspace", "Un dominio, dos puntos de entrada"],
], [3500, 5860])

heading(doc, "44. Pruebas de límites")
bullets(doc, [
    "Owning Workspace inválido retorna 422.", "ETag obsoleto retorna 412.", "Preview consecutivo mantiene el mismo número.",
    "Endpoint create-project-proposal no existe (404).", "Intento de UPDATE a IdeaEvaluation lanza ValueError y hace rollback.",
])

heading(doc, "45. Siguiente gate permitido")
doc.add_paragraph("Gate 07B podrá consumir únicamente Ideas ACCEPTED con READY_FOR_PROJECT_PROPOSAL para crear un proceso Project Proposal separado. Debe respetar accepted_evaluation_id y source_idea_id, sin convertir Idea en Workspace ni alterar snapshots de Gate 07A.")

heading(doc, "46. No iniciado")
bullets(doc, ["Project Proposal model/API/UI", "SECOP Bidder", "Idea Workspace", "Portfolio prioritization/gate decision", "Binary attachment repository empresarial"])

heading(doc, "47. Estado operacional")
add_table(doc, ["Componente", "Estado"], [
    ["Domain/API", "Operacional"], ["USER UI", "Operacional"], ["ADMIN UI", "Operacional"], ["Migration", "Aplicada en localhost"],
    ["Gate tests", "PASS"], ["Localhost", "5173 frontend / 8000 API"],
], [3600, 5760])

heading(doc, "48. Cierre")
callout(doc, "Estado final exclusivo", "READY_FOR_PROJECT_PROPOSAL", PALE_GREEN)
doc.add_paragraph("Gate 07A termina aquí. El ciclo de Idea está listo para recibir una futura integración Project Proposal, pero dicha integración no fue iniciada. Si una Idea no cumple la decisión y readiness, el resultado permitido es GATE07A_REWORK_REQUIRED.")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
