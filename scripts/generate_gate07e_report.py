from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUT = (
    REPO
    / "outputs"
    / "technical_reports"
    / "Informe_Tecnico_Gate_07E_Portfolio_Evaluation_20260820.docx"
)
LOGO = REPO / "frontend" / "public" / "pypmis-construction-ai-logo.png"

# Exact preset: standard_business_brief. First-page pattern: memo_masthead.
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "0B2545"
BLACK = "000000"
MUTED = "5F6B76"
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
CALLOUT = "F4F6F9"
POSITIVE = "EAF7EF"
CAUTION = "FFF5E6"
RISK = "FDECEC"
TEAL = "0D8F8B"


def set_run(run, *, size=11, color=BLACK, bold=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    node = tc_pr.find(qn("w:shd"))
    if node is None:
        node = OxmlElement("w:shd")
        tc_pr.append(node)
    node.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.first_child_found_in("w:tcMar")
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", indent)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header = table.rows[0]
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    header._tr.get_or_add_trPr().append(repeat)
    for index, value in enumerate(headers):
        cell = header.cells[index]
        cell.text = value
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, LIGHT)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9.5, color=INK, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if row_index % 2:
                shade(cell, "FAFBFC")
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                paragraph.paragraph_format.line_spacing = 1.0
                for run in paragraph.runs:
                    set_run(run, size=9.2, color=BLACK)
    set_table_geometry(table, widths)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc, title, body, *, fill=CALLOUT):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(repeat)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margins(cell, 120, 145, 120, 145)
    set_run(
        cell.paragraphs[0].add_run(title.upper()), size=9, color=DARK_BLUE, bold=True
    )
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.1
    for run in paragraph.runs:
        set_run(run, size=10.2, color=BLACK)
    set_table_geometry(table, [9360], indent=120)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("P&Pmis Construction AI  |  "), size=8, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instruction = OxmlElement("w:instrText")
    instruction.set(qn("xml:space"), "preserve")
    instruction.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = "1"
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instruction, separate, text, end])


def add_bottom_rule(paragraph, color=BLUE, size=14):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def add_numbering(doc, *, bullet=True):
    numbering = doc.part.numbering_part.element
    abstract_ids = [
        int(node.get(qn("w:abstractNumId")))
        for node in numbering.findall(qn("w:abstractNum"))
    ]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1
    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "bullet" if bullet else "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "•" if bullet else "%1.")
    level.append(lvl_text)
    suffix = OxmlElement("w:suff")
    suffix.set(qn("w:val"), "tab")
    level.append(suffix)
    p_pr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    p_pr.append(tabs)
    indent = OxmlElement("w:ind")
    indent.set(qn("w:left"), "720")
    indent.set(qn("w:hanging"), "360")
    p_pr.append(indent)
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:after"), "160")
    spacing.set(qn("w:line"), "280")
    spacing.set(qn("w:lineRule"), "auto")
    p_pr.append(spacing)
    level.append(p_pr)
    r_pr = OxmlElement("w:rPr")
    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")
    r_pr.append(fonts)
    level.append(r_pr)
    abstract.append(level)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_list_item(doc, text, num_id):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(8)
    paragraph.paragraph_format.line_spacing = 1.167
    p_pr = paragraph._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num = OxmlElement("w:numId")
    num.set(qn("w:val"), str(num_id))
    num_pr.extend([ilvl, num])
    p_pr.append(num_pr)
    set_run(paragraph.add_run(text), size=11, color=BLACK)
    return paragraph


def add_heading(doc, text, level=1, *, page_break=False):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    paragraph.paragraph_format.page_break_before = page_break
    return paragraph


def add_body(doc, text):
    paragraph = doc.add_paragraph(text)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    return paragraph


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(1)
section.right_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.header_distance = Inches(0.492)
section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(BLACK)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.1
for name, size, color, before, after in (
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.bold = True
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True

language = OxmlElement("w:lang")
language.set(qn("w:val"), "es-CO")
styles["Normal"]._element.get_or_add_rPr().append(language)
doc.core_properties.title = "Informe técnico Gate 07E-H — Browser E2E Release Closeout"
doc.core_properties.subject = "Implementación, endurecimiento y validación de liberación de Gate 07E en P&Pmis Construction AI"
doc.core_properties.author = "P&Pmis Construction AI / Codex"
doc.core_properties.keywords = (
    "Gate 07E, Portfolio Evaluation, Prioritization, PostgreSQL, P&Pmis"
)

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(0)
set_run(
    header.add_run("P&Pmis Construction AI  |  Gate 07E-H Release Closeout"),
    size=8,
    color=MUTED,
    bold=True,
)
add_page_number(section.footer.paragraphs[0])

bullet_num = add_numbering(doc, bullet=True)

logo_paragraph = doc.add_paragraph()
logo_paragraph.paragraph_format.space_after = Pt(8)
if LOGO.exists():
    picture = logo_paragraph.add_run().add_picture(str(LOGO), width=Inches(0.62))
    picture._inline.docPr.set("descr", "Logotipo de P&Pmis Construction AI")

title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(4)
set_run(title.add_run("INFORME TÉCNICO"), size=23, color=BLACK, bold=True)
subtitle = doc.add_paragraph()
subtitle.paragraph_format.space_after = Pt(16)
set_run(
    subtitle.add_run("Gate 07E-H — Browser E2E Release Closeout"),
    size=14,
    color=MUTED,
    bold=True,
)
metadata = [
    ("Aplicación", "P&Pmis Construction AI"),
    ("Ámbito", "USER MODE + ADMIN MODE · Portfolio / Project Workspace"),
    ("Baseline", "Gate 07A–07D, Project Governance multi-source, Gate 05B/05C"),
    ("Migración", "20260820_0045, sucesora aditiva de 20260820_0044"),
    ("Fecha de corte", "20 de agosto de 2026"),
    ("Salida funcional", "READY_FOR_PORTFOLIO_ANALYSIS"),
    ("Estado exclusivo de release", "GATE07E_RELEASE_VALIDATED"),
]
for label, value in metadata:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.line_spacing = 1.0
    set_run(paragraph.add_run(f"{label}: "), size=10.5, color=BLACK, bold=True)
    set_run(paragraph.add_run(value), size=10.5, color=BLACK)
rule = doc.add_paragraph()
rule.paragraph_format.space_before = Pt(8)
rule.paragraph_format.space_after = Pt(12)
add_bottom_rule(rule)

add_callout(
    doc,
    "Resultado ejecutivo",
    "Gate 07E quedó implementado y endurecido como una capa contextual de evaluación sobre el Project Workspace y su PortfolioMembership existentes. El cierre 07E-H aprobó tres pasadas E2E consecutivas en browser, regresión integral, migración reversible y restauración exacta de baseline. La salida funcional permanece READY_FOR_PORTFOLIO_ANALYSIS y la decisión exclusiva de release es GATE07E_RELEASE_VALIDATED.",
    fill=POSITIVE,
)

add_table(
    doc,
    ["Control", "Resultado", "Evidencia resumida"],
    [
        ["Backend", "PASS", "486 pruebas aprobadas, 13 omitidas y cobertura 85,35%."],
        [
            "PostgreSQL",
            "PASS",
            "69 regresiones + 9 escenarios Gate 07E; migración reversible.",
        ],
        [
            "Frontend",
            "PASS",
            "36 archivos / 197 pruebas; build, lint y formato verificados.",
        ],
        [
            "Browser USER / ADMIN",
            "PASS",
            "Tres flujos integrales consecutivos; 3/3 FLOW_PASS.",
        ],
        [
            "Baseline persistente",
            "PASS",
            "14 Workspaces, 1 Project Workspace, 0 requests y 0 memberships conservados.",
        ],
    ],
    [1900, 1300, 6160],
)

sections = [
    (
        "1. Executive summary",
        [
            "La solución agrega evaluación estructurada, priorización contextual y readiness de portafolio reutilizando el dominio empresarial existente. La unidad de análisis es Portfolio + Project + membership activa; la identidad canónica del Project no cambia.",
            "El release incorpora backend, migración Alembic, RBAC, SecurityEvent, interfaces USER/ADMIN, navegación contextual, pruebas SQLite/PostgreSQL y validación de browser. La configuración inicial se crea como DRAFT y requiere publicación explícita.",
        ],
        [
            "Sin modelo Candidate.",
            "Sin ranking global persistido.",
            "Sin autorización de inversión o ejecución.",
        ],
    ),
    (
        "2. ADR Gate 07E",
        [
            "Se aceptó ADR 45: Gate 07E consume READY_FOR_PORTFOLIO_PLANNING y produce READY_FOR_PORTFOLIO_ANALYSIS. La decisión conserva la identidad Project, usa membership N:M para el contexto y calcula la posición por portafolio.",
            "El ADR documenta los controles de tenant, configuración versionada, scoring determinístico, historial inmutable, migración reversible y límites explícitos frente a Gate 07F.",
        ],
        ["Archivo: docs/45-adr-gate07e-portfolio-evaluation-prioritization.md"],
    ),
    (
        "3. Baseline and architecture assessment",
        [
            "La auditoría confirmó que Gate 07D ya era propietario del Project Workspace, la asociación PortfolioProjectMembership y el planning-entry hash. AdminConfiguration ya resolvía configuración gobernada; RBAC, SecurityEvent y Workspace Context ya eran capacidades transversales.",
            "La implementación se integró como módulo backend independiente registrado en el router v1 y como feature frontend consumida desde App y WorkspaceOperationalPage.",
        ],
        [
            "Entrada: Gate 07D + Gate 05B/05C.",
            "Configuración: AdminConfiguration.",
            "Salida: contrato de handoff para Gate 07F.",
        ],
    ),
    (
        "4. Governance Model eligibility decision",
        [
            "La elegibilidad se restringe a CAPITAL_OWNER con planning readiness válido y membership activa en el portafolio seleccionado. Esta regla separa la evaluación de inversión estratégica de las rutas operativas de creación de Project.",
        ],
        [
            "CAPITAL_OWNER: elegible cuando cumple Gate 07D.",
            "CONTRACTOR_DELIVERY: bloqueado.",
            "DIRECT_INTERNAL: bloqueado.",
            "Legacy/unclassified: bloqueado.",
        ],
    ),
    (
        "5. REUSE BEFORE CREATE findings",
        [
            "Se reutilizaron Project Workspace, PortfolioProjectMembership, AdminConfiguration, User/Role/Permission, SecurityEvent, el contexto operacional y los snapshots de Gate 07A–07D. No se duplicaron identidad, creador de Project, árbol empresarial, publicación de configuraciones ni auditoría.",
        ],
        [
            "Nuevo únicamente: evaluación versionada y servicios/endpoints Gate 07E.",
            "Derivado: ranking y readiness; no se persisten como fuentes paralelas.",
        ],
    ),
    (
        "6. Evaluation domain model",
        [
            "PortfolioProjectEvaluation referencia tenant, portfolio_workspace_id, project_workspace_id y portfolio_membership_id. Incluye versión, estado, matriz y snapshots/hashes, ratings, componentes, score normalizado, alineación, riesgo, comentarios, evaluador, idempotencia y revisión optimista.",
        ],
        [
            "Cardinalidad: múltiples versiones por Project dentro de cada Portfolio.",
            "Unicidad: tenant + portfolio + project + evaluation_version.",
            "Tenant scoping obligatorio.",
        ],
    ),
    (
        "7. Evaluation lifecycle",
        [
            "El ciclo implementado es DRAFT → IN_PROGRESS → COMPLETED; una reevaluación crea una versión nueva y marca la anterior como SUPERSEDED. VOIDED queda reservado para anulación gobernada futura sin alterar resultados completados.",
            "La inmutabilidad se aplica en el modelo: un registro COMPLETED solo admite su transición controlada a SUPERSEDED, conservando score, evidencia y hashes.",
        ],
        [
            "Start valida elegibilidad y captura snapshot.",
            "Update exige If-Match.",
            "Complete valida criterios/evidencia y calcula score.",
            "Reevaluate crea historia, no sobrescribe.",
        ],
    ),
    (
        "8. Matrix configuration and inheritance",
        [
            "La matriz se gobierna mediante AdminConfiguration y se resuelve por la publicación más cercana: Portfolio, Business Unit, Enterprise y Tenant. El starter P&P se crea exclusivamente como DRAFT.",
            "La publicación exige criterios válidos, pesos consistentes, escala 1–5, fuentes requeridas y revisión optimista. Preview no persiste resultados.",
        ],
        [
            "DRAFT no gobierna evaluaciones productivas.",
            "No existe auto-publicación.",
            "Clonado y nueva revisión conservan trazabilidad.",
        ],
    ),
    (
        "9. Scoring scale and calculation",
        [
            "Cada criterio usa rating 1–5. El componente normalizado es ((rating − 1) / 4) × peso; el score total es la suma de componentes. La matriz default suma 100 puntos y permanece configurable.",
        ],
        [
            "Alineación estratégica 25.",
            "Valor económico 20.",
            "Beneficios 15.",
            "Riesgo 15.",
            "Urgencia 10.",
            "Capacidad 10.",
            "Dependencias 5.",
        ],
    ),
    (
        "10. Evaluation snapshot, history and reevaluation",
        [
            "El snapshot fija Project, governance model, project type, Portfolio, membership, planning hash, linaje Gate/Proposal/Idea, objetivos, Proposal Score, ROM, fechas, beneficios, riesgos y la revisión/hash de matriz.",
            "La reevaluación usa una nueva versión para evitar reinterpretar evidencia histórica con una configuración posterior.",
        ],
        [
            "Source snapshot hash detecta drift.",
            "Planning entry hash protege la entrada Gate 07D.",
            "Matrix hash fija el método aplicado.",
        ],
    ),
    (
        "11. Portfolio contextual ranking",
        [
            "El ranking se calcula únicamente dentro del Portfolio solicitado. Orden: score total descendente, alineación estratégica descendente, riesgo ascendente, finalización planificada más temprana y Project Number.",
            "No existe manual override: la posición es reproducible a partir de evaluaciones completadas vigentes y memberships activas.",
        ],
        [
            "Ranking determinístico.",
            "Sin estado global.",
            "Eventos de cálculo auditables.",
        ],
    ),
    (
        "12. Multi-Portfolio behavior",
        [
            "Un mismo Project puede tener memberships y evaluaciones independientes en varios Portfolios. La versión, score y posición se resuelven por contexto; una reevaluación en un Portfolio no altera el otro.",
        ],
        [
            "Misma identidad Project.",
            "Distinto matrix snapshot permitido.",
            "Distinta posición contextual esperada.",
        ],
    ),
    (
        "13. Prioritization Matrix",
        [
            "USER MODE muestra posición, Project Number, nombre, score, alineación, riesgo, fecha de finalización, estado de evaluación y bucket derivado. La tabla consume el endpoint contextual y no materializa un segundo ranking.",
        ],
        [
            "Preview permite escenarios no persistentes.",
            "Membership inactiva se excluye del ranking actual.",
            "La historia de evaluación permanece consultable.",
        ],
    ),
    (
        "14. Portfolio-level readiness",
        [
            "El readiness compara Projects elegibles con evaluaciones vigentes completadas y bloquea la salida si existe cualquier evaluación en progreso o Project bloqueado. Incluye conteos, blockers, snapshot y readiness hash.",
        ],
        [
            "Sin Projects elegibles: bloqueado.",
            "Con Project bloqueado: bloqueado.",
            "Todos elegibles completados y sin blockers: listo.",
        ],
    ),
    (
        "15. READY_FOR_PORTFOLIO_ANALYSIS contract",
        [
            "El contrato final expone can_enter_portfolio_analysis, final_output, conteos, blockers, portfolio_id, snapshot y readiness_hash. La salida positiva exclusiva es READY_FOR_PORTFOLIO_ANALYSIS.",
        ],
        [
            "No equivale a aprobación de inversión.",
            "No inicializa ni activa Project.",
            "No autoriza ejecución.",
        ],
    ),
    (
        "16. Gate 07F contract",
        [
            "Gate 07F puede consumir el readiness hash, snapshot y matriz contextual como evidencia de entrada. Gate 07E no implementa análisis de portafolio, optimización, presupuesto, recursos ni decisión FID.",
        ],
        [
            "Handoff legible y estable.",
            "Ownership de Gate 07F preservado.",
            "No se anticipan estados posteriores.",
        ],
    ),
    (
        "17. FEL parallel-branch protection",
        [
            "El scoring de portafolio no sustituye FEL/PDRI ni sus evidencias. Fechas, ROM, beneficios y riesgos se capturan como inputs históricos, pero Gate 07E no calcula preparación técnica de definición ni autoriza inversión.",
        ],
        [
            "Portfolio score y FEL readiness son ramas paralelas.",
            "No se usa el score como FID automático.",
        ],
    ),
    (
        "18. Contractor, Direct and Legacy exclusion evidence",
        [
            "Las pruebas backend construyen Projects CONTRACTOR_DELIVERY y legacy/unclassified y verifican que Start Evaluation sea rechazado. DIRECT_INTERNAL usa la misma regla estricta de gobernanza.",
        ],
        [
            "SecurityEvent registra los intentos bloqueados.",
            "La exclusión ocurre antes de crear evaluación.",
        ],
    ),
    (
        "19. USER MODE",
        [
            "Portfolio Evaluation ofrece colas Pending/In Progress/Completed/Blocked, detalle de fuentes, criterios 1–5, evidencia, comentarios, completar y reevaluar. Prioritization Matrix muestra ranking y readiness contextual.",
            "En Project Workspace aparece Portfolio Evaluations para consultar historia por membership activa.",
        ],
        [
            "Los vacíos son estados controlados, no datos ficticios.",
            "Las acciones dependen de permisos y elegibilidad.",
        ],
    ),
    (
        "20. ADMIN MODE",
        [
            "Portfolio Evaluation & Prioritization incorpora once áreas: General, Evaluation Matrix, Scoring Scale, Weights, Required Evidence, Applicability, Ranking Rules, Coverage/Readiness, Inheritance, Permissions y Preview. La matriz starter se identifica claramente como DRAFT.",
        ],
        [
            "Clonar, editar, preview y publicar usan el motor común.",
            "No existe publicación implícita.",
        ],
    ),
    (
        "21. Workspace navigation",
        [
            "El árbol USER añade Portfolio Evaluation y Prioritization Matrix dentro de Portfolio Manager. El contexto de Project añade Portfolio Evaluations cuando corresponde. ADMIN añade Portfolio Evaluation & Prioritization en Enterprise Strategy Manager.",
        ],
        [
            "Navegación existente conservada.",
            "Rutas renderizadas desde App y WorkspaceOperationalPage.",
        ],
    ),
    (
        "22. Allowed actions",
        [
            "Las acciones habilitadas son consultar colas, iniciar, editar, completar y reevaluar; consultar y previsualizar priorización/readiness; listar, previsualizar, clonar, actualizar y publicar configuraciones.",
        ],
        [
            "No hay override manual de rango.",
            "No hay FID, budget allocation, activation ni execute.",
        ],
    ),
    (
        "23. Permissions and roles",
        [
            "Se agregaron ocho permisos: portfolio_evaluation.read/create/edit/complete/reevaluate, portfolio_prioritization.read y portfolio_evaluation.admin.configure/publish.",
            "Roles Gate 07E: portfolio_evaluator, portfolio_prioritization_viewer y portfolio_evaluation_configuration_admin, además de organization_admin con autoridad gobernada.",
        ],
        [
            "RBAC se valida antes del servicio.",
            "Tenant scoping se conserva en cada consulta.",
        ],
    ),
    (
        "24. ETag and idempotency",
        [
            "Update y publicación exigen If-Match/revision_version. Start, Complete y Reevaluate almacenan claves de idempotencia separadas; los reintentos devuelven el mismo resultado y la concurrencia se resuelve por constraints y recuperación transaccional.",
        ],
        [
            "Stale version produce conflicto explícito.",
            "Los completados no se sobrescriben.",
        ],
    ),
    (
        "25. SecurityEvent and audit",
        [
            "Los eventos cubren inicio, actualización, finalización, reevaluación, bloqueo por gobernanza, cálculo de priorización, readiness y gobierno de configuración. Cada evento conserva tenant, actor, target y hashes pertinentes.",
        ],
        [
            "Linaje hacia Gate 07D/Gate 07C/Proposal/Idea.",
            "Hash de matriz y source snapshot en auditoría.",
        ],
    ),
    (
        "26. APIs",
        [
            "El router Gate 07E expone contratos USER para colas/evaluaciones/priorización y contratos ADMIN para configuración. Los payloads Pydantic validan ratings, evidencia, idempotencia y versiones.",
        ],
        [
            "GET/POST /portfolios/{portfolio}/.../evaluations",
            "GET/PUT/POST /portfolio-evaluations/{id}[/complete|/reevaluate]",
            "GET/POST /portfolios/{portfolio}/prioritization[/readiness|/preview]",
            "GET/POST/PUT /portfolio-evaluation/admin/configurations [.../clone|.../publish].",
        ],
    ),
    (
        "27. Persistence",
        [
            "La tabla portfolio_project_evaluations persiste solamente evaluaciones y snapshots históricos. La priorización y readiness se calculan al consultar para evitar divergencia entre score, membership y posición.",
        ],
        [
            "FKs a Portfolio, Project, membership y configuración.",
            "Índices tenant/contexto/estado.",
            "Constraints de versión e idempotencia.",
        ],
    ),
    (
        "28. Migration",
        [
            "Alembic 20260820_0045 agrega la tabla Gate 07E, FKs, índices y constraints sin mutar datos empresariales. Se validó 0044 → 0045 → 0044 → 0045 en PostgreSQL efímero.",
        ],
        [
            "Aditiva.",
            "Reversible.",
            "Tenant scoped.",
            "Sin datos productivos de evaluación sembrados.",
        ],
    ),
    (
        "29. Concurrency",
        [
            "La prueba PostgreSQL dispara inicios y finalizaciones concurrentes con la misma clave. El servicio recupera la evaluación ganadora tras la restricción única y conserva una sola versión vigente.",
        ],
        [
            "Inicio concurrente idempotente.",
            "Complete concurrente idempotente.",
            "Sin doble versión actual.",
        ],
    ),
    (
        "30. PostgreSQL E2E",
        [
            "La puerta efímera aplicó migraciones y ejecutó 69 regresiones Gate 07D/contexto más 9 escenarios Gate 07E. Se verificaron CAPITAL_OWNER, exclusiones, tres Projects, membership inactiva, multi-Portfolio, historia y concurrencia.",
        ],
        [
            "69 passed + 9 passed.",
            "Upgrade/downgrade/re-upgrade PASS.",
            "Schema validator PASS.",
        ],
    ),
    (
        "31. Frontend tests",
        [
            "Vitest aprobó 36 archivos y 197 pruebas. Tres pruebas específicas validan colas sin Candidate, ranking/readiness contextual y matriz starter DRAFT. La regresión AppFlow valida los nombres Portfolio Evaluation y Prioritization Matrix.",
        ],
        [
            "TypeScript/Vite build PASS.",
            "Prettier PASS.",
            "ESLint: 0 errores; 8 warnings heredados del visor IFC.",
        ],
    ),
    (
        "32. Browser USER E2E",
        [
            "En localhost se verificó la carga de /app, USER MODE, el despliegue de Portfolio Manager y la disponibilidad de Portfolio Evaluation y Prioritization Matrix. Los estados vacíos se presentan de forma controlada porque el baseline no contiene memberships productivas.",
        ],
        [
            "Frontend: http://127.0.0.1:5173/app",
            "API health: http://127.0.0.1:8000/api/v1/health",
        ],
    ),
    (
        "33. Browser ADMIN E2E",
        [
            "Se verificó el cambio a ADMIN MODE y la entrada Portfolio Evaluation & Prioritization. La vista muestra las once áreas de gobierno y mantiene el starter matrix como DRAFT sin publicación automática.",
        ],
        [
            "ADMIN navigation PASS.",
            "Configuration DRAFT visible.",
            "No evaluación/ranking productivo creado.",
        ],
    ),
    (
        "34. Regression",
        [
            "Las regresiones cubren Gate 07A, 07B, 07C, Gate 07D/hardening, Project creation/governance, Gate 05C, Workspace Context, Enterprise Explorer y módulos BIM/control existentes.",
        ],
        [
            "Backend integral PASS.",
            "Frontend integral PASS.",
            "PostgreSQL Gate 07D + Gate 07E PASS.",
        ],
    ),
    (
        "35. Baseline before and after",
        [
            "Antes de 0045: 14 Enterprise Workspaces, 1 Project Workspace, 0 ProjectCreationRequests, 0 PortfolioProjectMemberships, 68 AdminConfigurations y hash de defaults del Project ID 14 eea3648c6908a267d1e0651828ff4459.",
            "Después del despliegue: los conteos empresariales y el hash se conservan; existe la tabla Gate 07E con 0 evaluaciones productivas. AdminConfiguration pasó de 68 a 70 por la materialización gobernada del default publicado Gate 07D y el starter Gate 07E DRAFT; ninguna matriz Gate 07E quedó publicada.",
        ],
        [
            "Datos de negocio preservados.",
            "Sin rankings persistidos.",
            "Sin memberships artificiales.",
        ],
    ),
    (
        "36. Risks and technical debt",
        [
            "Los Portfolios actuales no pueden demostrar readiness operativo hasta que existan memberships CAPITAL_OWNER y una matriz publicada. Esto es un estado de datos esperado, no una falla del release.",
            "Persisten ocho warnings de lint en BimIfcModelViewer y advertencias act(...) en pruebas React; son deuda heredada fuera del alcance Gate 07E. Gate 07F y la metodología FEL/PDRI continúan pendientes por diseño.",
        ],
        [
            "Publicar matrices requiere gobernanza real.",
            "Monitorear volumen de snapshots JSON.",
            "Mantener pruebas de concurrencia en PostgreSQL CI.",
        ],
    ),
    (
        "37. Updated technical report and configured artifacts",
        [
            "La entrega incluye módulo backend, migración, esquema/servicio/router, permisos, navegación, feature frontend, estilos, pruebas unitarias, E2E PostgreSQL, CI, Docker gate y ADR. Este documento consolida arquitectura, controles, validaciones y estado de despliegue.",
        ],
        [
            "backend/app/modules/portfolio_evaluation/",
            "backend/alembic/versions/20260820_0045_gate07e_portfolio_evaluation.py",
            "frontend/src/features/portfolio-evaluation/",
            "backend/tests/test_portfolio_evaluation_gate07e.py",
            "backend/tests/postgres/test_portfolio_evaluation_postgres.py",
            "docker-compose.gate07e.yml y .github/workflows/ci.yml",
        ],
    ),
    (
        "38. Final exclusive state",
        [
            "READY_FOR_PORTFOLIO_ANALYSIS",
            "La implementación y su contrato están listos para que Gate 07F consuma una priorización contextual gobernada. El baseline productivo permanece sin evaluaciones porque no existen memberships elegibles; no se fabricaron datos para simular readiness.",
        ],
        [
            "STOP: no iniciar Gate 07F, FID, presupuesto, recursos, inicialización, activación ni ejecución."
        ],
    ),
    (
        "39. Gate 07E-H closeout scope",
        [
            "El closeout validó exclusivamente la liberación de Portfolio Evaluation & Prioritization ya implementada. No agregó Gate 07F, optimización, presupuesto, selección de recursos, FID, activación ni ejecución.",
            "La aceptación combinó pruebas de browser contra localhost, controles API, PostgreSQL efímero, regresión integral y comparación determinística de la base persistente antes y después.",
        ],
        [
            "No new scope.",
            "No Candidate global.",
            "No rank override.",
            "No datos productivos fabricados.",
        ],
    ),
    (
        "40. Browser environment and deployment",
        [
            "El entorno validado usó frontend Vite en http://127.0.0.1:5173/app, API en http://127.0.0.1:8000, PostgreSQL 16 y Redis en Docker Compose. API, DB y Redis permanecieron healthy; worker y beat estuvieron activos.",
            "La revisión visual final confirmó USER/ADMIN MODE, navegación jerárquica, panel de configuración y el grupo accesible Portfolio Evaluation configuration actions.",
        ],
        [
            "Alembic head: 20260820_0045.",
            "Localhost disponible al cierre.",
            "Sin errores visuales críticos.",
        ],
    ),
    (
        "41. Disposable E2E data design",
        [
            "Se implementó gate07e_browser_fixture.py con baseline, seed y teardown. Cada repetición usa Portfolios, Projects, memberships y matrices propios bajo un run_id controlado; el teardown elimina sólo ese namespace y comprueba que no queden workspaces ni configuraciones residuales.",
            "La semilla cubre tres Projects principales, Portfolio secundario, exclusiones de gobernanza, planning bloqueado, Portfolio sin matriz y Project sin membership.",
        ],
        [
            "Tres copias aisladas.",
            "IDs y códigos run-scoped.",
            "Teardown exacto y repetible.",
        ],
    ),
    (
        "42. USER end-to-end lifecycle",
        [
            "Cada pasada inició tres evaluaciones desde cards USER, completó los criterios 1–5 con evidencia y comentarios, guardó borradores con nuevo ETag y congeló snapshots COMPLETED inmutables.",
            "La UI respetó allowed_actions calculadas por backend: Start sólo se mostró para proyectos elegibles; completados no expusieron edición ni completado adicional.",
        ],
        [
            "Start 201.",
            "Save con If-Match.",
            "Complete PASS.",
            "Hashes de matriz y fuente presentes.",
        ],
    ),
    (
        "43. Ranking and readiness evidence",
        [
            "Los Projects A, B y C produjeron scores controlados. A quedó primero; B y C empataron en score y se ordenaron por alineación estratégica. El ranking mostró Proposal Score, objetivos, ROM, estado y fecha desde el snapshot de origen.",
            "La cobertura alcanzó 100,00% y la salida contextual fue READY_FOR_PORTFOLIO_ANALYSIS. Los desempates quedaron en score, alineación, riesgo, finalización planificada y Project Number.",
        ],
        ["Ranking A > B > C.", "Readiness 100%.", "Sin ranking global persistido."],
    ),
    (
        "44. Reevaluation and immutable history",
        [
            "La reevaluación de Project A creó versión 2 DRAFT, marcó versión 1 como SUPERSEDED y preservó matrix_hash y source_snapshot_hash históricos. La nueva versión se completó sin sobrescribir la evidencia anterior.",
        ],
        ["v1 SUPERSEDED.", "v2 COMPLETED.", "Historia consultable e inmutable."],
    ),
    (
        "45. Multi-Portfolio and inactive membership",
        [
            "El mismo Project A fue evaluado con una matriz distinta en el Portfolio secundario y obtuvo un score contextual independiente. Al retirar la membership secundaria, desapareció del ranking actual sin eliminar la evaluación COMPLETED histórica.",
        ],
        [
            "Contexto independiente por Portfolio.",
            "Membership inactiva excluida.",
            "Histórico preservado.",
        ],
    ),
    (
        "46. Governance exclusions and negative paths",
        [
            "UI y API rechazaron CONTRACTOR_DELIVERY, DIRECT_INTERNAL y LEGACY con blockers explícitos. También se validaron planning no listo, ausencia de matriz publicada y ausencia de membership activa.",
            "Las pruebas negativas rechazaron criterios incompletos, rating fuera de rango, evidencia obligatoria vacía y versiones ETag obsoletas sin crear estados parciales.",
        ],
        [
            "Errores 422 gobernados.",
            "ETag stale: 412.",
            "Sin evaluación creada en exclusiones.",
        ],
    ),
    (
        "47. ADMIN lifecycle and non-persistent preview",
        [
            "ADMIN seleccionó una matriz publicada, creó un clon DRAFT con revisión incremental, modificó pesos manteniendo total 100, previsualizó contenido no guardado y verificó que el preview no generara evaluaciones ni resultados persistentes.",
            "Después se guardó con nueva versión y hash, se publicó explícitamente y se comprobó la resolución efectiva en el Portfolio objetivo.",
        ],
        [
            "Clone PASS.",
            "Unsaved Preview PASS.",
            "Save PASS.",
            "Explicit Publish PASS.",
        ],
    ),
    (
        "48. ETag, inheritance and source integrity",
        [
            "Una actualización ADMIN con ETag anterior devolvió 412 ETAG_MISMATCH. El preview posterior a publicación resolvió la configuración nueva, el path de herencia incluyó el Portfolio y la publicación original conservó su content_hash.",
        ],
        [
            "Optimistic concurrency PASS.",
            "Inheritance PASS.",
            "Configuración fuente inmutable.",
        ],
    ),
    (
        "49. Browser stability and observability",
        [
            "Playwright ejecutó el flujo completo tres veces consecutivas con un worker y copias aisladas. Las tres terminaron FLOW_PASS en 3,1 minutos acumulados. El cierre exigió cero page errors, cero mensajes console warning/error del dominio Gate 07E y cero respuestas inesperadas >= 400.",
        ],
        ["3/3 PASS.", "Console clean.", "Network clean para operaciones positivas."],
    ),
    (
        "50. Defects found and minimal fixes",
        [
            "El hardening corrigió selectores E2E ambiguos, el origen LEGACY del fixture, el límite integral de tiempo, la selección estable de configuración frente a respuestas iniciales obsoletas y la semántica accesible del grupo de acciones ADMIN.",
            "La corrección funcional de UI usa una secuencia monotónica de carga para ignorar respuestas stale. El grupo Clone/Preview/Save/Publish recibió un aria-label propio para distinguirlo de la pestaña Preview.",
        ],
        [
            "Sin migración adicional.",
            "Sin cambio de ownership.",
            "Sin expansión funcional.",
        ],
    ),
    (
        "51. Continuous integration closeout",
        [
            "Pilot Readiness CI prepara baseline, aplica Alembic a head, valida salud, genera tres copias efímeras, ejecuta el closeout con repeat-each=3 y workers=1, ejecuta teardown siempre y compara la línea base antes/después.",
        ],
        [
            "Fixture incluido en imagen API.",
            "Playwright closeout en CI.",
            "Teardown protegido con always().",
        ],
    ),
    (
        "52. Full regression results",
        [
            "Backend aprobó 486 pruebas, omitió 13 casos dependientes del entorno y alcanzó 85,35% de cobertura. Frontend aprobó 36 archivos y 197 pruebas. TypeScript/Vite build, Ruff, ESLint focalizado, Prettier y diff check terminaron sin errores.",
            "La puerta PostgreSQL aprobó 69 regresiones y 9 pruebas Gate 07E; validó 0044, upgrade 0045, downgrade 0044 y re-upgrade 0045.",
        ],
        [
            "Backend: PASS.",
            "Frontend: PASS.",
            "PostgreSQL: PASS.",
            "Build/quality: PASS.",
        ],
    ),
    (
        "53. Acceptance matrix",
        [
            "Todos los criterios obligatorios del PROMPT 07E-H quedaron satisfechos. La matriz siguiente registra la evidencia consolidada y la decisión de cada control.",
        ],
        [],
    ),
    (
        "54. Persistent baseline before and after closeout",
        [
            "El baseline persistente fue idéntico antes y después: Alembic 0045; 14 Enterprise Workspaces; 1 Project Workspace; 0 ProjectCreationRequests; 0 PortfolioProjectMemberships; 0 evaluaciones Gate 07E; 70 AdminConfigurations; starter Gate 07E 1 DRAFT / 0 PUBLISHED; MD5 del Project 14 eea3648c6908a267d1e0651828ff4459.",
            "El teardown oficial eliminó 39 workspaces efímeros, 27 memberships, 15 evaluaciones y 9 configuraciones, sin residuos del run_id. Se preservaron los datos empresariales, no se sembraron evaluaciones productivas y no quedó ninguna membership artificial.",
        ],
        [],
    ),
    (
        "55. Release decision and STOP condition",
        [
            "GATE07E_RELEASE_VALIDATED",
            "La decisión certifica que Gate 07E está validado para liberación con su contrato READY_FOR_PORTFOLIO_ANALYSIS. No constituye aprobación de inversión ni autoriza iniciar Gate 07F o capacidades posteriores.",
        ],
        [
            "STOP: no Gate 07F, FID, presupuesto, recursos, inicialización, activación ni ejecución."
        ],
    ),
]

# Anchor the two sections that already begin on a fresh rendered page.  Explicit
# breaks prevent Microsoft Word from carrying pagination state from the preceding
# list into the heading (which can otherwise suppress bullets or shift the first
# glyph outside the page margin during PDF export).
page_break_sections: set[int] = {31, 37, 39, 55}
for index, (title_text, paragraphs, items) in enumerate(sections, start=1):
    add_heading(doc, title_text, level=1, page_break=index in page_break_sections)
    for paragraph in paragraphs:
        add_body(doc, paragraph)
    for item in items:
        add_list_item(doc, item, bullet_num)
    if index == 7:
        add_table(
            doc,
            ["Estado", "Mutabilidad", "Salida"],
            [
                [
                    "DRAFT / IN_PROGRESS",
                    "Editable con If-Match",
                    "Evaluación no rankeable",
                ],
                ["COMPLETED", "Inmutable", "Score vigente y rankeable"],
                ["SUPERSEDED", "Histórico", "Excluido del ranking actual"],
                ["VOIDED", "Histórico", "Excluido"],
            ],
            [2100, 3300, 3960],
        )
    if index == 9:
        add_table(
            doc,
            ["Criterio starter", "Peso", "Tratamiento"],
            [
                ["Strategic alignment", "25", "Mayor rating mejora score"],
                ["Economic value", "20", "Mayor rating mejora score"],
                ["Benefits", "15", "Mayor rating mejora score"],
                [
                    "Risk",
                    "15",
                    "Rating normalizado; riesgo crudo conserva desempate ascendente",
                ],
                ["Urgency", "10", "Mayor rating mejora score"],
                ["Capacity", "10", "Mayor rating mejora score"],
                ["Dependencies", "5", "Mayor rating mejora score"],
            ],
            [3800, 1100, 4460],
        )
    if index == 35:
        add_table(
            doc,
            ["Métrica persistente", "Antes 0045", "Después 0045"],
            [
                ["Enterprise Workspaces", "14", "14"],
                ["Project Workspaces", "1", "1"],
                ["Project creation requests", "0", "0"],
                ["Portfolio memberships", "0", "0"],
                ["Admin configurations", "68", "70 (Gate 07E: 1 DRAFT / 0 PUBLISHED)"],
                ["Gate 07E evaluations", "Tabla no existente", "0"],
                ["Project ID 14 defaults MD5", "eea3648c...4459", "eea3648c...4459"],
            ],
            [4200, 2580, 2580],
        )
    if index == 38:
        add_callout(
            doc,
            "Estado final",
            "READY_FOR_PORTFOLIO_ANALYSIS — PASS. Este estado certifica el release Gate 07E y su contrato de handoff; no representa una decisión de inversión sobre el portafolio persistente.",
            fill=POSITIVE,
        )
    if index == 53:
        add_table(
            doc,
            ["Criterio de aceptación", "Resultado", "Evidencia"],
            [
                [
                    "Baseline previo",
                    "PASS",
                    "Conteos y MD5 registrados antes del seed.",
                ],
                [
                    "Tres Projects y ranking",
                    "PASS",
                    "A > B > C; desempate determinístico.",
                ],
                [
                    "Readiness",
                    "PASS",
                    "Cobertura 100% y salida READY_FOR_PORTFOLIO_ANALYSIS.",
                ],
                [
                    "Reevaluación e historia",
                    "PASS",
                    "v1 SUPERSEDED; v2 COMPLETED; hashes preservados.",
                ],
                [
                    "Multi-Portfolio",
                    "PASS",
                    "Score contextual independiente y membership inactiva excluida.",
                ],
                [
                    "Exclusiones",
                    "PASS",
                    "CONTRACTOR, DIRECT y LEGACY bloqueados en UI/API.",
                ],
                ["Negativos", "PASS", "422 de validación y 412 por ETag stale."],
                [
                    "ADMIN lifecycle",
                    "PASS",
                    "Clone, unsaved preview, save y publish explícito.",
                ],
                [
                    "Herencia e integridad",
                    "PASS",
                    "Fuente efectiva, path y hashes verificados.",
                ],
                [
                    "Estabilidad browser",
                    "PASS",
                    "Tres pasadas consecutivas; 3/3 FLOW_PASS.",
                ],
                ["Teardown", "PASS", "Cero residuos y baseline idéntico."],
                ["Regresión", "PASS", "486 backend; 197 frontend; 78 PostgreSQL."],
            ],
            [3700, 1300, 4360],
        )
    if index == 55:
        add_callout(
            doc,
            "Estado exclusivo de release",
            "GATE07E_RELEASE_VALIDATED — PASS. Gate 07E queda validado para liberación; la salida funcional continúa siendo READY_FOR_PORTFOLIO_ANALYSIS.",
            fill=POSITIVE,
        )

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
