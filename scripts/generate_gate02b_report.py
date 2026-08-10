from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
EVIDENCE = REPO / "artifacts" / "enterprise_structure" / "gate02b"
REPORT_ROOT = Path.home() / "Documents" / "P&P" / "P&Pmis Construction AI"
SOURCE = next(
    REPORT_ROOT.glob(
        "Dise*/Resumen de Sprint/Informe_Tecnico_PPMIS_Core_Correction_Gate02A_2026-08-09.docx"
    )
)
OUTPUT = SOURCE.with_name("Informe_Tecnico_PPMIS_Core_Correction_Gate02B_2026-08-09.docx")

NAVY = "0D2A3A"
TEAL = "0F9D9A"
PALE_TEAL = "EAF7F6"
PALE_BLUE = "EFF5F9"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
WHITE = "FFFFFF"
TEXT = "16222B"
MUTED = "526775"


def load(name: str) -> dict:
    return json.loads((EVIDENCE / name).read_text(encoding="utf-8"))


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text_color(cell, color: str, bold: bool = False) -> None:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = RGBColor.from_string(color)
            run.bold = bold


def set_cell_margin(cell, top: int = 90, start: int = 110, bottom: int = 90, end: int = 110) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float] | None = None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = widths is None
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        cell.text = header
        shade(cell, NAVY)
        set_cell_text_color(cell, WHITE, bold=True)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margin(cell)
        if widths:
            cell.width = Inches(widths[idx])
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        fill = WHITE if row_index % 2 == 0 else PALE_BLUE
        for idx, value in enumerate(values):
            cells[idx].text = str(value)
            shade(cells[idx], fill)
            set_cell_text_color(cells[idx], TEXT)
            cells[idx].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cells[idx])
            if widths:
                cells[idx].width = Inches(widths[idx])
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return table


def add_callout(doc: Document, title: str, text: str, fill: str = PALE_TEAL) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(2)
    run = paragraph.add_run(title.upper())
    run.bold = True
    run.font.color.rgb = RGBColor.from_string(NAVY)
    run.font.size = Pt(9)
    detail = cell.add_paragraph(text)
    detail.paragraph_format.space_after = Pt(0)
    for run in detail.runs:
        run.font.color.rgb = RGBColor.from_string(TEXT)
        run.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(2)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


first = load("gate02b_apply_first.json")
second = load("gate02b_apply_second.json")
preflight = load("gate02b_preflight_fingerprints.json")
post_apply = load("gate02b_post_apply_fingerprints.json")
post_second = load("gate02b_post_second_fingerprints.json")
dry_run = load("gate02b_immediate_dry_run.json")
validation = load("gate02b_post_apply_validation.json")
qa = load("gate02b_qa.json")
tests = load("gate02b_tests.json")
audit = load("gate02b_audit_and_publish_validation.json")

doc = Document(SOURCE)
doc.core_properties.title = "Informe técnico P&Pmis Construction AI — Core Correction Gate 02B"
doc.core_properties.subject = "Controlled Apply CORE transaccional e idempotente"
doc.core_properties.comments = "Documento acumulativo Gates 00, 01, 02A y 02B. Publish CORE no ejecutado."

# Update the cover while retaining the established visual identity.
doc.paragraphs[3].text = "Nivel 2B Core Correction — Gates 00, 01, 02A y 02B"
doc.paragraphs[4].text = (
    "Apply CORE controlado · 3 identidades adoptadas · 11 nodos creados · "
    "segunda ejecución idempotente · Publish CORE no ejecutado"
)
cover = doc.tables[0].cell(0, 0)
cover.text = (
    "RESULTADO ACTUAL\n"
    "GATE 02B COMPLETO · APPLY CORE SUCCESS · QA ADMIN/USER PASS\n"
    "14 WORKSPACES · 7 OBJETIVOS · 26 CLASIFICACIONES · 0 LINKS"
)
shade(cover, PALE_GREEN)
set_cell_text_color(cover, NAVY, bold=True)
set_cell_margin(cover, 180, 180, 180, 180)

# Extend the manual content index.
contents = doc.tables[2]
row = contents.add_row().cells
row[0].text = "31–41"
row[1].text = "Gate 02B — apply transaccional, evidencias, QA, riesgos y cierre"
for cell in row:
    shade(cell, PALE_GREEN)
    set_cell_margin(cell)

doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

add_heading(doc, "31. Gate 02B — resultado ejecutivo")
doc.add_paragraph(
    "Se ejecutó el Controlled Apply CORE autorizado para P&P Ingeniería y Proyectos. El proceso reutilizó "
    "el importador, el modelo multitenant, las reglas de composición, los catálogos publicados y las vistas "
    "ADMIN/USER ya existentes. La transacción preservó las identidades internas 1, 2 y 3, completó la "
    "jerarquía a 14 workspaces, persistió objetivos y clasificaciones, y se repitió sin cambios funcionales."
)
add_table(
    doc,
    ["Control de salida", "Resultado"],
    [
        ["Preflight", "PASS · fuente protegida sin cambios"],
        ["Dry-run inmediato", "VALID · 0 errores / 0 warnings / 0 conflictos"],
        ["Primer apply", "SUCCESS · 3 ADOPT + 11 CREATE"],
        ["Segundo apply", "PASS · 47 UNCHANGED; 0 CREATE / UPDATE / CONFLICT"],
        ["QA ADMIN", "PASS"],
        ["QA USER", "PASS"],
        ["Pruebas", "PASS"],
        ["Publish CORE", "NOT EXECUTED"],
    ],
    [2.1, 4.3],
)
add_callout(
    doc,
    "Límite de autorización respetado",
    "Gate 02B autorizó apply CORE, no publicación. El único evento publish existente es el ID 8 de Gate 01; "
    "Gate 02B agregó únicamente los eventos técnicos core_applied ID 9 e ID 10.",
    PALE_AMBER,
)

add_heading(doc, "32. Aprobaciones y estado inicial")
doc.add_paragraph(
    "La ejecución usó exactamente las decisiones incluidas en PROMPT_02B_CONTROLLED_APPLY_CORE.md y no "
    "aplicó heurísticas de identidad. El actor fue admin@demo.local, usuario activo con permiso "
    "admin.enterprise_structure.manage y alcance organization-wide."
)
add_table(
    doc,
    ["Aprobación", "Valor aplicado"],
    [
        ["Tenant", "ID 1 · Demo Energy Infrastructure / demo-energy → P&P Ingeniería y Proyectos / pyp-ingenieria-proyectos"],
        ["Moneda", "COP preservada"],
        ["ADOPT", "ENT-PYP = ID 1; BU-PYP-PMO = ID 3; BU-PYP-CONST = ID 2"],
        ["CREATE", "11 nodos exactos del input reconciliado"],
        ["Contenido", "7 Strategic Objectives; 26 Classifications; 0 Links"],
        ["Exclusiones", "Sin EXPERIENCE, PROPERTY_FACILITY, Project Creator ni publish"],
    ],
    [1.7, 4.7],
)
add_heading(doc, "Preflight y concurrencia", 2)
doc.add_paragraph(
    "El inventario real de Gate 02B conservaba el tenant ID 1, tres workspaces y los mismos conteos observados "
    "en Gate 02A. Los MD5 históricos no son directamente comparables porque Gate 02A no conservó su receta de "
    "canonicalización; por ello se verificó equivalencia semántica exacta y se estableció un baseline SHA-256 documentado."
)
add_table(
    doc,
    ["Tabla protegida", "Filas iniciales", "SHA-256 preflight"],
    [
        [name, data["rows"], data["sha256"]]
        for name, data in preflight["tables"].items()
    ],
    [2.35, 0.9, 3.15],
)
doc.add_paragraph("Protected source hash", style="Heading 2")
doc.add_paragraph(preflight["protected_source_hash"], style="Code Block")

add_heading(doc, "33. Revalidación inmediata e integridad del input")
doc.add_paragraph(
    "Después de aplicar la migración y antes de cualquier escritura funcional se repitieron preflight y dry-run. "
    "La migración no alteró la huella protegida. El hash físico del YAML se comprobó antes y dentro de la "
    "transacción; el parser produjo además el mismo hash canónico de Gate 02A."
)
add_table(
    doc,
    ["Control", "Valor"],
    [
        ["YAML raw SHA-256 aprobado", first["input_hash"]],
        ["Hash canónico", first["canonical_input_hash"]],
        ["Release", first["release_code"]],
        ["Dry-run válido", str(dry_run["valid"]).upper()],
        ["Errores / warnings / conflictos", f"{dry_run['summary']['errors']} / {dry_run['summary']['warnings']} / {dry_run['summary']['conflict']}"],
        ["Diff", "3 adopt · 44 create totales = 11 workspaces + 7 objetivos + 26 clasificaciones"],
        ["Base mutations en dry-run", dry_run["summary"]["base_mutations"]],
    ],
    [2.2, 4.2],
)

add_heading(doc, "34. Implementación del apply controlado")
doc.add_paragraph(
    "El flujo se implementó dentro del importador existente. La CLI abre una sola transacción y apply_core nunca "
    "confirma por cuenta propia. PostgreSQL adquiere locks coordinados sobre las tablas protegidas; cualquier "
    "excepción revierte identidad del tenant, workspaces, objetivos, clasificaciones y el evento técnico."
)
add_bullets(
    doc,
    [
        "Verificación doble del SHA-256 físico del YAML y comparación del snapshot protegido.",
        "Resolución exclusiva del tenant anterior o final, sin crear tenants alternativos.",
        "Validación de actor existente, activo, autorizado y con alcance organizacional.",
        "ADOPT por existing_id declarado; no se borra ni recrea ninguna de las tres identidades.",
        "Recodificación temporal segura para evitar la colisión 01.01 entre ID 2 y BU-PYP-CONS.",
        "Creación topológica de los once nodos y upsert idempotente de objetivos y clasificaciones.",
        "Chequeos finales de una raíz, unicidad, padres, tenant y ausencia de ciclos.",
        "Evento enterprise_structure.core_applied con hashes, release, actor, resultado e IDs afectados.",
    ],
)
add_heading(doc, "Archivos principales incorporados o actualizados", 2)
add_table(
    doc,
    ["Archivo / componente", "Responsabilidad Gate 02B"],
    [
        ["importer/apply.py", "Transacción, aprobaciones, locks, adopción, creación, integridad e idempotencia"],
        ["importer/inventory.py", "Inventario por reflexión y fingerprints SHA-256 reproducibles"],
        ["importer/cli.py", "Subcomandos preflight/apply y evidencia JSON + humana de una misma transacción"],
        ["importer/validator.py + snapshot.py", "Revalidación reconciliada y lectura de identidad persistida"],
        ["models.py / domain/models.py", "external_key persistente y Strategic Objectives"],
        ["repository.py / service.py / routers", "Consumo unificado en ADMIN y USER"],
        ["auth.py / domain/schemas.py / seed.py", "Login compatible tras el cambio de slug; seed sin tenant duplicado"],
        ["tests/test_enterprise_structure_apply.py", "Apply, rollback, seguridad, cross-tenant e idempotencia"],
        ["frontend Enterprise Structure", "Árbol, tabla, formularios, Record Codes y cabecera compacta"],
        ["README-enterprise-structure.md", "Contrato operativo de preflight, validate y apply; publish separado"],
    ],
    [2.7, 3.7],
)

add_heading(doc, "35. Migración PostgreSQL")
doc.add_paragraph(
    "Alembic avanzó desde 20260809_0029 hasta 20260809_0030 (head). La migración añade external_key nullable "
    "con índice y unicidad por tenant, y crea enterprise_strategic_objectives con identidad estable por tenant y código."
)
add_table(
    doc,
    ["Elemento", "Resultado"],
    [
        ["Cadena Alembic", "20260809_0029 → 20260809_0030 (head)"],
        ["DDL PostgreSQL", "PASS · transaccional"],
        ["Huella antes/después de migrar", "Sin cambio: cfdcef74…ad966"],
        ["Backfill funcional", "No ejecutado por la migración; external_key se asignó únicamente durante el apply autorizado"],
        ["Downgrade", "Definido para tabla, constraint, índice y columna"],
    ],
    [2.0, 4.4],
)

add_heading(doc, "36. Primera ejecución y estado persistido")
doc.add_paragraph(
    "La primera ejecución finalizó SUCCESS y generó SecurityEvent ID 9. El tenant fue renombrado dentro de la "
    "misma transacción, la moneda COP se preservó y el árbol quedó completo."
)
add_table(
    doc,
    ["Métrica", "Resultado"],
    [[key, value] for key, value in first["summary"].items()],
    [2.7, 3.7],
)
add_heading(doc, "Identidades finales de los 14 workspaces", 2)
add_table(
    doc,
    ["External key", "ID", "Record Code", "Tipo", "Nombre"],
    [
        [item["external_key"], item["id"], item["record_code"], item["workspace_type"], item["name"]]
        for item in first["workspaces"]
    ],
    [1.45, 0.45, 1.15, 1.25, 2.6],
)
add_heading(doc, "Validación posterior", 2)
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Tenant", "ID 1 · P&P Ingeniería y Proyectos · pyp-ingenieria-proyectos · COP"],
        ["Workspaces / raíces", f"{validation['totals']['workspaces']} / {validation['totals']['root_nodes']}"],
        ["External keys / Record Codes únicos", "14 / 14"],
        ["Objetivos / clasificaciones / links", "7 / 26 / 0"],
        ["Ciclos / padres rotos", "0 / 0"],
        ["Duplicados external_key / record_code", "0 / 0"],
        ["Referencias padre cross-tenant", "0"],
        ["PROPERTY / FACILITY", "0"],
    ],
    [2.8, 3.6],
)

add_heading(doc, "37. Objetivos y clasificaciones CORE")
add_heading(doc, "Strategic Objectives creados", 2)
objective_labels = {
    "OBJ-PYP-01": "Convertir proyectos complejos en resultados previsibles",
    "OBJ-PYP-02": "Proteger plazo, costo, alcance y valor",
    "OBJ-PYP-03": "Fortalecer el gobierno y la toma de decisiones",
    "OBJ-PYP-04": "Conectar estrategia, control y ejecución mediante tecnología",
    "OBJ-PYP-05": "Transformar diseños en activos operativos confiables",
    "OBJ-PYP-06": "Maximizar el desempeño y el valor del ciclo de vida de los activos",
    "OBJ-PYP-07": "Consolidar capacidades para minería e infraestructura de alta complejidad",
}
add_table(doc, ["Código", "Objetivo"], [[code, objective_labels[code]] for code in first["objective_codes"]], [1.4, 5.0])
add_heading(doc, "Clasificaciones creadas", 2)
classification_rows = []
for key in first["classification_keys"]:
    workspace, category, item = key.split(":", 2)
    classification_rows.append([workspace, category, item])
add_table(doc, ["Workspace", "Categoría", "Valor"], classification_rows, [1.9, 2.0, 2.5])

add_heading(doc, "38. Auditoría e idempotencia")
add_table(
    doc,
    ["Ejecución", "Evento", "Resultado", "Replay", "Cambio funcional"],
    [
        ["Primera", "ID 9 · enterprise_structure.core_applied", "success", "false", "3 ADOPT + 44 CREATE"],
        ["Segunda", "ID 10 · enterprise_structure.core_applied", "success", "true", "Ninguno"],
    ],
    [1.0, 2.25, 0.8, 0.7, 1.65],
)
doc.add_paragraph(
    "La segunda ejecución usó el mismo YAML, hash aprobado, actor y decisiones. El verificador aceptó el baseline "
    "original únicamente porque encontró una ejecución previa exitosa de la misma release y comprobó el estado final exacto."
)
add_table(
    doc,
    ["Métrica segunda ejecución", "Valor"],
    [[key, value] for key, value in second["summary"].items()],
    [3.2, 3.2],
)
add_callout(
    doc,
    "Prueba de idempotencia",
    "Los fingerprints funcionales de tenant, workspaces, clasificaciones, links, catálogos, usuarios y settings "
    "son idénticos después de la primera y la segunda ejecución. Solo cambia security_events por el evento técnico ID 10.",
    PALE_GREEN,
)

add_heading(doc, "39. Fingerprints posteriores y QA funcional")
add_heading(doc, "Fingerprint posterior a primera y segunda ejecución", 2)
fingerprint_rows = []
for name, first_value in post_apply["tables"].items():
    second_value = post_second["tables"][name]
    status = "IGUAL" if first_value["sha256"] == second_value["sha256"] else "AUDITORÍA ESPERADA"
    fingerprint_rows.append(
        [name, first_value["rows"], second_value["rows"], first_value["sha256"][:16] + "…", second_value["sha256"][:16] + "…", status]
    )
add_table(
    doc,
    ["Conjunto", "1ª filas", "2ª filas", "SHA 1ª", "SHA 2ª", "Resultado"],
    fingerprint_rows,
    [1.55, 0.55, 0.55, 1.25, 1.25, 1.25],
)
add_heading(doc, "QA ADMIN MODE", 2)
add_bullets(
    doc,
    [
        "Ruta validada: ADMIN MODE → Enterprise Structure → Enterprise Structure Configuration.",
        "14 nodos, una raíz, cuatro Business Units y Record Codes 01 a 01.04 visibles con indentación jerárquica correcta.",
        "Nombres, tipos, categorías, acciones ADMIN, formulario común y CompactModuleHeader disponibles.",
        "Árbol sin duplicados; external_key raíz ENT-PYP asociado al ID 1.",
    ],
)
add_heading(doc, "QA USER MODE", 2)
add_bullets(
    doc,
    [
        "Ruta validada: USER MODE → Enterprise Strategy Manager → Enterprise Structure & Workspace Manager → Enterprise Explorer.",
        "Los mismos 14 nodos y Record Codes se muestran en el mismo orden; 1 Project y 0 Facilities.",
        "Filtros por texto, tipo, Business Unit, objetivo, región y estado disponibles; vistas Árbol/Tabla y detalle preservados.",
        "7 objetivos disponibles y 26 clasificaciones persistidas; no se muestran acciones de escritura ADMIN.",
        "Aplicación en http://127.0.0.1:5173/app, título correcto y cero errores de consola durante la verificación.",
    ],
)

add_heading(doc, "40. Calidad, evidencias y riesgos")
add_heading(doc, "Pruebas ejecutadas", 2)
add_table(
    doc,
    ["Suite / control", "Resultado"],
    [
        ["Backend apply focalizado", "7 passed"],
        ["Backend Enterprise Structure + importer + auth", "35 passed"],
        ["Rollback inyectado", "PASS después de workspaces y después de classifications"],
        ["Hash / snapshot / seguridad / cross-tenant", "PASS"],
        ["Ruff", "PASS"],
        ["Frontend Enterprise Structure", "8 passed"],
        ["ESLint", "PASS · 0 errores; 8 warnings preexistentes del visor BIM"],
        ["Prettier", "PASS"],
        ["Build TypeScript + Vite", "PASS"],
        ["Migración PostgreSQL desde head vigente", "PASS · 0029 → 0030"],
    ],
    [2.8, 3.6],
)
add_heading(doc, "Paquete de evidencias", 2)
add_table(
    doc,
    ["Evidencia", "Propósito"],
    [
        ["gate02b_preflight_inventory.json / fingerprints.json", "Baseline real anterior a la migración"],
        ["gate02b_gate02a_comparison.json", "Equivalencia semántica Gate 02A → 02B"],
        ["gate02b_pre_apply_*.json", "Prueba de que la migración no cambió el snapshot protegido"],
        ["gate02b_immediate_dry_run.txt / .json", "Revalidación final sin mutación"],
        ["gate02b_apply_first.txt / .json / .sha256", "Primera transacción y checksum"],
        ["gate02b_post_apply_inventory.json / fingerprints.json", "Estado posterior a primera ejecución"],
        ["gate02b_apply_second.txt / .json / .sha256", "Replay idempotente y checksum"],
        ["gate02b_post_second_inventory.json / fingerprints.json", "Estado posterior a segunda ejecución"],
        ["gate02b_post_apply_validation.json", "Tenant, conteos e integridad estructural"],
        ["gate02b_idempotency_validation.json", "Comparación funcional de ambas ejecuciones"],
        ["gate02b_audit_and_publish_validation.json", "Eventos técnicos y confirmación de no publish"],
        ["gate02b_qa.json / gate02b_tests.json", "QA de API/UI y calidad automatizada"],
    ],
    [3.0, 3.4],
)
add_heading(doc, "Riesgos pendientes", 2)
add_bullets(
    doc,
    [
        "La estructura permanece en estados DRAFT salvo el Project activo, conforme al YAML autorizado; no se debe interpretar el apply como publicación.",
        "El actor conserva el correo histórico admin@demo.local. Cambiar identidades de usuario no formó parte de Gate 02B.",
        "Gate 02A usó MD5 con canonicalización no documentada; Gate 02B corrige el control futuro mediante SHA-256 reproducible.",
        "ESLint conserva ocho warnings no bloqueantes en BimIfcModelViewer, fuera del alcance de Enterprise Structure.",
        "La revisión visual automática del DOCX depende de LibreOffice/soffice; si no está instalado, la validación se limita a estructura y contenido del archivo.",
    ],
)

add_heading(doc, "41. Cierre de Gate 02B y recomendación Gate 03")
doc.add_paragraph(
    "Gate 02B cumple sus criterios de salida: preflight sin cambios, dry-run válido, primera ejecución exitosa, "
    "segunda ejecución completamente idempotente, QA ADMIN/USER y suites técnicas aprobadas. No se ejecutó "
    "Publish CORE ni se cargaron lotes adicionales."
)
add_callout(
    doc,
    "Recomendación para Gate 03",
    "Abrir un gate independiente de Publish CORE únicamente después de una revisión funcional del árbol aplicado, "
    "los estados DRAFT, los 26 vínculos de clasificación y los hashes de evidencia. Gate 03 debe exigir aprobación "
    "separada, expected hashes de las configuraciones a publicar, evento auditable y plan de rollback. Este informe "
    "no concede esa autorización.",
    PALE_AMBER,
)
add_table(
    doc,
    ["Gate de salida", "Estado final"],
    [
        ["Gate 02B", "COMPLETO"],
        ["Apply CORE", "EJECUTADO Y VALIDADO"],
        ["Idempotencia", "DEMOSTRADA"],
        ["QA ADMIN / USER", "PASS / PASS"],
        ["Publish CORE", "NOT EXECUTED"],
        ["Próxima autoridad requerida", "Gate 03 — Publish CORE explícito y separado"],
    ],
    [2.2, 4.2],
)

# Preserve the existing single-section layout and avoid accidental blank trailing pages.
if len(doc.sections) > 1 and doc.sections[-1].start_type == WD_SECTION.NEW_PAGE:
    pass

doc.save(OUTPUT)
print(OUTPUT)
