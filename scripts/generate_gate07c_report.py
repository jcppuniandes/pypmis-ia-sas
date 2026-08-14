from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO = Path(__file__).resolve().parents[1]
OUTPUT = (
    REPO
    / "outputs"
    / "technical_reports"
    / "Informe_Tecnico_Gate_07C_Strategic_Gate_Decision_20260813.docx"
)

# Preset: standard_business_brief. Header archetype: memo_masthead.
NAVY = "0D2A3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "16222B"
MUTED = "5B6870"
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
PALE_TEAL = "EAF7F6"


def set_run(run, *, size=11, color=INK, bold=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=120):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", indent)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row = table.rows[0]
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_row._tr.get_or_add_trPr().append(header_flag)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, NAVY)
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=9.1)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(doc, title, body, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_flag)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    set_run(cell.paragraphs[0].add_run(title.upper()), size=9, color=NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run(run, size=9.5)
    set_table_geometry(table, [9360], indent=160)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def bullets(doc, values):
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("P&Pmis Construction AI  |  "), size=8.5, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, end])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = Inches(1)
section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.08
for name, size, color, before, after in (
    ("Title", 23, NAVY, 0, 4),
    ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 15, BLUE, 15, 7),
    ("Heading 2", 12.5, DARK_BLUE, 10, 5),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(10.5)
    styles[name].paragraph_format.space_after = Pt(4)
    styles[name].paragraph_format.left_indent = Inches(0.5)
    styles[name].paragraph_format.first_line_indent = Inches(-0.25)

doc.core_properties.title = "Informe técnico Gate 07C — Strategic Gate Decision"
doc.core_properties.subject = "Configuración y validación de Strategic Gate Decision"
doc.core_properties.author = "P&Pmis Construction AI / Codex"
doc.core_properties.keywords = "Gate 07C, Strategic Gate Decision, Portfolio Intake, P&Pmis"
lang = OxmlElement("w:lang")
lang.set(qn("w:val"), "es-CO")
styles["Normal"]._element.get_or_add_rPr().append(lang)

header = section.header.paragraphs[0]
set_run(
    header.add_run("P&Pmis Construction AI  |  Gate 07C Technical Report"),
    size=8.5,
    color=MUTED,
    bold=True,
)
add_page_number(section.footer.paragraphs[0])

title = doc.add_paragraph(style="Title")
set_run(title.add_run("Informe técnico — Gate 07C"), size=23, color=NAVY, bold=True)
subtitle = doc.add_paragraph(style="Subtitle")
set_run(subtitle.add_run("Strategic Gate Decision Foundation & Decision Lifecycle"), size=13, color=MUTED)

add_table(
    doc,
    ["Campo", "Detalle"],
    [
        ["Aplicación", "P&Pmis Construction AI"],
        ["Modo / módulo", "USER y ADMIN / Enterprise Strategy Manager"],
        ["Submódulo", "Strategic Gate Decision"],
        ["Baseline", "Gate 07B Project Proposal en READY_FOR_STRATEGIC_GATE"],
        ["Migración", "20260813_0042 (head), PostgreSQL 16"],
        ["Fecha de corte", "13 de agosto de 2026"],
        ["Estado", "Implementado, probado y operativo en localhost"],
    ],
    [2400, 6960],
)

callout(
    doc,
    "Resultado ejecutivo",
    "Gate 07C quedó implementado como un registro estratégico separado, multirregistro y auditable. Consume únicamente Proposals preparadas, soporta APPROVE, RETURN, REJECT y DEFER, y limita APPROVE a READY_FOR_PORTFOLIO_INTAKE con can_create_portfolio_candidate=false. No crea Candidate, Project ni Workspace.",
    PALE_GREEN,
)

heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph(
    "La solución incorpora el ciclo DRAFT → SUBMITTED → IN_REVIEW → DECIDED, además de VOIDED y retorno controlado a preparador. Las decisiones cerradas son inmutables, las nuevas rondas crean registros adicionales y la UI está disponible en USER MODE y ADMIN MODE."
)

heading(doc, "2. REUSE BEFORE CREATE")
bullets(doc, [
    "Se reutilizaron Gate07B readiness, AdminConfiguration, AdminNumberSequence, Workspace Context, AuthContext, el resolver de permisos, SecurityEvent y el patrón ETag/Idempotency-Key.",
    "La evidencia permanece como referencias JSON; no se creó repositorio binario, motor BPM, microservicio ni Workspace de Decision.",
    "Backend y frontend se encapsularon en app/modules/strategic_gate y features/strategic-gate.",
])

doc.add_page_break()
heading(doc, "3. Aclaración status vs readiness")
add_table(doc, ["Contrato", "Valor", "Uso"], [
    ["ProjectProposal.status", "READY_FOR_STRATEGIC_GATE", "Estado persistido de la Proposal"],
    ["StrategicGateReadiness.status", "READY_FOR_STRATEGIC_GATE_DECISION", "Contrato de entrada calculado por Gate 07B"],
], [2800, 3300, 3260])
doc.add_paragraph("La implementación corrigió referencias heredadas que fusionaban ambos conceptos y conserva explícitamente esta diferencia en esquemas, servicios, UI y pruebas.")

heading(doc, "4. Modelo de dominio")
doc.add_paragraph("StrategicGateDecision es la fuente formal de la decisión. No es EnterpriseWorkspace y no se redujo a un campo de ProjectProposal. Guarda identidad, snapshots, actor, authority, resultado, hashes y auditoría.")

heading(doc, "5. Multirregistro y cardinalidad")
doc.add_paragraph("La cardinalidad es ProjectProposal 1:N StrategicGateDecision. decision_number y gate_round identifican cada ronda histórica. Un índice parcial impide más de una decisión activa por Proposal y gate type; las rondas cerradas nunca se sobrescriben.")

heading(doc, "6. Context type y gate type")
bullets(doc, [
    "context_type está preparado, pero operacionalmente sólo admite PROJECT_PROPOSAL.",
    "gate_type por defecto es PROJECT_PROPOSAL_GATE y permanece separado del outcome.",
    "No se creó un motor universal de gates.",
])

heading(doc, "7. Numeración")
doc.add_paragraph("AdminNumberSequence emite SGD-00001 y siguientes con alcance tenant. Preview consulta el siguiente valor sin consumirlo; create reserva el consecutivo dentro de la transacción y la restricción tenant/decision_number actúa como última defensa.")

heading(doc, "8. Modelo de datos")
add_table(doc, ["Grupo", "Campos principales", "Control"], [
    ["Identidad", "tenant_id, decision_number, context_type/id, project_proposal_id", "FK, tenant scope, índices y uniques"],
    ["Ciclo", "gate_type, gate_round, state, outcome, revision_version", "ETag y una activa por defecto"],
    ["Entrada", "proposal/readiness snapshots, hashes, Idea/Evaluation IDs", "Snapshot histórico no reescribible"],
    ["Gobierno", "criteria, checklist, authority, committee, conditions, evidence", "Configuración publicada y evidencia"],
    ["Auditoría", "prepared/submitted/reviewed/decided/voided, created/updated", "Actores, fechas y SecurityEvent"],
], [1800, 4800, 2760])

heading(doc, "9. Lifecycle")
add_table(doc, ["Desde", "Acción", "Hacia"], [
    ["DRAFT", "submit", "SUBMITTED"],
    ["SUBMITTED", "start-review", "IN_REVIEW"],
    ["SUBMITTED / IN_REVIEW", "return-to-preparer", "DRAFT"],
    ["IN_REVIEW", "decide", "DECIDED"],
    ["DRAFT / SUBMITTED / IN_REVIEW", "void", "VOIDED"],
], [2500, 3100, 3760])

heading(doc, "10. Outcomes")
add_table(doc, ["Outcome", "Efecto", "Portfolio Intake"], [
    ["APPROVE", "Proposal → STRATEGIC_GATE_APPROVED", "READY_FOR_PORTFOLIO_INTAKE"],
    ["RETURN", "Proposal → RETURNED", "GATE07C_REWORK_REQUIRED"],
    ["REJECT", "Proposal → STRATEGIC_GATE_REJECTED", "GATE07C_REWORK_REQUIRED"],
    ["DEFER", "Proposal → STRATEGIC_GATE_DEFERRED", "GATE07C_REWORK_REQUIRED"],
], [1800, 4200, 3360])

heading(doc, "11. ADR del estado post-decisión de Proposal")
doc.add_paragraph("Se amplió el enum real de Proposal para que el resultado estratégico sea visible sin recalcular toda la historia: APPROVE, REJECT y DEFER tienen estados explícitos; RETURN reutiliza RETURNED. StrategicGateDecision sigue siendo la fuente de verdad y preserva el outcome y el snapshot.")

heading(doc, "12. Readiness de entrada")
bullets(doc, [
    "Exige Proposal READY_FOR_STRATEGIC_GATE y readiness READY_FOR_STRATEGIC_GATE_DECISION con can_enter_strategic_gate=true.",
    "Bloquea evaluation obsoleta, Idea/evaluación fuente inválida, tenant distinto, ownership/portfolio incompleto o blockers Gate 07B.",
    "Create vuelve a validar bajo bloqueo de Proposal; Preview nunca persiste ni consume número.",
])

heading(doc, "13. Snapshots y hashes")
doc.add_paragraph("Se congelan readiness, Proposal, source Idea, accepted Idea Evaluation, Proposal Evaluation, score, Workspaces, objetivos, configuración, checklist y criterios. decision_hash identifica el resultado cerrado; readiness_hash acompaña el contrato de Portfolio Intake.")

heading(doc, "14. Comportamiento stale")
doc.add_paragraph("Submit y Decide recalculan el readiness/hash. Si la Proposal cambió después del snapshot, retornan HTTP 412 con razón STALE_READINESS. If-Match obsoleto retorna 412 ETAG_MISMATCH; no existe decisión sobre payload vencido.")

heading(doc, "15. Configuración ADMIN MODE")
doc.add_paragraph("ADMIN MODE → Enterprise Strategy Manager → Strategic Gate Decision muestra configuración publicada y permite clone-to-draft, edición, publicación y preview. Incluye Gate Types, Outcomes, Required Fields, Checklist, Criteria, Authority, Committee, Quorum, SoD, reglas de outcomes, output, numbering, inheritance y permissions.")

heading(doc, "16. Herencia")
doc.add_paragraph("La resolución sigue Enterprise → Business Unit → Portfolio y selecciona el override publicado más cercano al owning Workspace de la Proposal. El origen, revision y hash de configuración quedan capturados en la Decision.")

heading(doc, "17. Decision checklist")
doc.add_paragraph("El checklist publicado conserva PASS/FAIL/WARNING, blocking y evidence. La configuración por defecto cubre readiness, evaluation, business case, alineación, portfolio, riesgos, ROM cost, calendario, sponsor y reconocimiento de que funding aún no es requisito.")

heading(doc, "18. Decision criteria")
doc.add_paragraph("Gate 07C consume el Proposal Score y añade criterios de Strategic Fit, Value/Benefit, Affordability ROM, Risk Acceptability, Organizational Capacity, Timing, Portfolio Fit y Decision Conditions, sin duplicar la evaluación completa de Gate 07B.")

heading(doc, "19. Authority, committee y quorum")
doc.add_paragraph("Se soportan SINGLE_DECISION_MAKER y COMMITTEE. En comité se persisten members, roles, chair, votes/recommendations, quorum_required y quorum_met como snapshot. Decide bloquea un comité sin chair o sin quórum conforme a la política publicada.")

heading(doc, "20. Four-Eyes / SoD")
doc.add_paragraph("La política puede impedir que el decisor sea creator o evaluator de la Proposal. El control se ejecuta en backend y no tiene bypass automático para organization_admin; una prueba dedicada confirma HTTP 403 cuando se infringe.")

heading(doc, "21. Preview")
doc.add_paragraph("POST /preview y la entrada contextual de Proposal retornan el próximo SGD, Proposal/Idea/Evaluations, readiness/hash, Workspaces, objetivos, gate/config/checklist/authority, blockers y warnings con persisted=false. Ninguna secuencia se consume.")

heading(doc, "22. Transacción Create")
doc.add_paragraph("Create bloquea Proposal, valida tenant/status/readiness/hash, comprueba que no exista otra Decision activa, resuelve configuración publicada, reserva número, crea DRAFT, persiste snapshots/hashes, registra SecurityEvent y confirma de forma atómica.")

heading(doc, "23. Submit, review y decide")
doc.add_paragraph("Cada transición exige permiso, If-Match e Idempotency-Key. Submit valida payload, checklist y fuente vigente; start-review registra actor/fecha; decide exige IN_REVIEW, authority, quorum, SoD, readiness actual, outcome, reason y condiciones aplicables.")

heading(doc, "24. Transacción APPROVE")
doc.add_paragraph("APPROVE bloquea Decision y Proposal, revalida hash, autoridad, quórum y SoD, cierra la Decision, actualiza la Proposal, calcula Portfolio Intake readiness, emite eventos decided/approved y hace commit. No crea Candidate, Project ni Workspace.")

heading(doc, "25. RETURN, REJECT y DEFER")
doc.add_paragraph("Los tres outcomes cierran la Decision como DECIDED, actualizan el estado de Proposal y devuelven GATE07C_REWORK_REQUIRED. DEFER admite deferred_until; una ronda posterior crea un nuevo SGD. RETURN requiere que Gate 07B complete el retrabajo antes de una nueva entrada.")

heading(doc, "26. READY_FOR_PORTFOLIO_INTAKE")
callout(doc, "Contrato deliberadamente limitado", "status=READY_FOR_PORTFOLIO_INTAKE y can_create_portfolio_candidate=false. Incluye Decision, Proposal, Idea, evaluaciones, ownership, portfolio, objetivos, score, conditions, decision_hash, readiness_hash, blockers y warnings.", PALE_GREEN)

heading(doc, "27. Contrato Gate 07D")
doc.add_paragraph("Un futuro Gate 07D sólo podrá consumir Decision DECIDED/APPROVE y readiness READY_FOR_PORTFOLIO_INTAKE, preservando strategic_gate_decision_id, project_proposal_id, source_idea_id, target_portfolio_workspace_id y decision_hash. Gate 07D no fue iniciado.")

heading(doc, "28. Historial")
doc.add_paragraph("El detalle de Proposal lista todas sus decisiones por número, ronda, state, outcome, fecha, maker, conditions, reason y hash. Los registros DECIDED/VOIDED son inmutables mediante guardas de modelo; History expone SecurityEvents ordenados.")

heading(doc, "29. Integración Proposal e Idea")
doc.add_paragraph("Project Proposal incorpora un panel contextual de Strategic Gate Decisions y endpoints de list/preview. La trazabilidad mantiene Idea → Proposal(s) → Decision(s) usando source_idea_id y los IDs de evaluación aceptada y Proposal Evaluation.")

heading(doc, "30. USER MODE y ADMIN MODE")
add_table(doc, ["Modo", "Ruta", "Capacidades verificadas"], [
    ["USER", "Enterprise Strategy Manager → Strategic Gate Decision", "9 colas, búsqueda, preview, create, detail y acciones gobernadas"],
    ["ADMIN", "Enterprise Strategy Manager → Strategic Gate Decision", "Lista/config preview, clone, edit y publish"],
], [1300, 3900, 4160])

heading(doc, "31. Integración Workspace")
doc.add_paragraph("Para Enterprise, Business Unit y Portfolio, Workspace Context incorpora Strategic Gate Decisions después de Ideas y Project Proposals. Se reutiliza el contexto operativo existente y no se materializa Decision Workspace.")

heading(doc, "32. Allowed actions")
doc.add_paragraph("Backend calcula can_edit, can_submit, can_start_review, can_return_to_preparer, can_decide, can_void y can_create_new_round según state, permiso, outcome y vigencia. El frontend consume estas banderas y no infiere autorización sólo por status.")

heading(doc, "33. Permisos y roles")
add_table(doc, ["Rol", "Permisos relevantes"], [
    ["gate_preparer", "read, create, edit, submit"],
    ["gate_reviewer", "read, review"],
    ["gate_decision_maker", "read, decide"],
    ["gate_committee_member", "read, review, decide según configuración"],
    ["gate_configuration_admin", "read, admin.configure, admin.publish"],
], [3100, 6260])
doc.add_paragraph("Se registraron nueve permisos strategic_gate.* y se integraron con los grants tenant/workspace existentes.")

heading(doc, "34. ETag, idempotencia y audit")
doc.add_paragraph("revision_version genera ETag y todas las mutaciones exigen If-Match. Idempotency-Key cubre create, submit, review, return, decide, void y new-round. SecurityEvent registra created, updated, submitted, review_started, returned_to_preparer, decided y outcome específico, voided, new_round y eventos de configuración.")

heading(doc, "35. Decisión sobre attachments")
doc.add_paragraph("Se mantuvo el patrón evidence_refs JSON de Gates 07A/07B. No se añadió almacenamiento binario, duplicación documental ni acoplamiento a Project Workspace; futuras integraciones podrán referenciar el Document Manager existente.")

heading(doc, "36. APIs")
add_table(doc, ["Grupo", "Endpoints"], [
    ["USER", "options, preview, create/list/get/update, submit, start-review, return-to-preparer, decide, void, new-round, history, portfolio-intake-readiness"],
    ["Proposal", "GET list y POST preview por project_proposal_id"],
    ["Idea", "GET trazabilidad de Strategic Gate Decisions"],
    ["ADMIN", "configuration list/preview/clone/update/publish"],
], [1900, 7460])

heading(doc, "37. Persistencia y constraints")
doc.add_paragraph("La tabla principal strategic_gate_decisions usa FKs explícitas, unique tenant/decision_number, unique tenant/proposal/gate_type/gate_round, índice parcial para una activa y índices por tenant/state, Proposal, owning Workspace, target Portfolio y actores. JSON snapshots evitan tablas prematuras de reuniones/votos.")

heading(doc, "38. Migración y concurrencia")
doc.add_paragraph("Alembic 20260813_0042 es aditiva y reversible. Normaliza el valor Gate07B incorrecto antes de crear la tabla. Locks de Proposal/Decision/sequence, versioning, idempotencia, uniques y el índice parcial protegen numbering, create, submit, decide, approve duplicado y new-round. La contención fue validada por restricciones PostgreSQL; una carga paralela sostenida permanece como endurecimiento futuro.")

heading(doc, "39. PostgreSQL E2E")
doc.add_paragraph("docker-compose.gate07c.yml levantó PostgreSQL 16 desechable, creó baseline 0041, aplicó 0042, ejecutó 13 pruebas Idea/Proposal/Gate, revirtió a 0041 y reaplicó 0042. Resultado: PASS. La base persistente de localhost también quedó en 0042 head.")

heading(doc, "40. Pruebas backend, frontend y navegador")
add_table(doc, ["Control", "Resultado"], [
    ["Ruff backend/app backend/tests", "PASS"],
    ["Pytest regresión focal local", "19 PASS"],
    ["PostgreSQL 16 contenedor", "13 PASS + upgrade/downgrade/re-upgrade"],
    ["TypeScript + ESLint", "PASS"],
    ["Vitest estratégico + Proposal", "4 PASS"],
    ["Vite production build", "PASS, 2367 módulos"],
    ["Browser USER/ADMIN", "PASS; colas, SGD-00001 preview y configuración visibles"],
    ["Console browser", "0 errores / 0 warnings"],
], [3900, 5460])

heading(doc, "41. Regresión")
doc.add_paragraph("La regresión focal cubrió Gate 07A Idea Lifecycle, Gate 07B Project Proposal, Gate 07C Strategic Gate y Workspace Context. No se modificó comportamiento de CORE, EWS, Project/Physical Workspace ni módulos operativos.")

heading(doc, "42. Baseline antes y después")
add_table(doc, ["Aspecto", "Antes", "Después"], [
    ["Gate07B", "Proposal ready; una referencia fusionaba status/readiness", "Status READY_FOR_STRATEGIC_GATE y readiness separado"],
    ["Gate07C", "No existía registro formal", "Decision 1:N, lifecycle y outcomes"],
    ["Portfolio", "Sin intake contract", "Readiness limitado, sin Candidate"],
    ["Navegación", "Sin submódulo Gate 07C", "USER/ADMIN y Workspace Context"],
    ["DB", "Alembic 0041", "Alembic 0042 head"],
], [1800, 3600, 3960])

heading(doc, "43. Riesgos y deuda técnica")
bullets(doc, [
    "El diseño de committee/votes usa snapshots JSON; si aparecen consultas analíticas complejas convendrá normalizar en Gate posterior.",
    "La suite prueba transacciones y constraints en PostgreSQL, pero una campaña de estrés multiworker con barreras simultáneas sería el siguiente nivel de resiliencia.",
    "Los bundles IFC continúan siendo grandes; no pertenecen a Gate 07C, pero conviene mantener lazy loading y presupuesto de performance.",
    "El renderizador LibreOffice no estaba disponible; se usó Microsoft Word como fallback y se inspeccionaron visualmente todas las páginas exportadas.",
])

heading(doc, "44. Recomendación para Gate 07D")
doc.add_paragraph("Iniciar Gate 07D sólo mediante un prompt explícito. Debe consumir exclusivamente APPROVE + READY_FOR_PORTFOLIO_INTAKE, volver a validar decision_hash y tenant/portfolio, y decidir de forma separada si corresponde crear Portfolio Candidate. No reinterpretar APPROVE como FID.")

heading(doc, "45. Estado final exclusivo")
callout(doc, "Contrato de salida", "Una Decision cerrada produce exactamente uno de dos contratos: READY_FOR_PORTFOLIO_INTAKE para APPROVE, o GATE07C_REWORK_REQUIRED para RETURN, REJECT y DEFER. can_create_portfolio_candidate permanece false en ambos casos.", PALE_AMBER)

heading(doc, "46. Informe técnico actualizado")
doc.add_paragraph("Este documento consolida alcance, arquitectura, modelo, ADR, configuración, seguridad, migración, pruebas, operación local, límites y deuda de Gate 07C. El trabajo se detiene en Strategic Gate Decision y no inicia Gate 07D.")

heading(doc, "Anexo A. Archivos principales", level=2)
add_table(doc, ["Capa", "Ruta"], [
    ["Backend", "backend/app/modules/strategic_gate/{models,schemas,service,router}.py"],
    ["Migración", "backend/alembic/versions/20260813_0042_gate07c_strategic_gate_decision.py"],
    ["Pruebas", "backend/tests/test_strategic_gate_gate07c.py"],
    ["Frontend", "frontend/src/features/strategic-gate/"],
    ["Navegación", "frontend/src/navigation/applicationNavigation.ts"],
    ["PostgreSQL E2E", "docker-compose.gate07c.yml"],
], [2400, 6960])

heading(doc, "Anexo B. Verificación operacional", level=2)
bullets(doc, [
    "API: http://127.0.0.1:8000/api/v1/health/ready → ready; database=ok; redis=ok.",
    "Frontend: http://127.0.0.1:5173/app → HTTP 200.",
    "Servicios: api healthy, db healthy, redis healthy, frontend, worker y beat activos.",
    "Usuario validado: admin@demo.local; navegación USER/ADMIN comprobada sin errores de consola.",
])

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
