from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUT = (
    REPO
    / "outputs"
    / "technical_reports"
    / "Informe_Tecnico_Gate_07D_Hardening_Release_Closeout_20260820.docx"
)

# Preset: standard_business_brief. Header archetype: memo_masthead.
NAVY = "0D2A3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "16222B"
MUTED = "5B6870"
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
PALE_RED = "FCECEC"
PALE_TEAL = "EAF7F6"


def set_run(run, *, size=10.5, color=INK, bold=None, italic=False):
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top=75, start=110, bottom=75, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (
        ("top", top),
        ("start", start),
        ("bottom", bottom),
        ("end", end),
    ):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths, indent=110):
    table.autofit = False
    total = sum(widths)
    tbl_pr = table._tbl.tblPr
    for tag, value in (("tblW", total), ("tblInd", indent)):
        node = tbl_pr.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tbl_pr.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        column = OxmlElement("w:gridCol")
        column.set(qn("w:w"), str(width))
        grid.append(column)
    for row in table.rows:
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    header_row = table.rows[0]
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    header_row._tr.get_or_add_trPr().append(header_flag)
    for index, header in enumerate(headers):
        cell = header_row.cells[index]
        cell.text = header
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        shade(cell, NAVY)
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=8.8, color=WHITE, bold=True)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            cell = cells[index]
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            shade(cell, WHITE if row_index % 2 == 0 else LIGHT)
            set_cell_margin(cell)
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=8.9)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)


def callout(doc, title, body, fill=PALE_TEAL):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    header_flag = OxmlElement("w:tblHeader")
    header_flag.set(qn("w:val"), "true")
    table.rows[0]._tr.get_or_add_trPr().append(header_flag)
    cell = table.cell(0, 0)
    shade(cell, fill)
    set_cell_margin(cell, 130, 150, 130, 150)
    set_run(cell.paragraphs[0].add_run(title.upper()), size=8.8, color=NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run(run, size=9.3)
    set_table_geometry(table, [9360], indent=150)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc, text, level=1):
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True
    return paragraph


def bullets(doc, values):
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(3)


def numbered(doc, values):
    for index, value in enumerate(values, start=1):
        paragraph = doc.add_paragraph()
        paragraph.paragraph_format.left_indent = Inches(0.5)
        paragraph.paragraph_format.first_line_indent = Inches(-0.25)
        paragraph.add_run(f"{index}.  {value}")
        paragraph.paragraph_format.space_after = Pt(3)


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(paragraph.add_run("P&Pmis Construction AI  |  "), size=8.2, color=MUTED)
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run = paragraph.add_run()
    run._r.extend([begin, instr, end])


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.8)
section.bottom_margin = Inches(0.95)
section.left_margin = section.right_margin = Inches(0.85)
section.header_distance = section.footer_distance = Inches(0.4)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.2)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(5)
normal.paragraph_format.line_spacing = 1.06
for name, size, color, before, after in (
    ("Title", 22, NAVY, 0, 4),
    ("Subtitle", 12.5, MUTED, 0, 14),
    ("Heading 1", 14.5, BLUE, 13, 6),
    ("Heading 2", 11.8, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(10.1)
    styles[name].paragraph_format.space_after = Pt(3)
    styles[name].paragraph_format.left_indent = Inches(0.5)
    styles[name].paragraph_format.first_line_indent = Inches(-0.25)

doc.core_properties.title = (
    "Informe técnico Gate 07D — Hardening, release validation y closeout"
)
doc.core_properties.subject = (
    "Cierre de release del stage entry de Portfolio Planning y preparación FEL"
)
doc.core_properties.author = "P&Pmis Construction AI / Codex"
doc.core_properties.keywords = (
    "Gate 07D, Portfolio Planning, Gate 05B, Gate 07C, P&Pmis"
)
lang = OxmlElement("w:lang")
lang.set(qn("w:val"), "es-CO")
styles["Normal"]._element.get_or_add_rPr().append(lang)

header = section.header.paragraphs[0]
set_run(
    header.add_run("P&Pmis Construction AI  |  Gate 07D Release Closeout"),
    size=8.2,
    color=MUTED,
    bold=True,
)
add_page_number(section.footer.paragraphs[0])

title = doc.add_paragraph(style="Title")
set_run(title.add_run("Informe técnico — Gate 07D"), size=22, color=NAVY, bold=True)
subtitle = doc.add_paragraph(style="Subtitle")
set_run(
    subtitle.add_run("Hardening · Release validation · Closeout"),
    size=12.5,
    color=MUTED,
)

add_table(
    doc,
    ["Campo", "Detalle"],
    [
        ["Aplicación", "P&Pmis Construction AI"],
        [
            "Modo / módulo",
            "USER MODE · Enterprise Strategy Manager · Portfolio Manager",
        ],
        ["Submódulos", "Strategic Project Planning Entry / Portfolio Projects"],
        ["ADMIN MODE", "Portfolio Planning Entry & Membership"],
        [
            "Baseline",
            "Gate 07C APPROVE + READY_FOR_PORTFOLIO_INTAKE; Gate 05B Project Creation",
        ],
        ["Migración", "20260820_0043, sucesora de 20260813_0042"],
        ["Fecha de corte", "20 de agosto de 2026"],
        ["Estado final", "GATE07D_RELEASE_VALIDATED"],
    ],
    [2300, 7060],
)

callout(
    doc,
    "Resultado ejecutivo",
    "Gate 07D quedó implementado y cerrado para release como una entrada de etapa, no como un gate metodológico ni un motor de priorización. La validación final cubrió PostgreSQL 16 real, migración up/down/up, concurrencia, idempotencia, flujo USER/ADMIN en navegador, regresiones y protección de la línea base persistente.",
    PALE_GREEN,
)

heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph(
    "La solución conserva las autoridades existentes y completa la trazabilidad Idea → Proposal → Strategic Gate Decision → ProjectCreationRequest → Project Workspace → Portfolio membership. El estado terminal es exclusivamente READY_FOR_PORTFOLIO_PLANNING o GATE07D_REWORK_REQUIRED; ninguna ruta inicializa, activa o autoriza ejecución."
)

heading(doc, "2. Alcance implementado")
bullets(
    doc,
    [
        "Entrada estratégica gobernada desde USER MODE con preview no persistente y hashes de precondición.",
        "Extensión aditiva de ProjectCreationRequest para lineage, mapping y target Portfolio.",
        "Relación analítica Portfolio–Project N:M sin modificar el parent del árbol empresarial.",
        "Snapshot/hash de entrada de planificación y dos contratos de readiness desacoplados.",
        "Registro Portfolio Projects, contexto estratégico de Project y configuración ADMIN heredable.",
        "RBAC, ETag, four-eyes, idempotencia reutilizada y SecurityEvent.",
    ],
)

heading(doc, "3. Exclusiones explícitas")
callout(
    doc,
    "Límite del Gate 07D",
    "No se implementó PortfolioCandidate, Candidate Evaluation, Portfolio Evaluation, ranking, priorización, PDRI/FEL scoring, FID, Project initialization, activation ni módulos operativos. Las referencias PDRI/FEL son sugerencias de framework, nunca un score o decisión.",
    PALE_AMBER,
)

heading(doc, "4. REUSE BEFORE CREATE").paragraph_format.page_break_before = True
add_table(
    doc,
    ["Capacidad", "Decisión de reutilización", "Resultado"],
    [
        [
            "Gate 07C",
            "Consumir Decision e intake readiness vigentes",
            "Sin duplicar gate o evaluación",
        ],
        [
            "Gate 05B",
            "ProjectCreationRequest sigue siendo autoridad",
            "Mismo lifecycle, número y materialización",
        ],
        [
            "Workspace",
            "EnterpriseWorkspace sigue siendo Project identity",
            "No hay segunda identidad Project",
        ],
        [
            "Configuración",
            "AdminConfiguration versionada",
            "Default + override heredable",
        ],
        [
            "Seguridad",
            "Permission resolver y SecurityEvent",
            "Tenant scope, RBAC y auditoría",
        ],
        [
            "Concurrencia",
            "revision_version, If-Match y constraints",
            "Control optimista y unicidad",
        ],
    ],
    [1700, 4050, 3610],
)

heading(doc, "5. ADR arquitectónico")
doc.add_paragraph(
    "El ADR 43 formaliza Gate 07D como stage-entry bridge. La consecuencia principal es que APPROVE no equivale a FID: habilita planificación de portafolio y definición del proyecto, manteniendo el Project PENDING y la ejecución bloqueada."
)

heading(doc, "6. Flujo extremo a extremo")
numbered(
    doc,
    [
        "Seleccionar una StrategicGateDecision DECIDED/APPROVE cuyo intake esté READY_FOR_PORTFOLIO_INTAKE.",
        "Previsualizar lineage, número/record code, parent, template, manager, mapping y readiness sin persistir ni consumir secuencias.",
        "Crear una única ProjectCreationRequest de Gate 05B con snapshot y hashes Gate 07D.",
        "Ejecutar submit, start-review, approve por actor distinto y materialize usando Gate 05B.",
        "Crear Project PENDING y, en la misma transacción, su membresía target STRATEGIC_INTAKE.",
        "Persistir planning entry snapshot/hash y publicar READY o REWORK según los contratos de readiness.",
    ],
)

heading(doc, "7. Estados y contratos")
add_table(
    doc,
    ["Contrato", "Estado válido", "Significado"],
    [
        [
            "Gate 07C input",
            "READY_FOR_PORTFOLIO_INTAKE",
            "La decisión estratégica puede ingresar",
        ],
        [
            "ProjectCreationRequest",
            "draft → submitted → under_review → approved → created",
            "Lifecycle Gate 05B",
        ],
        [
            "Project Workspace",
            "pending",
            "Definición/planificación permitida; ejecución no",
        ],
        [
            "Portfolio readiness",
            "READY / BLOCKED",
            "Prerequisitos para evaluación futura",
        ],
        [
            "Definition readiness",
            "READY / BLOCKED",
            "Prerequisitos para definición FEL futura",
        ],
        [
            "Gate 07D",
            "READY_FOR_PORTFOLIO_PLANNING / GATE07D_REWORK_REQUIRED",
            "Salida exclusiva",
        ],
    ],
    [2250, 3350, 3760],
)

heading(doc, "8. Identidad y cardinalidad")
doc.add_paragraph(
    "ProjectCreationRequest continúa siendo registro de proceso y EnterpriseWorkspace continúa siendo identidad Project. Una unique tenant_id/strategic_gate_decision_id impide duplicar la entrada estratégica. PortfolioProjectMembership representa Portfolio N:M Project y permite múltiples portafolios activos sin re-parenting."
)

heading(doc, "9. Modelo de persistencia")
add_table(
    doc,
    ["Entidad", "Campos Gate 07D", "Controles"],
    [
        [
            "ProjectCreationRequest",
            "source context, Decision/Proposal/Idea IDs, hashes, target Portfolio, mapping revision/hash, source snapshot",
            "FKs, índices, unique por Decision",
        ],
        [
            "PortfolioProject\nMembership",
            "Portfolio, Project, source, target flag, effective dates, status, revision",
            "N:M, unique active, tenant scope",
        ],
        [
            "EnterpriseWorkspace\n.defaults_json",
            "_project planning context + _portfolio_planning snapshot/hash/status",
            "Project permanece PENDING",
        ],
        [
            "AdminConfiguration",
            "portfolio_planning_configuration",
            "published immutable, clone/update/publish",
        ],
        [
            "SecurityEvent",
            "linkage, materialization, membership, readiness, configuration",
            "Actor, tenant, timestamp, metadata",
        ],
    ],
    [2450, 4450, 2460],
)

heading(doc, "10. Membresía de Portfolio")
bullets(
    doc,
    [
        "La membresía inicial usa source=STRATEGIC_INTAKE e is_target_portfolio=true.",
        "La membresía target estratégica no puede eliminarse mediante la operación normal.",
        "Membresías adicionales MANUAL o RULE_BASED son compatibles y no cambian Project.parent_id.",
        "La remoción es lógica: ACTIVE → INACTIVE con effective_to y revision_version.",
        "If-Match protege creación sobre Project y remoción sobre Membership.",
    ],
)

heading(doc, "11. Snapshot de entrada de planificación")
doc.add_paragraph(
    "El snapshot consolida Decision, Proposal, Idea, target Portfolio, Project request/workspace, membresía, business need, preliminary scope, benefits, ROM cost, fechas, objetivos, riesgos, sponsor, manager y strategic conditions. planning_entry_hash permite verificar integridad y evita reinterpretar información futura como si hubiera existido al entrar."
)

heading(doc, "12. Readiness de Portfolio Evaluation")
doc.add_paragraph(
    "El contrato verifica membresía activa y requisitos configurables como strategic objectives, proposal score, ROM cost, expected benefits, risk summary y target dates. Devuelve required_source_data, available_source_data, blocking_issues y warnings; no calcula score, ranking ni prioridad."
)

heading(doc, "13. Readiness de Project Definition")
doc.add_paragraph(
    "El contrato verifica Project Type, parent, template, contexto estratégico, business need, preliminary scope, costos ROM, fechas, riesgos, sponsor y Project Manager según política. Puede sugerir un framework de definición por tipo de proyecto, pero no ejecuta PDRI ni FEL."
)

heading(doc, "14. Integración con Gate 05B").paragraph_format.page_break_before = True
doc.add_paragraph(
    "Gate 07D llama ProjectCreationService.create_request con metadata estratégica; no implementa otro ciclo de aprobación. La materialización ejecuta un hook local y atómico que valida nuevamente Gate 07C, crea la membresía y fija el estado final. Los hashes de submission y approval Gate 05B ahora incluyen el lineage Gate 07D."
)

heading(doc, "15. Four-eyes y segregación de funciones")
doc.add_paragraph(
    "Submit y review conservan los permisos Gate 05B. Approve exige un actor distinto del requestor y del último modificador. Materialize exige organization_admin o project_materialization_service. La UI presenta acciones calculadas por backend y no autoriza por inferencia de estado."
)

heading(doc, "16. Stale source y precondiciones")
add_table(
    doc,
    ["Control", "Momento", "Respuesta"],
    [
        [
            "expected_decision_hash",
            "Create Gate 07D",
            "HTTP 412 STALE_STRATEGIC_SOURCE",
        ],
        [
            "expected_readiness_hash",
            "Create Gate 07D",
            "HTTP 412 STALE_STRATEGIC_SOURCE",
        ],
        [
            "Gate 07C hashes persistidos",
            "Materialize Gate 05B",
            "Rollback transaccional si cambiaron",
        ],
        ["If-Match revision", "Mutaciones", "HTTP 412 ETAG_MISMATCH"],
        [
            "Unique Decision linkage",
            "Create concurrente",
            "Una única request estratégica",
        ],
        [
            "Unique active membership",
            "Membership concurrente",
            "Una relación activa por par",
        ],
    ],
    [2500, 3000, 3860],
)

heading(doc, "17. Idempotencia")
doc.add_paragraph(
    "Create retorna la entrada ya existente para la misma Decision sin reservar otra request. Materialize retorna ALREADY_CREATED para una request creada y vuelve a asegurar el snapshot/membership sin duplicar Project. Gate 05B conserva sus claves y constraints de idempotencia."
)

heading(doc, "18. Configuración ADMIN MODE")
bullets(
    doc,
    [
        "Política de membresía: STRATEGIC_INTAKE_ONLY, RULE_BASED o HYBRID.",
        "Fuentes: STRATEGIC_INTAKE obligatoria; RULE_BASED y MANUAL disponibles.",
        "Filtros declarativos con operadores limitados y sin ejecución de scoring.",
        "Mapping de Decision/Proposal a Project planning context.",
        "Recomendaciones de template, Project Type y framework de definición.",
        "Requisitos independientes de Portfolio Evaluation y Project Definition.",
        "Default tenant + overrides workspace heredables; clone/update/publish versionado.",
    ],
)

heading(
    doc, "19. Validación de configuración"
).paragraph_format.page_break_before = True
doc.add_paragraph(
    "El servicio rechaza membership policy desconocida, ausencia de STRATEGIC_INTAKE, operadores no permitidos y cualquier clave pdri_threshold o fel_score. Una configuración publicada es inmutable; los cambios se realizan en un draft clonado y se protegen con version/ETag."
)

heading(doc, "20. RBAC")
add_table(
    doc,
    ["Rol", "Capacidades principales"],
    [
        ["portfolio_intake_planner", "read, preview, create, readiness"],
        ["portfolio_membership_manager", "read/create/remove memberships, register"],
        ["portfolio_configuration_admin", "read/configure/publish ADMIN configuration"],
        ["organization_admin", "Cobertura organizacional y administración"],
        ["Gate 05B roles", "review, approve y materialize sin duplicación"],
    ],
    [3300, 6060],
)

heading(doc, "21. Eventos de seguridad")
add_table(
    doc,
    ["Fase", "Eventos principales"],
    [
        ["Entrada", "strategic_project_planning.request_created / request_linked"],
        ["Materialización", "strategic_project_planning.project_materialized"],
        ["Membresía", "portfolio_membership.created / removed"],
        [
            "Ready",
            "portfolio_planning.entry_ready / portfolio_project.ready_for_planning",
        ],
        [
            "Readiness",
            "portfolio_project.ready_for_evaluation / project_definition.ready",
        ],
        ["Rework", "portfolio_planning.entry_rework_required"],
        ["ADMIN", "configuration_cloned / updated / published"],
    ],
    [2550, 6810],
)

heading(doc, "22. APIs USER MODE")
add_table(
    doc,
    ["Grupo", "Rutas"],
    [
        [
            "Entrada",
            "GET options; GET portfolio-options; POST preview; POST create; GET decision entry/readiness",
        ],
        ["Portfolio", "GET /portfolios/{id}/projects"],
        ["Membership", "GET/POST /projects/{id}/portfolio-memberships; POST remove"],
        [
            "Readiness",
            "GET portfolio-planning, portfolio-evaluation y project-definition readiness",
        ],
    ],
    [2200, 7160],
)

heading(doc, "23. APIs ADMIN MODE")
doc.add_paragraph(
    "La configuración expone list, preview, clone, update y publish bajo /strategic-project-planning/admin/configurations. Los endpoints aplican permisos de administración, tenant scope, validación de contenido y control de versión."
)

heading(doc, "24. UI USER MODE")
bullets(
    doc,
    [
        "Portfolio Manager → Strategic Project Planning Entry: selector de decisions, lineage, prefill gobernado, readiness y acciones Gate 05B.",
        "Portfolio Manager → Portfolio Projects: registro de Projects por Portfolio y sus readiness.",
        "No hay botones Activate, FID, score, Candidate o priorización.",
        "Estados blocked/ready, issues y warnings se muestran desde el backend.",
    ],
)

heading(doc, "25. UI ADMIN MODE")
doc.add_paragraph(
    "Enterprise Strategy Manager → Portfolio Planning Entry & Membership permite revisar la configuración efectiva, clonar un draft, editar JSON gobernado y publicar. La pantalla reutiliza CompactModuleHeader y la estructura visual existente."
)

heading(doc, "26. Contexto operativo de Workspace")
doc.add_paragraph(
    "Un Portfolio muestra Portfolio Projects. Un Project de origen STRATEGIC_GATE muestra Strategic Context, Portfolio Memberships, Portfolio Planning Readiness y Project Definition Readiness. Projects ordinarios conservan su navegación anterior. Para un Project PENDING, sólo estas vistas de planificación y Home/Overview son seguras; Scope/Schedule/Cost siguen ocultas."
)

heading(doc, "27. Migración 20260820_0043")
bullets(
    doc,
    [
        "Añade once campos de lineage Gate 07D a project_creation_requests.",
        "Crea FKs PostgreSQL, índices de consulta y unique tenant/Decision.",
        "Crea portfolio_project_memberships con índices y unique parcial ACTIVE.",
        "Downgrade elimina tabla, índices, FKs y columnas en orden seguro.",
        "Upgrade es idempotente frente a columnas/tablas preexistentes.",
    ],
)

heading(doc, "28. PostgreSQL E2E y CI")
doc.add_paragraph(
    "docker-compose.gate07d.yml ejecutó un PostgreSQL 16.14 desechable. El cierre preparó el baseline 0042, validó su esquema, aplicó 0043, validó columnas/FKs/índices, ejecutó la suite, revirtió a 0042, volvió a validar, reaplicó 0043 y repitió los casos Gate 07D clave. El workflow Pilot Readiness CI incorpora el mismo gate reproducible."
)
callout(
    doc,
    "Resultado PostgreSQL",
    "PASS: 59 pruebas en la primera ejecución, downgrade 0043→0042 y re-upgrade 0042→0043 sin pérdida; esquema validado en cada estado; 2 pruebas Gate 07D repetidas al final. Alembic conserva una sola cabeza: 20260820_0043.",
    PALE_GREEN,
)

heading(doc, "29. Defectos detectados y corregidos")
add_table(
    doc,
    ["Hallazgo", "Corrección aplicada", "Evidencia"],
    [
        [
            "Índices Gate 07D superaban 63 caracteres",
            "Nombres PostgreSQL abreviados y explícitos",
            "Upgrade 0042→0043 PASS",
        ],
        [
            "Create concurrente propagaba UniqueViolation",
            "Captura exclusiva del constraint estratégico, rollback y relectura autoritativa",
            "Un request; sin número huérfano",
        ],
        [
            "Fixture Gate 07C sensible a Four-Eyes",
            "Decisor organization_admin distinto del autor",
            "Reejecución estable",
        ],
        [
            "Navegador E2E apuntaba a localhost interno",
            "Reescritura al API del stack y origen CORS permitido",
            "Tres pasadas completas",
        ],
        [
            "Project Type no pertenecía al catálogo",
            "Uso del código publicado capital",
            "Creación Gate 05B PASS",
        ],
        [
            "Chequeo no-alcance incluía todo el menú",
            "Aserción acotada al workspace Gate 07D",
            "Sin falso positivo",
        ],
    ],
    [2350, 4550, 2460],
)

heading(doc, "30. Validación backend")
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Ruff check/format: app, tests y migración 0043", "PASS · 188 archivos"],
        ["PostgreSQL 16: Gate 07D + 05B + 07C + Context", "PASS · 59 pruebas"],
        ["Migración up/down/up y esquema", "PASS · 0042/0043/0042/0043"],
        ["Reejecución Gate 07D post-migración", "PASS · 2 pruebas"],
        [
            "Concurrencia create/materialize/membership",
            "PASS · unicidad e idempotencia",
        ],
        ["Stale decision/readiness/ETag", "PASS · HTTP 412 esperado"],
        [
            "Regresión 07A/07B/Enterprise/Initialization",
            "75 PASS, 2 SKIP; 2 aislamientos PASS",
        ],
        ["No portfolio_candidates table", "PASS"],
    ],
    [4800, 4560],
)

heading(doc, "31. Validación frontend")
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Prettier src + spec E2E", "PASS"],
        ["ESLint max-warnings=10", "PASS · 0 errores, 9 advertencias heredadas"],
        ["TypeScript + Vite production build", "PASS · 2.373 módulos"],
        ["Vitest Gate 07D", "PASS · 1 prueba"],
        ["Vitest global", "193 PASS; 1 timeout por recursos, repetido PASS en 62 s"],
        [
            "Playwright full-stack USER/ADMIN",
            "PASS ×3 consecutivas: 46,5 s / 40,6 s / 40,4 s",
        ],
        ["Red/consola/page errors Gate 07D", "0 errores inesperados"],
    ],
    [4550, 4810],
)

heading(doc, "32. Protección de la línea base persistente")
add_table(
    doc,
    ["Indicador", "Antes", "Después"],
    [
        ["Alembic persistente", "20260813_0042", "20260813_0042"],
        ["Enterprise workspaces", "14", "14"],
        ["Project workspaces", "1 activo", "1 activo"],
        ["Requests / Ideas / Proposals / Decisions", "0 / 0 / 0 / 0", "0 / 0 / 0 / 0"],
        ["Record codes", "14 / 14 únicos", "14 / 14 únicos"],
        ["Secuencias gobernadas", "12 intactas", "12 intactas"],
        ["Columnas/tabla Gate 07D", "0 / ausente", "0 / ausente"],
        ["Core release", "ES-PYP-CORE-RECONCILED-20260809:published", "Sin cambio"],
    ],
    [3800, 2780, 2780],
)

heading(doc, "33. Criterios de aceptación cubiertos")
add_table(
    doc,
    ["Criterio", "Evidencia"],
    [
        [
            "Sin nueva Project identity",
            "Gate 05B materializa EnterpriseWorkspace una vez",
        ],
        ["Sin Candidate", "Modelo, APIs, UI y test explícito"],
        ["Project PENDING", "Policy + test de materialización"],
        ["Target membership", "STRATEGIC_INTAKE active/target"],
        ["Múltiples Portfolios", "Segunda membresía manual sin cambiar parent"],
        ["Readiness separado", "Dos contratos y rutas independientes"],
        ["No scoring/FID", "Config validator + ausencia de controles"],
        ["Trazabilidad completa", "IDs, snapshots, hashes y SecurityEvent"],
    ],
    [3900, 5460],
)

heading(doc, "34. Observaciones residuales")
bullets(
    doc,
    [
        "Una prueba histórica AppFlow agotó el límite al compartir recursos con el build; aprobó en aislamiento dentro del mismo límite. No afecta Gate 07D.",
        "Dos pruebas heredadas de Enterprise Structure comparten datos si siguen a Gate 07A en el mismo proceso; ambas aprobaron en aislamiento. Es deuda de aislamiento de fixtures, no defecto funcional.",
        "ESLint conserva nueve advertencias preexistentes del visor BIM y ProjectCreationWorkspace; Gate 07D no añadió advertencias.",
        "Las futuras decisiones de Portfolio Evaluation, FEL/PDRI, FID, initialization y activation siguen fuera de alcance y requieren un prompt independiente.",
    ],
)

heading(doc, "35. Recomendaciones operativas")
numbered(
    doc,
    [
        "Publicar al menos un Project Template Gate 05A aplicable a Portfolio/Program.",
        "Mantener Project Types, managers y objetivos estratégicos activos antes del intake.",
        "Configurar overrides sólo cuando un Portfolio necesite requisitos distintos del default tenant.",
        "Usar el registro Portfolio Projects como entrada a módulos futuros, no como motor de priorización actual.",
        "No activar el Project hasta que un gate posterior formalice FID y se complete la inicialización correspondiente.",
    ],
)

heading(doc, "36. Estado de entrega")
callout(
    doc,
    "GATE07D_RELEASE_VALIDATED",
    "La implementación, el hardening y el closeout quedaron validados. Gate 07D termina en Portfolio Planning stage entry y preparación FEL; no inicia Gate 07E, Portfolio Evaluation, PDRI/FEL assessment, FID, initialization ni activation.",
    PALE_GREEN,
)

heading(doc, "Anexo A. Archivos principales", level=2)
add_table(
    doc,
    ["Capa", "Ruta"],
    [
        [
            "Backend",
            "backend/app/modules/portfolio_planning/{models,schemas,service,router}.py",
        ],
        [
            "Gate 05B",
            "backend/app/modules/project_creation/{models,schemas,service}.py",
        ],
        [
            "Migración",
            "backend/alembic/versions/20260820_0043_gate07d_portfolio_planning.py",
        ],
        ["Navigator", "backend/app/modules/workspace_context/navigator.py"],
        ["Backend tests", "backend/tests/test_portfolio_planning_gate07d*.py"],
        ["Frontend", "frontend/src/features/portfolio-planning/"],
        ["Navegación", "frontend/src/navigation/applicationNavigation.ts"],
        ["Browser E2E", "frontend/e2e/portfolio-planning-entry.spec.ts"],
        [
            "PostgreSQL",
            "docker-compose.gate07d.yml + backend/tests/postgres/*gate07d*.py",
        ],
        ["ADR", "docs/43-adr-gate07d-portfolio-planning-stage-entry.md"],
    ],
    [2300, 7060],
)

heading(
    doc, "Anexo B. Matriz de salida", level=2
).paragraph_format.page_break_before = True
add_table(
    doc,
    ["Condición", "Salida Gate 07D", "Acción permitida"],
    [
        [
            "Workspace PENDING + target membership + ambos readiness READY",
            "READY_FOR_PORTFOLIO_PLANNING",
            "Entrar a Portfolio Evaluation y Project Definition futuros",
        ],
        [
            "Falta requisito Portfolio o Definition",
            "GATE07D_REWORK_REQUIRED",
            "Corregir configuración/datos, sin ejecución",
        ],
        [
            "Gate 07C stale/no APPROVE",
            "Precondition/eligibility error",
            "Refrescar o resolver en Gate 07C",
        ],
        [
            "Project ACTIVE",
            "Ineligible for Gate 07D",
            "No reinterpretar como stage entry",
        ],
    ],
    [3350, 2950, 3060],
)

heading(doc, "Anexo C. Referencia de verificación", level=2)
bullets(
    doc,
    [
        "Backend focal: pytest test_portfolio_planning_gate07d.py test_portfolio_planning_gate07d_hardening.py test_project_creation_process.py test_strategic_gate_gate07c.py test_workspace_context.py -q --no-cov.",
        "Frontend: Prettier check, ESLint max-warnings=10, Vitest, TypeScript/Vite build y tres ejecuciones Playwright contra el stack PostgreSQL aislado.",
        "PostgreSQL: docker compose -p pypmis-gate07d -f docker-compose.gate07d.yml up --build --abort-on-container-exit --exit-code-from gate07d-test.",
        "Resultado de release exclusivo: GATE07D_RELEASE_VALIDATED.",
    ],
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
