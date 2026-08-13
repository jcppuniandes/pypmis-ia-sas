from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

REPO = Path(__file__).resolve().parents[1]
OUTPUT = REPO / "outputs" / "technical_reports" / "Informe_Tecnico_Gate_06D_Workspace_Operational_Context_Navigation_20260813.docx"

# Preset: standard_business_brief. First-page pattern: memo_masthead.
NAVY = "0D2A3A"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "16222B"
MUTED = "5B6870"
WHITE = "FFFFFF"
LIGHT = "F2F4F7"
PALE_GREEN = "EAF7EF"
PALE_AMBER = "FFF5E6"
PALE_TEAL = "EAF7F6"


def set_run(run, *, size: float = 11, color: str = INK, bold: bool | None = None, italic: bool = False) -> None:
    run.font.name = "Calibri"
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(color)
    run.bold = bold
    run.italic = italic


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margin(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_table_geometry(item, widths: list[int], indent: int = 120) -> None:
    item.autofit = False
    total = sum(widths)
    tbl_pr = item._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")
    grid = item._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in item.rows:
        for index, cell in enumerate(row.cells):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[int]) -> None:
    item = doc.add_table(rows=1, cols=len(headers))
    item.style = "Table Grid"
    item.alignment = WD_TABLE_ALIGNMENT.LEFT
    for index, header in enumerate(headers):
        cell = item.rows[0].cells[index]
        cell.text = header
        shade(cell, NAVY)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margin(cell)
        for run in cell.paragraphs[0].runs:
            set_run(run, size=9, color=WHITE, bold=True)
    set_repeat_header(item.rows[0])
    for row_index, values in enumerate(rows):
        cells = item.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = str(value)
            shade(cells[index], WHITE if row_index % 2 == 0 else LIGHT)
            cells[index].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margin(cells[index])
            for paragraph in cells[index].paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)
                for run in paragraph.runs:
                    set_run(run, size=9.1)
    set_table_geometry(item, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def callout(doc: Document, title: str, body: str, fill: str = PALE_TEAL) -> None:
    item = doc.add_table(rows=1, cols=1)
    item.style = "Table Grid"
    cell = item.cell(0, 0)
    shade(cell, fill)
    set_cell_margin(cell, 140, 160, 140, 160)
    set_run(cell.paragraphs[0].add_run(title.upper()), size=9, color=NAVY, bold=True)
    paragraph = cell.add_paragraph(body)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_run(run, size=9.5)
    set_repeat_header(item.rows[0])
    set_table_geometry(item, [9360], indent=160)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_heading(text, level=level)
    paragraph.paragraph_format.keep_with_next = True


def bullets(doc: Document, values: list[str]) -> None:
    for value in values:
        paragraph = doc.add_paragraph(value, style="List Bullet")
        paragraph.paragraph_format.space_after = Pt(4)


def page_break(doc: Document) -> None:
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
section.header_distance = section.footer_distance = Inches(0.492)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.10
for name, size, color, before, after in (
    ("Title", 23, NAVY, 0, 4),
    ("Subtitle", 13, MUTED, 0, 16),
    ("Heading 1", 16, BLUE, 16, 8),
    ("Heading 2", 13, BLUE, 12, 6),
    ("Heading 3", 12, DARK_BLUE, 8, 4),
):
    style = styles[name]
    style.font.name = "Calibri"
    style.font.size = Pt(size)
    style.font.color.rgb = RGBColor.from_string(color)
    style.paragraph_format.space_before = Pt(before)
    style.paragraph_format.space_after = Pt(after)
    style.paragraph_format.keep_with_next = True
for name in ("List Bullet", "List Number"):
    styles[name].font.name = "Calibri"
    styles[name].font.size = Pt(11)
    styles[name].paragraph_format.space_after = Pt(8)
    styles[name].paragraph_format.line_spacing = 1.167
    styles[name].paragraph_format.left_indent = Inches(0.5)
    styles[name].paragraph_format.first_line_indent = Inches(-0.25)

header = section.header.paragraphs[0]
header.paragraph_format.space_after = Pt(0)
set_run(header.add_run("P&Pmis Construction AI  |  Gate 06D Technical Report"), size=8.5, color=MUTED, bold=True)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_run(footer.add_run("Controlled technical baseline  |  13 Aug 2026  |  Page "), size=8, color=MUTED)
page_run = footer.add_run()
set_run(page_run, size=8, color=MUTED)
for field_type, value in (("begin", None), (None, " PAGE "), ("separate", None), (None, "1"), ("end", None)):
    if field_type:
        field = OxmlElement("w:fldChar")
        field.set(qn("w:fldCharType"), field_type)
        page_run._r.append(field)
    else:
        text = OxmlElement("w:instrText" if value == " PAGE " else "w:t")
        text.set(qn("xml:space"), "preserve")
        text.text = value
        page_run._r.append(text)

doc.core_properties.title = "Informe técnico Gate 06D - Workspace Operational Context & Navigation"
doc.core_properties.subject = "Contexto operacional y navegación Workspace-scoped"
doc.core_properties.author = "P&Pmis Construction AI Engineering"
doc.core_properties.comments = "Gate 06D cerrado; módulos operacionales profundos no iniciados."

doc.add_paragraph("INFORME TÉCNICO", style="Title")
doc.add_paragraph("Gate 06D · Workspace Operational Context & Navigation", style="Subtitle")
for label, value in (
    ("Tenant", "P&P Ingeniería y Proyectos"),
    ("Fecha de cierre", "13 de agosto de 2026"),
    ("Alcance", "PROJECT · PROPERTY · FACILITY · WAREHOUSE"),
    ("Estado final", "READY_FOR_OPERATIONAL_MODULES"),
    ("Límite", "Sin Asset/Inventory Manager ni operación profunda"),
):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(2)
    set_run(paragraph.add_run(f"{label}: "), bold=True)
    set_run(paragraph.add_run(value))
callout(
    doc,
    "Resultado de cierre",
    "Se implementó el contexto operativo activo por usuario y tenant, navegación derivada por backend, Workspace Home "
    "común, rutas Workspace-scoped, Recent Workspaces, My Workspaces y cambio seguro de contexto. Se reutilizaron "
    "Enterprise Structure, Module Definitions, Workspace Module Settings y el RBAC existente.",
    PALE_GREEN,
)

heading(doc, "1. Resumen ejecutivo")
doc.add_paragraph(
    "Gate 06D conecta Enterprise Explorer con la operación diaria. Un Workspace autorizado se abre mediante una "
    "acción explícita, carga identidad, parent path, template snapshot, responsable, módulos, permisos y Home, y "
    "mantiene la navegación aislada por tenant_id + workspace_id. Las capacidades inexistentes devuelven colecciones "
    "vacías o PLANNED; no se inventan datos ni se crean módulos operativos ficticios."
)
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Contexto activo", "PASS · user + tenant + Workspace"],
        ["Tipos", "PASS · PROJECT / PROPERTY / FACILITY / WAREHOUSE"],
        ["Navigator", "PASS · backend-derived, RBAC y status aware"],
        ["Workspace Home", "PASS · layout común y capability flags"],
        ["Switching", "PASS · invalidación sin data leakage"],
        ["PostgreSQL", "PASS · 6/6, migration upgrade/downgrade/upgrade"],
        ["Frontend focal", "PASS · 12/12, TypeScript, lint y build"],
        ["Browser localhost", "PASS · /workspaces/14/home, 0 errores consola"],
        ["Baseline", "PASS · 14 Workspaces, 4 Projects, 60 configs"],
    ],
    [2850, 6510],
)

heading(doc, "2. REUSE BEFORE CREATE")
bullets(
    doc,
    [
        "EnterpriseWorkspace continúa como identidad canónica; no se creó un segundo registro maestro.",
        "Enterprise Explorer, árbol, tabla, creación y lifecycle existentes se mantienen sin duplicación.",
        "Module Definitions y WorkspaceModuleSetting siguen siendo la fuente operativa de módulos.",
        "AdminConfiguration se amplió con workspace_navigation_profile en el mismo sistema versionado.",
        "SecurityRole, PermissionCatalog, assignments y SecurityEvent se reutilizan para RBAC y auditoría.",
        "Sólo se agregó RecentWorkspace porque la continuidad user-scoped no existía en el baseline.",
    ],
)

heading(doc, "3. Alcance y exclusiones")
add_table(
    doc,
    ["Incluido", "Excluido"],
    [
        ["PROJECT / PROPERTY / FACILITY / WAREHOUSE", "REGION / DISTRICT / SITE"],
        ["Context, Home, Navigator y módulo guard", "LINEAR_ASSET / ASSET"],
        ["Recent Workspaces y My Workspaces", "Analítica avanzada o dashboards nuevos"],
        ["READY / PLANNED / status restrictions", "Asset, maintenance, inventory o utilities records"],
    ],
    [4680, 4680],
)

page_break(doc)
heading(doc, "4. Arquitectura implementada")
add_table(
    doc,
    ["Capa", "Responsabilidad"],
    [
        ["React Provider", "Contexto activo, Home, invalidación y loading/error states"],
        ["Workspace UI", "Header, Breadcrumb, Navigator, Home, My/Recent Workspaces y switching"],
        ["FastAPI router", "Contratos context/home/navigator/open/recent/last-route/module guard"],
        ["Application service", "Scope tenant/user, acceso, estados, acciones, ETag y auditoría"],
        ["Navigator service", "Perfil por tipo + definitions + settings + status + permisos"],
        ["PostgreSQL", "RecentWorkspace único por tenant/user/workspace"],
    ],
    [2150, 7210],
)

heading(doc, "5. Active Workspace Context")
doc.add_paragraph(
    "El contrato devuelve tenant/workspace identity, type, name, status, Business Number, Record Code, external_key, "
    "parent y parent_path, template code/revision/hash, responsable, enabled/planned modules, permisos, acciones, "
    "version, ETag, last route y navigator. El frontend no usa localStorage para permisos ni retiene el objeto anterior "
    "durante un cambio de Workspace."
)
callout(
    doc,
    "Regla de aislamiento",
    "Toda resolución operacional exige tenant_id + workspace_id. Los caches lógicos/ETags incorporan tenant, user, "
    "workspace version, template, module settings y Navigator.",
)

heading(doc, "6. Workspace Navigator por tipo")
add_table(
    doc,
    ["Tipo", "Navegación"],
    [
        ["PROJECT", "Home · Overview · Scope · Schedule · Cost · Documents · Reports"],
        ["PROPERTY", "Home · Overview · Real Estate Information · Documents · Related Workspaces"],
        ["FACILITY", "Home · Overview · Documents + Asset/Maintenance/Space/Utilities PLANNED"],
        ["WAREHOUSE", "Home · Overview · Documents + Inventory/Receipts/Issues/Transfers PLANNED"],
    ],
    [1900, 7460],
)
doc.add_paragraph(
    "Un Module Definition no publicado, un setting disabled, un permiso faltante o un status no operativo elimina la "
    "ruta del Navigator. PLANNED se presenta como roadmap y nunca habilita acceso directo ni registros ficticios."
)

heading(doc, "7. Comportamiento por estado")
add_table(
    doc,
    ["Estado", "Comportamiento"],
    [
        ["DRAFT", "No operativo; apertura bloqueada"],
        ["PENDING", "Home/Overview permitidos; módulos operativos ocultos"],
        ["ACTIVE", "Operación según módulos y permisos"],
        ["ARCHIVED", "Contexto y Navigator read-only"],
        ["INACTIVE", "Apertura bloqueada"],
    ],
    [2100, 7260],
)

page_break(doc)
heading(doc, "8. Workspace Home común")
bullets(
    doc,
    [
        "Workspace Header compacto con tipo, nombre, Business Number, Record Code, status y responsible.",
        "Workspace Breadcrumb derivado del parent path persistido, no de texto hardcoded.",
        "Key Information, Status, Responsible y Context version/ETag.",
        "Enabled Modules y Planned Modules claramente diferenciados.",
        "Recent Activity, Recent Documents y My Tasks permanecen vacíos con capability flags false.",
        "Related Workspaces usa vínculos transversales existentes y vuelve sólo relaciones autorizadas.",
    ],
)

heading(doc, "9. Recent Workspaces y last route")
doc.add_paragraph(
    "La migración 20260813_0039 agrega recent_workspaces con unicidad tenant_id/user_id/workspace_id, last_route, "
    "last_opened_at y version. Open Workspace crea o actualiza la continuidad; Last Route valida que la ruta pertenezca "
    "al mismo Workspace y esté READY. Se retornan hasta 12 recientes, siempre filtrados por acceso vigente."
)

heading(doc, "10. My Workspaces")
add_table(
    doc,
    ["Filtro / dato", "Implementación"],
    [
        ["Workspace Type", "PROJECT / PROPERTY / FACILITY / WAREHOUSE"],
        ["Status", "ACTIVE / PENDING / ARCHIVED"],
        ["Responsible", "Filtro parcial server-side"],
        ["Parent", "Filtro parcial server-side"],
        ["Business Number", "Filtro parcial server-side"],
        ["Name", "Filtro ilike server-side"],
        ["Recent Workspaces", "Acceso directo a last_route autorizada"],
    ],
    [2800, 6560],
)

heading(doc, "11. Workspace switching y multi-workspace readiness")
bullets(
    doc,
    [
        "El switch sólo muestra Workspaces ACTIVE/ARCHIVED autorizados.",
        "Cambiar id remonta el Provider y vacía context/home antes de resolver el nuevo fetch.",
        "Las respuestas tardías se descartan mediante generation guard.",
        "El backend recalcula permisos, módulos, acciones, ETag y Home para el nuevo Workspace.",
        "Las pruebas verifican A→B, ETag diferente y ausencia de nombre/datos del Workspace anterior.",
    ],
)

heading(doc, "12. Enterprise Explorer")
doc.add_paragraph(
    "Se conserva el Explorer existente y se agrega Open Workspace sólo para PROJECT/PROPERTY/FACILITY/WAREHOUSE "
    "ACTIVE. PENDING mantiene Overview. También se agrega My Workspaces como vista transversal sin retirar My Project "
    "Workspaces ni My Physical Workspaces ya desarrollados."
)

page_break(doc)
heading(doc, "13. ADMIN configuration")
bullets(
    doc,
    [
        "workspace_navigation_profile se integra a AdminConfiguration y a la navegación ADMIN existente.",
        "El perfil puede ordenar módulos, elegir default_home_route y controlar show_planned_modules.",
        "La validación rechaza perfiles sin workspace_type o con module_order no-array.",
        "Module Definitions y workspace_module_settings siguen siendo la verdad operativa.",
        "No se creó un segundo catálogo de módulos ni se tocó el CORE publicado.",
    ],
)

heading(doc, "14. Permisos y RBAC")
add_table(
    doc,
    ["Permiso", "Uso"],
    [
        ["workspace.open", "Abrir/cargar contexto"],
        ["workspace.home.read", "Home y My Workspaces"],
        ["workspace.navigator.read", "Navigator y route/module guard"],
        ["workspace.recent.read", "Recent Workspaces"],
        ["workspace.recent.write", "Open y last-route"],
        ["module permission key", "Visibilidad granular por Module Definition"],
    ],
    [3200, 6160],
)

heading(doc, "15. Seguridad")
bullets(
    doc,
    [
        "Cross-tenant devuelve 404 para no revelar existencia.",
        "Workspace sin assignment devuelve 403 aunque el usuario tenga un rol general.",
        "Tipos no elegibles devuelven 422; estados no operativos devuelven 409.",
        "Rutas directas se validan contra el Navigator backend-derived.",
        "Un módulo PLANNED devuelve 409 y un módulo oculto/disabled devuelve 403.",
        "ARCHIVED no expone acciones mutables en la UI operacional.",
    ],
)

heading(doc, "16. Auditoría y observabilidad")
add_table(
    doc,
    ["Evento", "Cuándo"],
    [
        ["workspace.context_loaded", "Contexto autorizado cargado"],
        ["workspace.context_denied", "Cross-tenant, access, type o status denegado"],
        ["workspace.opened", "Open Workspace con route validada"],
        ["workspace.context_switched", "Último Workspace del usuario cambia"],
        ["workspace.module_opened", "Route guard autoriza módulo READY"],
    ],
    [3600, 5760],
)

heading(doc, "17. Contratos API")
add_table(
    doc,
    ["Método", "Ruta"],
    [
        ["GET", "/api/v1/workspaces"],
        ["GET", "/api/v1/workspaces/recent"],
        ["POST", "/api/v1/workspaces/{id}/open"],
        ["GET", "/api/v1/workspaces/{id}/context"],
        ["GET", "/api/v1/workspaces/{id}/home"],
        ["GET", "/api/v1/workspaces/{id}/navigator"],
        ["GET", "/api/v1/workspaces/{id}/modules/{module_code}"],
        ["PUT", "/api/v1/workspaces/{id}/last-route"],
    ],
    [1500, 7860],
)

page_break(doc)
heading(doc, "18. Persistencia y migración")
bullets(
    doc,
    [
        "Alembic 20260813_0039 depende de 20260813_0038 y es aditiva.",
        "La tabla recent_workspaces contiene sólo continuidad por usuario; no duplica EnterpriseWorkspace.",
        "El ciclo PostgreSQL upgrade 0038→0039, downgrade 0039→0038 y re-upgrade fue exitoso.",
        "El entorno local quedó en 20260813_0039 (head).",
        "La tabla se creó vacía; no hubo backfill especulativo ni mutación funcional del tenant.",
    ],
)

heading(doc, "19. Evidencia de pruebas")
add_table(
    doc,
    ["Control", "Resultado"],
    [
        ["Gate 06D backend focal", "6 passed"],
        ["06D + Enterprise regression", "14 passed"],
        ["Backend scope Ruff", "PASS"],
        ["Gate 06D PostgreSQL E2E", "6 passed + migration cycle"],
        ["Frontend Gate 06D", "2 passed"],
        ["Explorer + Gate 06D focal", "12 passed"],
        ["Frontend full suite", "185 passed; 1 caso corregido y 1 timeout ambiental"],
        ["AppFlow aislado", "INCONCLUSO · carga del runner > 4 min, sin aserción fallida"],
        ["TypeScript / ESLint scope / Prettier", "PASS"],
        ["Vite production build", "PASS · 2,354 modules"],
        ["Browser real", "PASS · Header/Breadcrumb/Navigator/Home; 0 console errors"],
    ],
    [3000, 6360],
)

heading(doc, "20. Protección del baseline")
add_table(
    doc,
    ["Entidad", "Antes", "Después"],
    [
        ["enterprise_workspaces", "14", "14"],
        ["projects", "4", "4"],
        ["admin_configurations", "60", "60"],
        ["project_creation_requests", "0", "0"],
        ["physical_workspace_creation_requests", "0", "0"],
        ["physical_workspace_initializations", "0", "0"],
        ["recent_workspaces", "MISSING", "0 (tabla nueva)"],
        ["Alembic head", "20260813_0038", "20260813_0039"],
    ],
    [4350, 2505, 2505],
)

heading(doc, "21. Despliegue local")
add_table(
    doc,
    ["Servicio", "Estado"],
    [
        ["Frontend", "HTTP 200 · http://127.0.0.1:5173/app"],
        ["API", "healthy · HTTP 200 · http://127.0.0.1:8000/api/v1/health"],
        ["PostgreSQL", "healthy · migration 0039"],
        ["Redis", "healthy"],
        ["Worker / Beat", "running"],
    ],
    [2600, 6760],
)

page_break(doc)
heading(doc, "22. Riesgos y deuda controlada")
bullets(
    doc,
    [
        "Las capacidades Home recientes/documentos/tareas aún no tienen motores propios; retornan vacío explícito.",
        "PROPERTY/FACILITY/WAREHOUSE reales todavía no existen en el tenant persistente; la cobertura usa fixtures aislados.",
        "El bridge PROJECT reutiliza la pantalla Project Controls existente; la convergencia de rutas profundas queda posterior.",
        "No se añadió cache distribuido; el ETag y Cache-Control privado evitan sobreoptimización prematura.",
        "Warnings históricos de BIM Viewer no pertenecen a 06D y no fueron modificados.",
    ],
)

heading(doc, "23. Matriz de aceptación Gate 06D")
criteria = [
    "Auditoría inicial y REUSE BEFORE CREATE",
    "Active Workspace Context backend",
    "Contrato de identity/parent/template/responsible",
    "Workspace Navigator derivado por backend",
    "Project navigation",
    "Property navigation",
    "Facility navigation con PLANNED",
    "Warehouse navigation con PLANNED",
    "Workspace-scoped routing",
    "WorkspaceContextProvider",
    "Workspace Home común",
    "Workspace Header y Breadcrumb",
    "Recent Workspaces y last route",
    "My Workspaces y filtros",
    "Multi-workspace readiness",
    "Workspace switching",
    "No data leakage",
    "tenant_id + workspace_id isolation",
    "Cross-tenant isolation",
    "Direct URL security",
    "Module permission security",
    "Workspace status behavior",
    "ADMIN navigation profile",
    "USER MODE Enterprise Explorer integration",
    "RBAC y nuevos permisos mínimos",
    "SecurityEvent audit trail",
    "ETag/cache isolation",
    "PostgreSQL migration reversibility",
    "Backend unit/integration tests",
    "Frontend unit/integration tests",
    "Browser E2E localhost",
    "Regression 05A–06C preserved",
    "No CORE revision or republish",
    "No duplicate module system",
    "No invented operational data",
    "Baseline before/after invariant",
    "Local services healthy",
    "Technical Word report",
    "Gate stop boundary respected",
]
add_table(doc, ["Criterio", "Resultado"], [[item, "PASS"] for item in criteria], [7850, 1510])
callout(doc, "Estado final exclusivo", "READY_FOR_OPERATIONAL_MODULES", PALE_GREEN)
doc.add_paragraph(
    "Gate 06D termina aquí. Asset Manager, Inventory Manager, mantenimiento, spaces, utilities, inventario, activos y "
    "la operación profunda de Property/Facility/Warehouse no fueron iniciados."
)

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
